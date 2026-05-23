try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass
"""
app.py — Companies Act 2013 Compliance CRM  v2.0
Flask REST API + JWT + RBAC
All 7 new features: Company Master PDF, MCA Credentials, DSC Records,
  Auditor Nature, All Statutory Registers, Document Template Store
"""
import os, uuid, json, io
import xlsxwriter
import openpyxl
from pathlib import Path
from datetime import datetime, date, timedelta
try:
    from flask_cors import CORS
    _CORS_AVAILABLE = True
except ImportError:
    _CORS_AVAILABLE = False
    class CORS:  # stub
        def __init__(self, *a, **kw): pass
from flask import Flask, request, jsonify, send_file, g, send_from_directory
from werkzeug.utils import secure_filename

from database import get_db, row, rows, hash_pw, init_db, write_audit_log, ensure_custom_placeholders_table, ensure_permission_tables, ensure_custom_placeholders_table
from auth import login_required, require_role, platform_admin_required, make_token, hash_pw as auth_hash, verify_pw, can, get_token, DEFAULT_PERMISSIONS, ALL_MODULES, tenant_scope, rate_limit_login
from compliance import (run_compliance_checks, generate_document, build_context,
                        extract_placeholders, get_active_entities, AUTO_FILLED,
                        generate_company_master_pdf, generate_register_pdf,
                        REGISTER_DEFINITIONS)

BASE_DIR   = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
ALLOWED_EXT = {".pdf",".doc",".docx",".jpg",".jpeg",".png"}

import os as _os
_ROOT = _os.path.dirname(_os.path.abspath(__file__))
# Try multiple candidate paths so deployment works regardless of folder structure
_FRONTEND_CANDIDATES = [
    _os.path.join(_ROOT, '..', 'frontend'),   # backend/app.py → ../frontend/
    _os.path.join(_ROOT, 'frontend'),          # app.py → frontend/  (flat layout)
    _os.path.join(_ROOT),                      # same folder as app.py
    _os.path.join(_ROOT, '..'),               # parent of app.py
]
_FRONTEND = next(
    (p for p in _FRONTEND_CANDIDATES
     if _os.path.isfile(_os.path.join(p, 'index.html'))),
    _os.path.join(_ROOT, '..', 'frontend')    # fallback (original)
)
app = Flask(__name__, static_folder=str(BASE_DIR/"static"))

@app.after_request
def _security_headers(response):
    """Add security headers to every response."""
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
    response.headers.setdefault("X-XSS-Protection", "1; mode=block")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    # Only add HSTS in production (not on localhost)
    if not app.debug:
        response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    return response

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Resource not found"}), 404

@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({"error": "Method not allowed"}), 405

@app.errorhandler(500)
def internal_error(e):
    import logging
    logging.getLogger(__name__).exception("Unhandled 500 error")
    return jsonify({"error": "Internal server error"}), 500

import os as _cors_os
_ALLOWED_ORIGINS = _cors_os.environ.get('ALLOWED_ORIGINS', '*').split(',')
CORS(app, resources={r'/api/*': {'origins': _ALLOWED_ORIGINS}})
app.config["MAX_CONTENT_LENGTH"] = 10*1024*1024

# ══ INPUT SANITIZERS (PostgreSQL strict typing safety) ════════════════════════
def _dt(v):
    """Return None for empty/None date strings; PostgreSQL DATE columns reject ''."""
    if v is None: return None
    v = str(v).strip()
    return v if v else None


def _pd(v):
    """Parse date string to date object, return None on failure."""
    if not v: return None
    try:
        from datetime import date as _d
        s = str(v)[:10]
        return _d.fromisoformat(s)
    except Exception:
        return None

def _num(v, default=0, cast=float):
    """Safe numeric conversion — returns default on empty/None/invalid."""
    try: return cast(v) if v not in (None, '', 'null') else cast(default)
    except (ValueError, TypeError): return cast(default)

def _str(v):
    """Return None for empty strings — needed for UNIQUE TEXT columns (CIN, PAN, DIN)."""
    if v is None: return None
    v = str(v).strip()
    return v if v else None

def _safe(fn, conn=None):
    """Call fn(); on any exception rollback conn and re-raise."""
    try:
        return fn()
    except Exception as _e:
        if conn:
            try: conn.rollback()
            except: pass
        raise _e

def _ds(v):
    """Date to ISO string — handles str, datetime.date, datetime.datetime, None."""
    if v is None: return None
    if isinstance(v, str): return v[:10] if len(v) >= 10 else v
    try: return v.isoformat()[:10]   # datetime.date / datetime.datetime
    except: return str(v)[:10]

def _to_date(v):
    """Parse anything to a datetime.date object, or return None."""
    if v is None: return None
    if isinstance(v, date): return v
    try: return date.fromisoformat(str(v)[:10])
    except: return None



@app.errorhandler(500)
def handle_500(e):
    import traceback as _tb
    tb = _tb.format_exc()
    app.logger.error("500 ERROR:\n" + tb)
    # Return the last meaningful line of the traceback for frontend display
    last_line = [l.strip() for l in tb.strip().splitlines() if l.strip()][-1]
    return jsonify({"error": last_line, "detail": tb[-1200:]}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    import traceback as _tb
    tb = _tb.format_exc()
    app.logger.error("UNHANDLED EXCEPTION:\n" + tb)
    last_line = [l.strip() for l in tb.strip().splitlines() if l.strip()][-1]
    return jsonify({"error": last_line, "detail": tb[-1200:]}), 500

# ══ HELPERS ══════════════════════════════════════════════════════════════════
def _kyc_due():
    today = date.today()
    yr = today.year if today.month>=4 else today.year-1
    return f"{yr+1}-09-30"

def _kyc_status(due_str):
    if not due_str: return "pending"
    d = _to_date(due_str)
    if d is None: return "pending"
    days = (d-date.today()).days
    if days < 0: return "overdue"
    if days <= 30: return "due_soon"
    return "compliant"

def _enrich_auditor(a):
    if a.get("end_date"):
        d=_to_date(a["end_date"]); days=(d-date.today()).days if d else 0
        a["days_to_expiry"]=days
        a["expiry_status"]="expired" if days<0 else ("expiring_soon" if days<=30 else "valid")
    return a

# ══ STATIC / INDEX ═══════════════════════════════════════════════════════════


def _rv(r, idx_or_key):
    """Safe row value accessor — works for both dict rows and tuple rows."""
    if isinstance(r, dict):
        if isinstance(idx_or_key, int):
            return list(r.values())[idx_or_key]
        return r.get(idx_or_key)
    return r[idx_or_key]

def _count(c):
    """Extract COUNT(*) result from either dict row or tuple row."""
    r = c.fetchone()
    if r is None: return 0
    if isinstance(r, dict): return int(list(r.values())[0] or 0)
    return int(r[0] or 0)

@app.route("/")
def index():
    import os
    from flask import Response
    fp = os.path.join(_FRONTEND, 'index.html')
    try:
        with open(fp, 'r', encoding='utf-8') as f:
            html = f.read()
        resp = Response(html, status=200, mimetype='text/html')
    except FileNotFoundError:
        resp = Response("<h1>Frontend not found</h1><p>Place index.html in the frontend/ folder.</p>", status=404, mimetype='text/html')
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"]        = "no-cache"
    resp.headers["Expires"]       = "0"
    return resp

@app.route("/static/<path:fn>")
def statics(fn): return send_from_directory(str(BASE_DIR/"static"),fn)

# ══ AUTH ═════════════════════════════════════════════════════════════════════
@app.route("/api/auth/login", methods=["POST"])
def login():
    _ip = request.headers.get("X-Forwarded-For", request.remote_addr or "?").split(",")[0].strip()
    if not rate_limit_login(_ip):
        return jsonify({"error": "Too many login attempts — try again in 1 minute"}), 429
    d=request.get_json(silent=True, force=True) or {}
    email=(d.get("email") or "").strip().lower(); pw=d.get("password") or ""
    if not email or not pw: return jsonify({"error":"Email and password required"}),400
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM users WHERE email=%s AND is_active=1",(email,))
    user=row(c.fetchone())
    if not user or not verify_pw(pw, user["password"]): conn.close(); return jsonify({"error":"Invalid credentials"}),401
    # ── Auto-upgrade legacy SHA-256 hash to bcrypt on first successful login ──
    stored = user["password"]
    if not (stored.startswith("$2b$") or stored.startswith("$2a$")):
        try:
            c.execute("UPDATE users SET password=%s WHERE id=%s", (hash_pw(pw), user["id"]))
            conn.commit()
        except Exception:
            pass
    is_pa = bool(user.get("is_platform_admin", 0))
    if not is_pa and user.get("tenant_id"):
        c.execute("SELECT status FROM tenants WHERE id=%s", (user["tenant_id"],))
        t = c.fetchone()
        if t:
            _t_status = t.get('status') if isinstance(t,dict) else t[0]
            if _t_status not in ("active",):
                conn.close(); return jsonify({"error":f"Account is {_t_status}. Contact support."}),403
    c.execute("UPDATE users SET last_login=NOW() WHERE id=%s",(user["id"],))
    conn.commit(); conn.close()
    return jsonify({
        "token": make_token(user["id"],user["role"],user["name"],user.get("tenant_id"),is_pa),
        "user":  {"id":user["id"],"name":user["name"],"email":user["email"],"role":user["role"]},
        "is_platform_admin": is_pa,
        "tenant_id": user.get("tenant_id"),
    })

@app.route("/api/auth/me")
@login_required
def me(): return jsonify({"id":g.user["id"],"name":g.user["name"],"email":g.user["email"],"role":g.user["role"],"is_platform_admin":bool(g.user.get("is_platform_admin",0)),"tenant_id":g.user.get("tenant_id")})

@app.route("/api/auth/change-password", methods=["POST"])
@login_required
def change_password():
    d=request.get_json(silent=True, force=True) or {}
    old=d.get("old_password",""); new=d.get("new_password","")
    if not old or not new or len(new)<6: return jsonify({"error":"Valid passwords required (min 6 chars)"}),400
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT password FROM users WHERE id=%s",(g.user_id,))
    r=c.fetchone()
    _rpw = r['password'] if isinstance(r,dict) else r[0]
    if not r or not verify_pw(old, _rpw): conn.close(); return jsonify({"error":"Current password incorrect"}),400
    c.execute("UPDATE users SET password=%s WHERE id=%s",(hash_pw(new),g.user_id))
    conn.commit(); conn.close(); return jsonify({"success":True})

# ══ USERS ════════════════════════════════════════════════════════════════════

@app.route("/api/auth/check-email", methods=["POST"])
def check_email():
    """Check if an email exists (for login UX feedback). Never reveals password info."""
    d = request.get_json(silent=True, force=True) or {}
    email = (d.get("email") or "").strip().lower()
    if not email:
        return jsonify({"exists": False}), 200
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT is_active FROM users WHERE email=%s", (email,))
    r = c.fetchone()
    conn.close()
    if not r:
        return jsonify({"exists": False}), 200
    is_active = (r.get("is_active") if isinstance(r, dict) else r[0])
    return jsonify({"exists": True, "active": bool(is_active)}), 200

@app.route("/api/users")
@require_role("superadmin", "manager")
def list_users():
    search = request.args.get("q","").strip()
    role_f = request.args.get("role","")
    conn = get_db(); c = conn.cursor()
    q = "SELECT id,name,email,role,is_active,created_at,last_login FROM users WHERE 1=1"
    params = []
    if g.tenant_id:
        q += " AND (tenant_id=%s OR is_platform_admin=1)"; params.append(g.tenant_id)
    if role_f:
        q += " AND role=%s"; params.append(role_f)
    if search:
        q += " AND (name ILIKE %s OR email ILIKE %s)"
        like = f"%{search}%"
        params.extend([like, like])
    q += " ORDER BY name"
    c.execute(q, params)
    result = rows(c.fetchall())
    conn.close()
    return jsonify(result)

@app.route("/api/users", methods=["POST"])
@require_role("superadmin")
def create_user():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("name"):     return jsonify({"error": "Name is required"}), 400
    if not d.get("email"):    return jsonify({"error": "Email is required"}), 400
    if not d.get("password"): return jsonify({"error": "Password is required"}), 400
    if len(d["password"]) < 6: return jsonify({"error": "Password must be at least 6 characters"}), 400
    role = d.get("role", "staff")
    # Superadmin can create any role; validate role value
    if role not in ("staff", "manager", "superadmin"):
        return jsonify({"error": "Invalid role. Must be staff, manager, or superadmin"}), 400
    uid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    try:
        c.execute(
            "INSERT INTO users (id,name,email,password,role,is_active,tenant_id) VALUES (%s,%s,%s,%s,%s,1,%s)",
            (uid, d["name"].strip(), d["email"].strip().lower(), hash_pw(d["password"]), role, g.tenant_id)
        )
        conn.commit()
    except Exception as e:
        conn.close()
        if "UNIQUE" in str(e): return jsonify({"error": "Email already exists"}), 400
        return jsonify({"error": str(e)}), 400
    c.execute("SELECT id,name,email,role,is_active,created_at FROM users WHERE id=%s", (uid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/users/<uid>", methods=["PUT"])
@require_role("superadmin")
def update_user(uid):
    d = request.get_json(silent=True, force=True) or {}
    # Prevent superadmin from editing themselves into a lower role
    if uid == g.user_id and "role" in d and d["role"] != "superadmin":
        return jsonify({"error": "Cannot change your own role"}), 400
    conn = get_db(); c = conn.cursor()
    fields = {k: d[k] for k in ["name", "email", "role", "is_active"] if k in d}
    if "password" in d and d["password"]:
        if len(d["password"]) < 6:
            return jsonify({"error": "Password must be at least 6 characters"}), 400
        fields["password"] = hash_pw(d["password"])
    if not fields: return jsonify({"error": "Nothing to update"}), 400
    c.execute(
        f"UPDATE users SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",
        list(fields.values()) + [uid]
    )
    conn.commit()
    c.execute("SELECT id,name,email,role,is_active,created_at,last_login FROM users WHERE id=%s", (uid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/users/<uid>/pending-tasks")
@require_role("superadmin")
def user_pending_tasks(uid):
    """Return count and list of pending tasks assigned to/led/managed by this user."""
    conn = get_db(); c = conn.cursor()
    c.execute("""
        SELECT t.id, t.title, t.status, t.priority, t.due_date,
               co.name AS company_name,
               CASE
                 WHEN t.task_leader=%s  THEN 'leader'
                 WHEN t.task_manager=%s THEN 'manager'
                 ELSE 'assignee'
               END AS user_role_in_task
        FROM tasks t
        LEFT JOIN companies co ON t.company_id=co.id
        WHERE (t.task_leader=%s OR t.task_manager=%s OR t.assigned_to=%s)
          AND t.status NOT IN ('completed','cancelled')
        ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END
    """, (uid, uid, uid, uid, uid))
    cols  = [d[0] for d in c.description]
    tasks = [dict(r) if isinstance(r,dict) else dict(zip(cols,r)) for r in c.fetchall()]
    conn.close()
    return jsonify({"count": len(tasks), "tasks": tasks})

@app.route("/api/users/<uid>/reassign-tasks", methods=["POST"])
@require_role("superadmin")
def reassign_tasks(uid):
    """Reassign all pending tasks from outgoing user to a replacement user."""
    d   = request.get_json(silent=True, force=True) or {}
    to  = d.get("to_user_id", "").strip()
    if not to: return jsonify({"error": "to_user_id is required"}), 400
    if to == uid: return jsonify({"error": "Cannot reassign to the same user"}), 400

    conn = get_db(); c = conn.cursor()
    # Verify target user exists
    c.execute("SELECT id, role FROM users WHERE id=%s AND is_active=1", (to,))
    target = c.fetchone()
    if not target: conn.close(); return jsonify({"error": "Target user not found"}), 404

    # Reassign each role column
    c.execute("UPDATE tasks SET task_leader=%s  WHERE task_leader=%s  AND status NOT IN ('completed','cancelled')", (to, uid))
    leader_cnt = c.rowcount
    c.execute("UPDATE tasks SET task_manager=%s WHERE task_manager=%s AND status NOT IN ('completed','cancelled')", (to, uid))
    manager_cnt = c.rowcount
    c.execute("UPDATE tasks SET assigned_to=%s  WHERE assigned_to=%s  AND status NOT IN ('completed','cancelled')", (to, uid))
    assignee_cnt = c.rowcount
    conn.commit(); conn.close()

    return jsonify({
        "success":  True,
        "reassigned": leader_cnt + manager_cnt + assignee_cnt,
        "leader":   leader_cnt,
        "manager":  manager_cnt,
        "assignee": assignee_cnt,
    })

@app.route("/api/users/<uid>", methods=["DELETE"])
@require_role("superadmin")
def delete_user(uid):
    if uid == g.user_id:
        return jsonify({"error": "Cannot delete your own account"}), 400
    conn = get_db(); c = conn.cursor()
    # Check for pending tasks
    c.execute("""SELECT COUNT(*) FROM tasks
                 WHERE (task_leader=%s OR task_manager=%s OR assigned_to=%s)
                 AND status NOT IN ('completed','cancelled')""", (uid,uid,uid))
    pending = _count(c)
    if pending > 0:
        conn.close()
        return jsonify({"error": f"User has {pending} pending task(s). Reassign them first.", "pending": pending}), 409
    c.execute("DELETE FROM users WHERE id=%s", (uid,))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/users/<uid>/reset-password", methods=["POST"])
@require_role("superadmin")
def reset_user_password(uid):
    d = request.get_json(silent=True, force=True) or {}
    new_pw = d.get("password", "")
    if len(new_pw) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("UPDATE users SET password=%s WHERE id=%s", (hash_pw(new_pw), uid))
    conn.commit(); conn.close()
    return jsonify({"success": True})

# ══ DASHBOARD ════════════════════════════════════════════════════════════════
@app.route("/api/dashboard")
@login_required
def dashboard():
    run_compliance_checks()
    # Refresh board meeting compliance alerts on every dashboard load
    try:
        sync_board_meeting_alerts(tenant_id=getattr(g, "tenant_id", None))
    except Exception:
        pass  # Never block dashboard for alert sync failures
    uid  = g.user_id
    role = g.role
    tid  = g.tenant_id          # ← tenant scope
    conn = get_db(); c = conn.cursor()
    today = date.today()

    # ── Tenant scope ──────────────────────────────────────────────────────────
    # All queries scoped to current tenant
    if tid:
        t_scope   = "AND COALESCE(t.tenant_id,%s)=%s"
        t_params  = [tid, tid]
        co_scope  = "AND COALESCE(co.tenant_id,%s)=%s"
        co_params = [tid, tid]
        a_scope   = "AND COALESCE(a.tenant_id,%s)=%s"
        a_params  = [tid, tid]
        d_scope   = "AND COALESCE(d.tenant_id,%s)=%s"
        d_params  = [tid, tid]
        dsc_scope = "AND COALESCE(dr.tenant_id,%s)=%s"
        dsc_params= [tid, tid]
        m_scope   = "AND COALESCE(m.tenant_id,%s)=%s"
        m_params  = [tid, tid]
    else:
        t_scope = co_scope = a_scope = d_scope = dsc_scope = m_scope = ""
        t_params = co_params = a_params = d_params = dsc_params = m_params = []

    # ── Role-scoped task filter ───────────────────────────────────────────────
    if role == "superadmin":
        task_scope  = "1=1"
        task_params = []
        alert_scope = "1=1"
        alert_params= []
    elif role == "manager":
        task_scope  = "t.task_manager = %s"
        task_params = [uid]
        alert_scope = "a.company_id IN (SELECT DISTINCT company_id FROM tasks WHERE task_manager=%s AND status NOT IN ('completed','cancelled'))"
        alert_params= [uid]
    else:
        task_scope  = "t.assigned_to = %s"
        task_params = [uid]
        alert_scope = "a.company_id IN (SELECT DISTINCT company_id FROM tasks WHERE assigned_to=%s AND status NOT IN ('completed','cancelled'))"
        alert_params= [uid]

    # ── Stats (tenant-scoped) ─────────────────────────────────────────────────
    c.execute(f"SELECT COUNT(*) FROM companies co WHERE co.status='active' {co_scope}", co_params)
    total_co  = _count(c)
    c.execute(f"SELECT COUNT(*) FROM directors d WHERE d.is_active=1 {d_scope}", d_params)
    total_dir = _count(c)
    c.execute(f"SELECT COUNT(*) FROM alerts a WHERE a.status='active' AND {alert_scope} {a_scope}",
              alert_params + a_params)
    active_alerts = _count(c)
    c.execute(f"SELECT COUNT(*) FROM alerts a WHERE a.status='active' AND a.severity='critical' AND {alert_scope} {a_scope}",
              alert_params + a_params)
    crit = _count(c)
    c.execute(f"SELECT COUNT(*) FROM tasks t WHERE t.status NOT IN ('completed','cancelled') AND {task_scope} {t_scope}",
              task_params + t_params)
    open_tasks = _count(c)
    c.execute(f"SELECT COUNT(*) FROM tasks t WHERE t.status NOT IN ('completed','cancelled') AND t.due_date<=%s AND {task_scope} {t_scope}",
              [(today+timedelta(days=7)).isoformat()] + task_params + t_params)
    due_soon = _count(c)
    c.execute(f"SELECT COUNT(*) FROM dsc_records dr WHERE dr.is_active=1 {dsc_scope}", dsc_params)
    total_dsc = _count(c)
    c.execute(f"SELECT COUNT(*) FROM dsc_records dr WHERE dr.is_active=1 AND dr.valid_to<%s {dsc_scope}",
              [today.isoformat()] + dsc_params)
    dsc_exp = _count(c)
    c.execute(f"SELECT COUNT(*) FROM dsc_records dr WHERE dr.is_active=1 AND dr.valid_to BETWEEN %s AND %s {dsc_scope}",
              [today.isoformat(), (today+timedelta(days=30)).isoformat()] + dsc_params)
    dsc_soon = _count(c)

    # ── Alert panels ──────────────────────────────────────────────────────────
    def role_alerts(entity_type, limit=5):
        c.execute(f"""SELECT a.*, co.name AS company_name FROM alerts a
                     LEFT JOIN companies co ON a.company_id=co.id
                     WHERE a.entity_type=%s AND a.status='active' AND {alert_scope} {a_scope}
                     ORDER BY a.due_date LIMIT {limit}""",
                  [entity_type] + alert_params + a_params)
        return rows(c.fetchall())

    aud_alerts = role_alerts("auditor")
    din_alerts = role_alerts("director")
    dsc_alerts = role_alerts("dsc")

    c.execute(f"""SELECT m.*,co.name as company_name FROM meetings m
                 JOIN companies co ON m.company_id=co.id
                 WHERE m.meeting_date>=%s AND m.status='scheduled' {m_scope}
                 ORDER BY m.meeting_date LIMIT 5""",
              [today.isoformat()] + m_params)
    meetings = rows(c.fetchall())

    c.execute(f"""SELECT t.*, co.name AS company_name FROM tasks t
                 LEFT JOIN companies co ON t.company_id=co.id
                 WHERE t.status NOT IN ('completed','cancelled') AND {task_scope} {t_scope}
                 ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2
                          WHEN 'medium' THEN 3 ELSE 4 END, t.due_date LIMIT 8""",
              task_params + t_params)
    tasks = rows(c.fetchall())

    conn.close()
    return jsonify({
        "stats": {
            "total_companies": total_co, "total_directors": total_dir,
            "active_alerts": active_alerts, "critical_alerts": crit,
            "open_tasks": open_tasks, "due_soon_tasks": due_soon,
            "total_dsc": total_dsc, "dsc_expired": dsc_exp, "dsc_expiring_soon": dsc_soon,
        },
        "auditor_alerts":    aud_alerts,
        "din_alerts":        din_alerts,
        "director_alerts":   din_alerts,
        "dsc_alerts":        dsc_alerts,
        "upcoming_meetings": meetings,
        "pending_tasks":     tasks,
    })

# ══ COMPANIES ════════════════════════════════════════════════════════════════
@app.route("/api/companies")
@login_required
def list_companies():
    q=request.args.get("q","").strip(); conn=get_db(); c=conn.cursor()
    tid = g.tenant_id
    if q:
        if tid: c.execute("SELECT * FROM companies WHERE tenant_id=%s AND (name ILIKE %s OR cin ILIKE %s OR pan ILIKE %s) ORDER BY name",[tid,f"%{q}%",f"%{q}%",f"%{q}%"])
        else:   c.execute("SELECT * FROM companies WHERE name ILIKE %s OR cin ILIKE %s OR pan ILIKE %s ORDER BY name",[f"%{q}%",f"%{q}%",f"%{q}%"])
    else:
        if tid: c.execute("SELECT * FROM companies WHERE tenant_id=%s ORDER BY name",(tid,))
        else:   c.execute("SELECT * FROM companies ORDER BY name")
    return jsonify(rows(c.fetchall()))

@app.route("/api/companies/<cid>")
@login_required
def get_company(cid):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s",(cid,)); co=row(c.fetchone())
    if not co: return jsonify({"error":"Not found"}),404
    c.execute("""SELECT d.*,k.last_kyc_date,k.next_due_date,k.kyc_status FROM directors d
                 LEFT JOIN director_kyc k ON d.id=k.director_id WHERE d.company_id=%s AND d.is_active=1""",(cid,))
    co["directors"]=rows(c.fetchall())
    c.execute("SELECT * FROM auditors WHERE company_id=%s ORDER BY created_at DESC",(cid,))
    co["auditors"]=[_enrich_auditor(a) for a in rows(c.fetchall())]
    c.execute("SELECT * FROM shareholders WHERE company_id=%s AND is_active=1",(cid,))
    co["shareholders"]=rows(c.fetchall())
    c.execute("SELECT * FROM dsc_records WHERE company_id=%s AND is_active=1",(cid,))
    co["dsc_records"]=rows(c.fetchall())
    c.execute("SELECT * FROM meetings WHERE company_id=%s ORDER BY meeting_date DESC LIMIT 5",(cid,))
    co["recent_meetings"]=rows(c.fetchall())
    c.execute("SELECT * FROM alerts WHERE company_id=%s AND status='active' ORDER BY severity DESC",(cid,))
    co["alerts"]=rows(c.fetchall())
    conn.close(); return jsonify(co)

@app.route("/api/companies/<cid>/master-pdf")
@login_required
def company_master_pdf(cid):
    try:
        pdf=generate_company_master_pdf(cid)
        conn=get_db(); c=conn.cursor()
        c.execute("SELECT name FROM companies WHERE id=%s",(cid,)); co=row(c.fetchone()); conn.close()
        fname=f"CompanyMaster_{(co['name'] if co else cid)[:20].replace(' ','_')}.pdf"
        return send_file(io.BytesIO(pdf),mimetype="application/pdf",as_attachment=True,download_name=fname)
    except Exception as e:
        return jsonify({"error":str(e)}),400

@app.route("/api/companies", methods=["POST"])
@login_required
def create_company():
    if not can("company","create"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    if not d.get("name"): return jsonify({"error":"Company name required"}),400
    cid=str(uuid.uuid4()); tid=g.tenant_id; conn=get_db(); c=conn.cursor()
    try:
     c.execute("""INSERT INTO companies
        (id,name,cin,incorporation_date,registered_office,pan,tan,email,phone,
         authorized_capital,paid_up_capital,business_activity,company_type,roc,
         letterhead_address,letterhead_footer,created_by,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (cid,d["name"],_str(d.get("cin")),_dt(d.get("incorporation_date")),d.get("registered_office"),
         _str((d.get("pan") or "").upper()),_str((d.get("tan") or "").upper()),
         _str(d.get("email")),d.get("phone"),_num(d.get("authorized_capital")),
         _num(d.get("paid_up_capital")),d.get("business_activity"),
         d.get("company_type","Private Limited"),d.get("roc"),
         d.get("letterhead_address"),d.get("letterhead_footer"),g.user_id,tid))
     conn.commit()
    except Exception as _ex:
        conn.rollback(); conn.close()
        msg = str(_ex)
        if 'unique' in msg.lower() or 'duplicate' in msg.lower():
            return jsonify({"error": "A company with this CIN already exists"}), 409
        return jsonify({"error": f"Could not create company: {msg}"}), 400
    c.execute("SELECT * FROM companies WHERE id=%s",(cid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/companies/<cid>", methods=["PUT"])
@login_required
def update_company(cid):
    if not can("company","update"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    _raw = {k:d[k] for k in ["name","cin","incorporation_date","registered_office","pan","tan","email","phone",
             "authorized_capital","paid_up_capital","business_activity","company_type","roc","status",
             "letterhead_address","letterhead_footer"] if k in d}
    fields = {}
    for _k,_v in _raw.items():
        if _k == "incorporation_date": fields[_k] = _dt(_v)
        elif _k in ("authorized_capital","paid_up_capital"): fields[_k] = _num(_v)
        elif _k in ("cin","pan","tan"): fields[_k] = _str(_v)
        else: fields[_k] = _v
    fields["updated_at"]=datetime.utcnow().isoformat()
    conn=get_db(); c=conn.cursor()
    c.execute(f"UPDATE companies SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[cid])
    conn.commit()
    c.execute("SELECT * FROM companies WHERE id=%s",(cid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/companies/<cid>", methods=["DELETE"])
@require_role("superadmin")
def delete_company(cid):
    conn=get_db(); c=conn.cursor()
    c.execute("DELETE FROM companies WHERE id=%s",(cid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ══ DIRECTORS ════════════════════════════════════════════════════════════════
@app.route("/api/directors")
@login_required
def all_directors():
    conn=get_db(); c=conn.cursor()
    ia = request.args.get("is_active", "")
    # is_active="" → all, "1" → active only, "0" → resigned/ceased
    if ia == "1":
        where = "WHERE d.is_active=1"
    elif ia == "0":
        where = "WHERE d.is_active=0"
    else:
        where = "WHERE 1=1"   # return all (active + resigned)
    c.execute(f"""SELECT d.*,k.last_kyc_date,k.next_due_date,k.kyc_status,
                        co.name as company_name,co.cin as company_cin
                 FROM directors d LEFT JOIN director_kyc k ON d.id=k.director_id
                 LEFT JOIN companies co ON d.company_id=co.id {where} ORDER BY d.name""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/companies/<cid>/directors")
@login_required
def list_directors(cid):
    conn=get_db(); c=conn.cursor()
    c.execute("""SELECT d.*,k.last_kyc_date,k.next_due_date,k.kyc_status FROM directors d
                 LEFT JOIN director_kyc k ON d.id=k.director_id WHERE d.company_id=%s ORDER BY d.name""",(cid,))
    return jsonify(rows(c.fetchall()))


@app.route("/api/dir-kyc")
@login_required
def dir_kyc_list():
    """Return one KYC record per DIN — deduped across companies."""
    conn=get_db(); c=conn.cursor()
    c.execute("""
        SELECT d.id, d.name, d.din, d.mobile, d.email,
               MAX(k.last_kyc_date) AS last_kyc_date,
               MIN(k.next_due_date) AS next_due_date,
               k.kyc_status,
               GROUP_CONCAT(DISTINCT co.name ORDER BY co.name SEPARATOR '|||') AS company_names
        FROM directors d
        LEFT JOIN director_kyc k ON d.id=k.director_id
        LEFT JOIN companies co ON d.company_id=co.id
        WHERE d.is_active=1
        GROUP BY d.din, d.name, d.mobile, d.email
        ORDER BY d.name
    """)
    results = []
    for r in c.fetchall():
        rec = dict(r)
        cos = rec.pop('company_names','') or ''
        rec['companies'] = [c.strip() for c in cos.split('|||') if c.strip()]
        results.append(rec)
    conn.close()
    return jsonify(results)


@app.route("/api/directors", methods=["POST"])
@login_required
def create_director():
    if not can("director","create"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("name"): return jsonify({"error":"company_id and name required"}),400
    did=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    tid_dir = g.tenant_id
    c.execute("""INSERT INTO directors
        (id,company_id,name,din,pan,aadhaar,email,mobile,address,designation,
         date_of_appointment,mca_user_id,mca_password,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (did,d["company_id"],d["name"],_str(d.get("din")),_str((d.get("pan") or "").upper()),
         d.get("aadhaar"),d.get("email"),d.get("mobile"),d.get("address"),
         d.get("designation","Director"),_dt(d.get("date_of_appointment")),
         d.get("mca_user_id"),d.get("mca_password"),tid_dir))
    # mca_notes stored via SAVEPOINT — safe even if column doesn't exist yet
    if d.get("mca_notes"):
        try:
            c.execute("SAVEPOINT sp_mca_notes")
            c.execute("UPDATE directors SET mca_notes=%s WHERE id=%s", (d["mca_notes"], did))
            c.execute("RELEASE SAVEPOINT sp_mca_notes")
        except Exception:
            # Column doesn't exist yet — rollback to savepoint to keep transaction healthy
            try: c.execute("ROLLBACK TO SAVEPOINT sp_mca_notes")
            except Exception: pass
    last_kyc=_dt(d.get("last_kyc_date")); due=_kyc_due()
    c.execute("INSERT INTO director_kyc (id,director_id,last_kyc_date,next_due_date,kyc_status,tenant_id) VALUES (%s,%s,%s,%s,%s,%s)",
              (str(uuid.uuid4()),did,last_kyc,due,_kyc_status(due),g.tenant_id))
    conn.commit()
    c.execute("""SELECT d.*,k.last_kyc_date,k.next_due_date,k.kyc_status FROM directors d
                 LEFT JOIN director_kyc k ON d.id=k.director_id WHERE d.id=%s""",(did,))
    result=row(c.fetchone()); conn.close(); return jsonify(result),201

@app.route("/api/directors/<did>", methods=["PUT"])
@login_required
def update_director(did):
    if not can("director","update"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}; conn=get_db(); c=conn.cursor()
    _rdir = {k:d[k] for k in ["name","din","pan","aadhaar","email","mobile","address","designation",
             "date_of_appointment","date_of_cessation","is_active","mca_user_id","mca_password","mca_notes"] if k in d}
    fields = {}
    for _k,_v in _rdir.items():
        if _k in ("date_of_appointment","date_of_cessation"): fields[_k] = _dt(_v)
        elif _k in ("din","pan"): fields[_k] = _str(_v)
        else: fields[_k] = _v
    if fields:
        c.execute(f"UPDATE directors SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[did])
    # Update KYC record if any KYC field was sent
    if "last_kyc_date" in d or "next_due_date" in d:
        # Use manually provided due date if given, else auto-calculate
        due = _dt(d["next_due_date"]) if d.get("next_due_date") else _kyc_due()
        st  = _kyc_status(due)
        last_kyc = _dt(d["last_kyc_date"]) if d.get("last_kyc_date") else None
        if last_kyc:
            c.execute("UPDATE director_kyc SET last_kyc_date=%s,next_due_date=%s,kyc_status=%s,updated_at=NOW() WHERE director_id=%s",
                      (last_kyc, due, st, did))
        else:
            # Only due date changed, preserve existing last_kyc_date
            c.execute("UPDATE director_kyc SET next_due_date=%s,kyc_status=%s,updated_at=NOW() WHERE director_id=%s",
                      (due, st, did))
    conn.commit()
    c.execute("""SELECT d.*,k.last_kyc_date,k.next_due_date,k.kyc_status FROM directors d
                 LEFT JOIN director_kyc k ON d.id=k.director_id WHERE d.id=%s""",(did,))
    result=row(c.fetchone()); conn.close(); return jsonify(result)

@app.route("/api/directors/<did>", methods=["DELETE"])
@login_required
def delete_director(did):
    """Soft deactivate — keeps history."""
    if not can("director","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE directors SET is_active=0 WHERE id=%s",(did,)); conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/directors/<did>/permanent", methods=["DELETE"])
@login_required
def hard_delete_director(did):
    """Permanent delete — removes director and KYC record."""
    if not can("director","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("DELETE FROM director_kyc WHERE director_id=%s",(did,))
    c.execute("DELETE FROM directors WHERE id=%s",(did,))
    conn.commit(); conn.close()
    return jsonify({"success":True})

# ══ AUDITORS ═════════════════════════════════════════════════════════════════
@app.route("/api/auditors")
@login_required
def all_auditors():
    tid  = g.tenant_id
    conn = get_db(); c = conn.cursor()
    # Filters from frontend
    fco    = request.args.get("co", "")
    fnature= request.args.get("nature", "")
    ffrom  = request.args.get("from", "")
    fto    = request.args.get("to", "")

    sql    = """SELECT a.*,co.name as company_name,co.cin as company_cin
                FROM auditors a JOIN companies co ON a.company_id=co.id
                WHERE a.is_active=1"""
    params = []

    # Tenant isolation
    if tid:
        sql += " AND (a.tenant_id=%s OR a.tenant_id IS NULL)"; params.append(tid)

    if fco:     sql += " AND a.company_id=%s";             params.append(fco)
    if fnature: sql += " AND a.nature_of_appointment=%s";  params.append(fnature)
    if ffrom:   sql += " AND a.end_date>=%s";              params.append(ffrom)
    if fto:     sql += " AND a.end_date<=%s";              params.append(fto)

    fsearch = request.args.get("q","").strip()
    if fsearch:
        sql += " AND (a.name ILIKE %s OR a.firm_name ILIKE %s OR co.name ILIKE %s)"
        like = f"%{fsearch}%"
        params.extend([like, like, like])
    sql += " ORDER BY a.end_date"
    c.execute(sql, params)
    result = rows(c.fetchall())
    conn.close()
    return jsonify([_enrich_auditor(a) for a in result])

@app.route("/api/companies/<cid>/auditors")
@login_required
def list_auditors(cid):
    conn=get_db(); c=conn.cursor()
    tid = g.tenant_id
    if tid:
        c.execute("SELECT * FROM auditors WHERE company_id=%s AND (tenant_id=%s OR tenant_id IS NULL) ORDER BY created_at DESC",(cid,tid))
    else:
        c.execute("SELECT * FROM auditors WHERE company_id=%s ORDER BY created_at DESC",(cid,))
    result = rows(c.fetchall())
    conn.close()
    return jsonify([_enrich_auditor(a) for a in result])

@app.route("/api/auditors", methods=["POST"])
@login_required
def create_auditor():
    if not can("auditor","create"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("name"): return jsonify({"error":"company_id and name required"}),400
    aid=str(uuid.uuid4())
    end=d.get("end_date")
    if not end and d.get("start_date"):
        _sd=_to_date(d["start_date"])
        end=(_sd.replace(year=_sd.year+1).isoformat() if _sd else None)
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE auditors SET is_active=0 WHERE company_id=%s AND is_active=1",(d["company_id"],))
    tid = g.tenant_id
    c.execute("""INSERT INTO auditors
        (id,company_id,name,firm_name,membership_no,frn,pan,address,email,phone,
         appointment_date,nature_of_appointment,appointment_type,start_date,end_date,srn_adt1,notes,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (aid,d["company_id"],d["name"],d.get("firm_name"),d.get("membership_no"),d.get("frn"),
         _str((d.get("pan") or "").upper()),d.get("address"),d.get("email"),d.get("phone"),
         _dt(d.get("appointment_date")),d.get("nature_of_appointment","Regular Auditor"),
         d.get("appointment_type","AGM Appointment"),_dt(d.get("start_date")),_dt(end),d.get("srn_adt1"),d.get("notes"),tid))
    conn.commit()
    c.execute("SELECT * FROM auditors WHERE id=%s",(aid,)); result=row(c.fetchone()); conn.close()
    return jsonify(_enrich_auditor(result)),201

@app.route("/api/auditors/<aid>", methods=["PUT"])
@login_required
def update_auditor(aid):
    if not can("auditor","update"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    _raud = {k:d[k] for k in ["name","firm_name","membership_no","frn","pan","address","email","phone",
             "appointment_date","nature_of_appointment","appointment_type","start_date","end_date",
             "srn_adt1","is_active","notes"] if k in d}
    fields = {}
    for _k,_v in _raud.items():
        if _k in ("appointment_date","start_date","end_date"): fields[_k] = _dt(_v)
        elif _k == "pan": fields[_k] = _str(_v)
        else: fields[_k] = _v
    if not fields: return jsonify({"error":"Nothing to update"}),400
    conn=get_db(); c=conn.cursor()
    c.execute(f"UPDATE auditors SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[aid])
    conn.commit()
    c.execute("SELECT * FROM auditors WHERE id=%s",(aid,)); result=row(c.fetchone()); conn.close()
    return jsonify(_enrich_auditor(result))

@app.route("/api/auditors/<aid>", methods=["DELETE"])
@login_required
def delete_auditor(aid):
    """Soft deactivate — keeps history."""
    if not can("auditor","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE auditors SET is_active=0 WHERE id=%s",(aid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/auditors/<aid>/permanent", methods=["DELETE"])
@login_required
def hard_delete_auditor(aid):
    """Permanent delete — removes record entirely."""
    if not can("auditor","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("DELETE FROM auditors WHERE id=%s",(aid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/auditors/<aid>/upload", methods=["POST"])
@login_required
def upload_adt1(aid):
    if "file" not in request.files: return jsonify({"error":"No file"}),400
    f=request.files["file"]
    if not f.filename: return jsonify({"error":"No file selected"}),400
    ext=Path(f.filename).suffix.lower()
    if ext not in ALLOWED_EXT: return jsonify({"error":f"Type {ext} not allowed"}),400
    fname=f"adt1_{aid}{ext}"; f.save(str(UPLOAD_DIR/fname))
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE auditors SET adt1_file=%s WHERE id=%s",(fname,aid)); conn.commit(); conn.close()
    return jsonify({"success":True,"file":fname})

# ══ SHAREHOLDERS ═════════════════════════════════════════════════════════════
@app.route("/api/companies/<cid>/shareholders")
@login_required
def list_shareholders(cid):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM shareholders WHERE company_id=%s AND is_active=1 ORDER BY name",(cid,))
    return jsonify(rows(c.fetchall()))

@app.route("/api/shareholders", methods=["POST"])
@login_required
def create_shareholder():
    if not can("shareholder","create"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}; sid=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    c.execute("""INSERT INTO shareholders
        (id,company_id,name,folio_no,pan,email,mobile,address,share_class,shares_held,face_value,
         date_of_entry,mca_user_id,mca_password,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (sid,d["company_id"],d["name"],d.get("folio_no"),_str((d.get("pan") or "").upper()),
         d.get("email"),d.get("mobile"),d.get("address"),d.get("share_class","Equity"),
         _num(d.get("shares_held",0),cast=int),_num(d.get("face_value",10),default=10),_dt(d.get("date_of_entry")),
         d.get("mca_user_id"),d.get("mca_password"),g.tenant_id))
    conn.commit()
    c.execute("SELECT * FROM shareholders WHERE id=%s",(sid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/shareholders/<sid>", methods=["PUT"])
@login_required
def update_shareholder(sid):
    if not can("shareholder","update"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    fields={k:d[k] for k in ["name","folio_no","pan","email","mobile","address","share_class",
             "shares_held","face_value","is_active","mca_user_id","mca_password"] if k in d}
    conn=get_db(); c=conn.cursor()
    c.execute(f"UPDATE shareholders SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[sid])
    conn.commit(); c.execute("SELECT * FROM shareholders WHERE id=%s",(sid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/shareholders/<sid>", methods=["DELETE"])
@login_required
def delete_shareholder(sid):
    """Soft deactivate."""
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE shareholders SET is_active=0 WHERE id=%s",(sid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/shareholders/<sid>/permanent", methods=["DELETE"])
@login_required
def hard_delete_shareholder(sid):
    """Permanent delete."""
    if not can("shareholder","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("DELETE FROM shareholders WHERE id=%s",(sid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ══ DSC RECORDS ══════════════════════════════════════════════════════════════
@app.route("/api/dsc")
@login_required
def all_dsc():
    tid     = g.tenant_id
    cid     = request.args.get("company_id", "") or request.args.get("co", "")
    custody = request.args.get("custody", "")
    ffrom   = request.args.get("from", "")
    fto     = request.args.get("to", "")

    conn = get_db(); c = conn.cursor()
    sql    = """SELECT d.*,co.name as company_name FROM dsc_records d
                LEFT JOIN companies co ON d.company_id=co.id
                WHERE d.is_active=1"""
    params = []

    # Tenant isolation
    if tid:
        sql += " AND (d.tenant_id=%s OR d.tenant_id IS NULL)"; params.append(tid)

    if cid:     sql += " AND d.company_id=%s";      params.append(cid)
    if custody: sql += " AND d.custody_status=%s";  params.append(custody)
    if ffrom:   sql += " AND d.valid_to>=%s";        params.append(ffrom)
    if fto:     sql += " AND d.valid_to<=%s";        params.append(fto)

    fsearch = request.args.get("q","").strip()
    if fsearch:
        sql += " AND (d.holder_name ILIKE %s OR co.name ILIKE %s OR d.issued_by ILIKE %s)"
        like = f"%{fsearch}%"
        params.extend([like, like, like])
    sql += " ORDER BY d.valid_to"
    c.execute(sql, params)
    result = rows(c.fetchall())
    conn.close()

    today = date.today()
    for r in result:
        if r.get("valid_to"):
            try:
                days = (_to_date(r["valid_to"]) - today).days if _to_date(r["valid_to"]) else None
                r["days_to_expiry"] = days
                r["expiry_status"]  = "expired" if days < 0 else ("expiring_soon" if days <= 30 else "valid")
            except (ValueError, TypeError):
                r["days_to_expiry"] = None
                r["expiry_status"]  = "unknown"
    return jsonify(result)

@app.route("/api/dsc", methods=["POST"])
@login_required
def create_dsc():
    if not can("dsc","create"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    if not d.get("holder_name"): return jsonify({"error":"holder_name is required"}),400
    if not d.get("valid_to"):    return jsonify({"error":"valid_to (expiry date) is required"}),400
    dsc_id=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    tid = g.tenant_id
    c.execute("""INSERT INTO dsc_records
        (id,company_id,director_id,holder_name,holder_type,dsc_class,issued_by,
         valid_from,valid_to,token_type,custody_status,custody_date,custody_notes,notes,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (dsc_id,_str(d.get("company_id")),_str(d.get("director_id")),d["holder_name"],
         d.get("holder_type","Director"),d.get("dsc_class","Class 3"),d.get("issued_by"),
         _dt(d.get("valid_from")),_dt(d.get("valid_to")),d.get("token_type"),
         d.get("custody_status","With Client"),_dt(d.get("custody_date")),
         d.get("custody_notes"),d.get("notes"),tid))
    conn.commit()
    c.execute("SELECT * FROM dsc_records WHERE id=%s",(dsc_id,)); result=row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/dsc/<dsc_id>", methods=["PUT"])
@login_required
def update_dsc(dsc_id):
    if not can("dsc","update"): return jsonify({"error":"Insufficient permissions"}),403
    d=request.get_json(silent=True, force=True) or {}
    _rdsc = {k:d[k] for k in ["holder_name","holder_type","dsc_class","issued_by","valid_from","valid_to",
             "token_type","custody_status","custody_date","custody_notes","is_active","notes"] if k in d}
    fields = {}
    for _k,_v in _rdsc.items():
        if _k in ("valid_from","valid_to","custody_date"): fields[_k] = _dt(_v)
        else: fields[_k] = _v
    conn=get_db(); c=conn.cursor()
    if fields:
        c.execute(f"UPDATE dsc_records SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[dsc_id])
    # Log custody change
    if "custody_status" in d:
        c.execute("SELECT custody_status FROM dsc_records WHERE id=%s",(dsc_id,))
        old=c.fetchone()
        c.execute("""INSERT INTO dsc_custody_log (id,dsc_id,action,action_date,from_party,to_party,notes,recorded_by)
                     VALUES (%s,%s,%s,%s,%s,%s,%s,%s)""",
                  (str(uuid.uuid4()),dsc_id,"custody_change",date.today().isoformat(),
                   (old.get('custody_status') if isinstance(old,dict) else (old[0] if old else '')),d["custody_status"],d.get("custody_notes",""),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM dsc_records WHERE id=%s",(dsc_id,)); result=row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/dsc/<dsc_id>", methods=["DELETE"])
@login_required
def delete_dsc(dsc_id):
    """Soft deactivate."""
    if not can("dsc","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE dsc_records SET is_active=0 WHERE id=%s",(dsc_id,)); conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/dsc/<dsc_id>/permanent", methods=["DELETE"])
@login_required
def hard_delete_dsc(dsc_id):
    """Permanent delete."""
    if not can("dsc","delete"): return jsonify({"error":"Insufficient permissions"}),403
    conn=get_db(); c=conn.cursor()
    c.execute("DELETE FROM dsc_custody_log WHERE dsc_id=%s",(dsc_id,))
    c.execute("DELETE FROM dsc_records WHERE id=%s",(dsc_id,))
    conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/dsc/<dsc_id>/custody-log")
@login_required
def dsc_custody_log(dsc_id):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT l.*,u.name as recorded_by_name FROM dsc_custody_log l LEFT JOIN users u ON l.recorded_by=u.id WHERE l.dsc_id=%s ORDER BY l.created_at DESC",(dsc_id,))
    return jsonify(rows(c.fetchall()))

# ══ MEETINGS ═════════════════════════════════════════════════════════════════
@app.route("/api/meetings")
@login_required
def list_meetings():
    cid     = request.args.get("company_id", "")
    mtype   = request.args.get("meeting_type", "")
    status  = request.args.get("status", "")
    search  = request.args.get("q", "").strip()
    dfrom   = request.args.get("date_from", "")
    dto     = request.args.get("date_to", "")
    conn = get_db(); c = conn.cursor()

    q = "SELECT m.*, co.name as company_name FROM meetings m JOIN companies co ON m.company_id=co.id WHERE 1=1"
    params = []

    # Tenant isolation
    if g.tenant_id:
        q += " AND m.tenant_id=%s"; params.append(g.tenant_id)

    if cid:    q += " AND m.company_id=%s";   params.append(cid)
    if mtype:  q += " AND m.meeting_type=%s"; params.append(mtype)
    if status: q += " AND m.status=%s";        params.append(status)
    if dfrom:  q += " AND m.meeting_date>=%s"; params.append(dfrom)
    if dto:    q += " AND m.meeting_date<=%s"; params.append(dto)
    if search:
        q += " AND (co.name ILIKE %s OR m.meeting_no ILIKE %s OR m.venue ILIKE %s OR m.meeting_type ILIKE %s)"
        like = f"%{search}%"
        params.extend([like, like, like, like])

    q += " ORDER BY m.meeting_date DESC"
    c.execute(q, params)
    result = rows(c.fetchall())
    conn.close()
    return jsonify(result)

@app.route("/api/meetings", methods=["POST"])
@login_required
def create_meeting():
    d=request.get_json(silent=True, force=True) or {}; mid=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    c.execute("""INSERT INTO meetings (id,company_id,meeting_type,meeting_no,meeting_date,meeting_time,venue,agenda,status,created_by,tenant_id)
                 VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (mid,d["company_id"],d.get("meeting_type","Board"),d.get("meeting_no"),
         _dt(d["meeting_date"]),d.get("meeting_time",""),d.get("venue",""),d.get("agenda",""),
         d.get("status","scheduled"),g.user_id,g.tenant_id))
    conn.commit()
    c.execute("SELECT m.*,co.name as company_name FROM meetings m JOIN companies co ON m.company_id=co.id WHERE m.id=%s",(mid,))
    result=row(c.fetchone()); conn.close()
    try: sync_board_meeting_alerts(company_id=d.get("company_id"))
    except Exception: pass
    return jsonify(result),201

@app.route("/api/meetings/<mid>", methods=["PUT"])
@login_required
def update_meeting(mid):
    d=request.get_json(silent=True, force=True) or {}
    _rmtg = {k:d[k] for k in ["meeting_type","meeting_no","meeting_date","meeting_time","venue","agenda","notes","minutes_drafted","status"] if k in d}
    fields = {_k:(_dt(_v) if _k=="meeting_date" else _v) for _k,_v in _rmtg.items()}
    conn=get_db(); c=conn.cursor()
    c.execute(f"UPDATE meetings SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[mid])
    conn.commit()
    c.execute("SELECT m.*,co.name as company_name FROM meetings m JOIN companies co ON m.company_id=co.id WHERE m.id=%s",(mid,))
    result=row(c.fetchone()); conn.close()
    try: sync_board_meeting_alerts(company_id=(result or {}).get("company_id"))
    except Exception: pass
    return jsonify(result)

@app.route("/api/meetings/<mid>", methods=["DELETE"])
@login_required
def delete_meeting(mid):
    conn=get_db(); c=conn.cursor(); c.execute("DELETE FROM meetings WHERE id=%s",(mid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ══ STATUTORY REGISTERS ══════════════════════════════════════════════════════
@app.route("/api/registers/types")
@login_required
def register_types():
    return jsonify([{"key":k,"name":v["name"],"section":v["section"]} for k,v in REGISTER_DEFINITIONS.items()])

@app.route("/api/registers/<reg_type>/<cid>")
@login_required
def get_register(reg_type, cid):
    reg=REGISTER_DEFINITIONS.get(reg_type)
    if not reg: return jsonify({"error":f"Unknown register: {reg_type}"}),400
    conn=get_db(); c=conn.cursor()
    # Run the display query (column-based)
    c.execute(reg["query"],(cid,))
    rows_raw=c.fetchall()
    data=[(list(r.values()) if isinstance(r, dict) else list(r)) for r in rows_raw]
    # Convert date/Decimal values to strings for JSON serialization
    from datetime import date as _ddate, datetime as _ddt
    def _cell(v):
        if v is None: return ""
        if isinstance(v, (_ddate, _ddt)): return v.isoformat()[:10]
        try:
            from decimal import Decimal
            if isinstance(v, Decimal): return float(v)
        except ImportError: pass
        return v
    data = [[_cell(v) for v in row_vals] for row_vals in data]
    # Also fetch IDs for deletable registers
    id_queries = {
        "MBP-1": "SELECT id FROM director_interests WHERE company_id=%s ORDER BY date_of_disclosure",
        "CHG-1": "SELECT id FROM charges WHERE company_id=%s ORDER BY date_of_creation",
        "MBP-3": "SELECT id FROM investments WHERE company_id=%s ORDER BY date_of_investment",
        "MBP-2": "SELECT id FROM loans_guarantees WHERE company_id=%s ORDER BY date_of_transaction",
        "RPT-188": "SELECT id FROM related_party_transactions WHERE company_id=%s ORDER BY date_of_transaction",
        "SH-6": "SELECT id FROM esop_grants WHERE company_id=%s ORDER BY grant_date",
    }
    ids=[]
    if reg_type in id_queries:
        c.execute(id_queries[reg_type],(cid,))
        ids=[(r['id'] if isinstance(r,dict) and 'id' in r else list(r.values())[0] if isinstance(r,dict) else r[0]) for r in c.fetchall()]
    c.execute("SELECT name,cin FROM companies WHERE id=%s",(cid,))
    co=row(c.fetchone()); conn.close()
    return jsonify({"register":reg["name"],"section":reg["section"],"company":co,
                    "columns":reg["columns"],"data":data,"_ids":ids})

@app.route("/api/registers/<reg_type>/<cid>/pdf")
@login_required
def download_register_pdf(reg_type, cid):
    try:
        pdf=generate_register_pdf(cid,reg_type)
        reg=REGISTER_DEFINITIONS.get(reg_type,{})
        fname=f"{reg.get('name','Register').replace(' ','_')}_{cid[:8]}.pdf"
        return send_file(io.BytesIO(pdf),mimetype="application/pdf",as_attachment=True,download_name=fname)
    except Exception as e:
        return jsonify({"error":str(e)}),400

# ══ DOCUMENT TEMPLATES ═══════════════════════════════════════════════════════
@app.route("/api/document-templates")
@login_required
def list_doc_templates():
    cat=request.args.get("category",""); conn=get_db(); c=conn.cursor()
    q="SELECT id,name,category,description,placeholders,is_system,is_active,created_at FROM document_templates WHERE is_active=1"
    params=[]
    if cat: q+=" AND category=%s"; params.append(cat)
    _tpl_q = request.args.get("q","").strip()
    if _tpl_q:
        q += " AND (name ILIKE %s OR description ILIKE %s)"
        _like = f"%{_tpl_q}%"
        params.extend([_like, _like])
    q+=" ORDER BY is_system DESC,name"
    c.execute(q,params); return jsonify(rows(c.fetchall()))

@app.route("/api/document-templates/<tid>")
@login_required
def get_doc_template(tid):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM document_templates WHERE id=%s",(tid,))
    t=row(c.fetchone()); conn.close()
    if not t: return jsonify({"error":"Not found"}),404
    return jsonify(t)

@app.route("/api/document-templates", methods=["POST"])
@login_required
def create_doc_template():
    d=request.get_json(silent=True, force=True) or {}
    if not d.get("name") or not d.get("template_body"): return jsonify({"error":"name and template_body required"}),400
    import re
    placeholders=list(set(re.findall(r'\{\{([^}]+)\}\}',d["template_body"])))
    tid=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    c.execute("""INSERT INTO document_templates (id,name,category,description,template_body,placeholders,is_system,created_by,tenant_id)
                 VALUES (%s,%s,%s,%s,%s,%s,0,%s,%s)""",
        (tid,d["name"],d.get("category","resolution"),d.get("description",""),
         d["template_body"],json.dumps(placeholders),g.user_id,g.tenant_id))
    conn.commit()
    c.execute("SELECT * FROM document_templates WHERE id=%s",(tid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/document-templates/<tid>", methods=["PUT"])
@login_required
def update_doc_template(tid):
    d=request.get_json(silent=True, force=True) or {}; conn=get_db(); c=conn.cursor()
    c.execute("SELECT is_system FROM document_templates WHERE id=%s",(tid,))
    t=c.fetchone()
    _is_sys = (t.get('is_system') if isinstance(t,dict) else t[0]) if t else 0
    if _is_sys and g.role!="superadmin": conn.close(); return jsonify({"error":"System templates can only be edited by superadmin"}),403
    import re
    fields={k:d[k] for k in ["name","category","description","template_body","is_active"] if k in d}
    if "template_body" in d: fields["placeholders"]=json.dumps(list(set(re.findall(r'\{\{([^}]+)\}\}',d["template_body"]))))
    fields["updated_at"]=datetime.utcnow().isoformat()
    c.execute(f"UPDATE document_templates SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[tid])
    conn.commit()
    c.execute("SELECT * FROM document_templates WHERE id=%s",(tid,)); result=row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/document-templates/<tid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_doc_template(tid):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT is_system FROM document_templates WHERE id=%s",(tid,))
    t=c.fetchone()
    _is_sys2 = (t.get('is_system') if isinstance(t,dict) else t[0]) if t else 0
    if _is_sys2: conn.close(); return jsonify({"error":"Cannot delete system templates"}),400
    c.execute("UPDATE document_templates SET is_active=0 WHERE id=%s",(tid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── Document Generation from template ────────────────────────────────────────
@app.route("/api/companies/<cid>/active-entities")
@login_required
def active_entities(cid):
    """Return only active directors, current auditors, active shareholders for document generation."""
    try:
        data = get_active_entities(cid)
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route("/api/templates/placeholders", methods=["GET","POST"])
@login_required
def get_template_placeholders():
    """Parse placeholders from a template and classify auto-filled vs manual."""
    d = request.get_json(silent=True, force=True) or {}
    tid = d.get("template_id")
    if not tid: return jsonify({"error": "template_id required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT template_body, placeholders FROM document_templates WHERE id=%s", (tid,))
    tmpl = row(c.fetchone()); conn.close()
    if not tmpl: return jsonify({"error": "Template not found"}), 404
    all_phs = extract_placeholders(tmpl["template_body"])
    auto = [p for p in all_phs if p in AUTO_FILLED]
    manual = [p for p in all_phs if p not in AUTO_FILLED]
    return jsonify({"all": all_phs, "auto_filled": auto, "manual": manual})

@app.route("/api/documents/generate", methods=["POST"])
@login_required
def gen_document():
    d = request.get_json(silent=True, force=True) or {}
    template_id = d.get("template_id"); company_id = d.get("company_id")
    if not template_id or not company_id:
        return jsonify({"error": "template_id and company_id required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM document_templates WHERE id=%s", (template_id,))
    tmpl = row(c.fetchone()); conn.close()
    if not tmpl: return jsonify({"error": "Template not found"}), 404
    try:
        ctx = build_context(
            company_id,
            extra=d.get("extra_context", {}),
            director_id=d.get("director_id"),
            director2_id=d.get("director2_id"),
            auditor_id=d.get("auditor_id"),
        )
        content = generate_document(tmpl["template_body"], ctx)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    all_phs = extract_placeholders(tmpl["template_body"])
    doc_id = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("INSERT INTO documents (id,company_id,template_id,doc_type,doc_name,content,module,generated_by) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
              (doc_id, company_id, template_id, tmpl["category"], tmpl["name"], content, "document_engine", g.user_id))
    conn.commit(); conn.close()
    return jsonify({"id": doc_id, "template_name": tmpl["name"], "category": tmpl["category"],
                    "content": content, "context": ctx,
                    "placeholders": all_phs,
                    "auto_filled": [p for p in all_phs if p in AUTO_FILLED],
                    "manual": [p for p in all_phs if p not in AUTO_FILLED]})

@app.route("/api/documents/generate/pdf", methods=["POST"])
@login_required
def gen_document_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

    d=request.get_json(silent=True, force=True) or {}
    template_id=d.get("template_id"); company_id=d.get("company_id")
    if not template_id or not company_id: return jsonify({"error":"template_id and company_id required"}),400
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM document_templates WHERE id=%s",(template_id,))
    tmpl=row(c.fetchone()); conn.close()
    if not tmpl: return jsonify({"error":"Template not found"}),404
    ctx = build_context(
        company_id,
        extra=d.get("extra_context", {}),
        director_id=d.get("director_id"),
        director2_id=d.get("director2_id"),
        auditor_id=d.get("auditor_id"),
    )
    content = generate_document(tmpl["template_body"], ctx)

    # Get company for letterhead
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s",(company_id,)); co=row(c.fetchone()); conn.close()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=0.8*inch, bottomMargin=0.8*inch,
                            leftMargin=1.0*inch, rightMargin=1.0*inch)
    NAVY = colors.HexColor("#0f2d5c")
    BLUE = colors.HexColor("#1a56db")
    GREY = colors.HexColor("#64748b")

    def sty(n, sz=9.5, bold=False, col=None, align=TA_LEFT, leading=None):
        return ParagraphStyle(n,
            fontName   = "Helvetica-Bold" if bold else "Helvetica",
            fontSize   = sz,
            textColor  = col or colors.HexColor("#1e293b"),
            alignment  = align,
            leading    = leading or sz * 1.55,
            spaceAfter = sz * 0.4)

    story = []

    # ── Letterhead ─────────────────────────────────────────────────────────
    if co:
        story.append(Paragraph(co["name"].upper(), sty("ln", 15, True, NAVY, TA_CENTER)))
        lh = co.get("letterhead_address") or co.get("registered_office", "")
        if lh:
            for lh_part in [p.strip() for p in lh.replace(" | ","|").split("|") if p.strip()]:
                story.append(Paragraph(lh_part, sty("la", 8, False, GREY, TA_CENTER)))
        # Email and phone on same line
        _co_email = co.get("email") or ""
        _co_phone = co.get("phone") or ""
        _contact_parts = []
        if _co_email: _contact_parts.append(f"✉ {_co_email}")
        if _co_phone: _contact_parts.append(f"✆ {_co_phone}")
        if _contact_parts:
            story.append(Paragraph("  |  ".join(_contact_parts), sty("lc", 7.5, False, GREY, TA_CENTER)))
        story.append(Spacer(1, 5))
        story.append(HRFlowable(width="100%", thickness=2, color=BLUE, spaceAfter=2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=NAVY, spaceAfter=12))

    # ── Parse HTML content from Quill editor ──────────────────────────────
    def _quill_to_story(html_content, story_out):
        """Convert Quill-editor HTML to ReportLab story elements."""
        try:
            from bs4 import BeautifulSoup, NavigableString
        except ImportError:
            # Fallback: basic regex strip if bs4 not installed
            import re
            plain = re.sub(r'<[^>]+>', '', html_content)
            plain = plain.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
            for line in plain.split('\n'):
                s = line.strip()
                if s:
                    story_out.append(Paragraph(s, sty("fb")))
                else:
                    story_out.append(Spacer(1, 4))
            return

        soup = BeautifulSoup(html_content, "html.parser")

        # Quill wraps content in <p> tags with alignment classes
        ALIGN_MAP = {
            "ql-align-center":  TA_CENTER,
            "ql-align-right":   TA_RIGHT,
            "ql-align-justify": TA_JUSTIFY,
            "ql-align-left":    TA_LEFT,
        }

        def node_to_rl_text(node):
            """Recursively convert a BS4 node to ReportLab-safe XML."""
            if isinstance(node, NavigableString):
                t = str(node).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                return t
            tag = node.name
            inner = "".join(node_to_rl_text(c) for c in node.children)
            if not inner.strip():
                return "&nbsp;"
            if tag in ("strong", "b"):
                return f"<b>{inner}</b>"
            if tag in ("em", "i"):
                return f"<i>{inner}</i>"
            if tag == "u":
                return f"<u>{inner}</u>"
            if tag == "br":
                return "<br/>"
            if tag in ("span", "a", "div"):
                return inner
            return inner

        for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "li", "br"]):
            if tag.name == "br":
                story_out.append(Spacer(1, 4))
                continue

            # Determine alignment
            classes = tag.get("class") or []
            align = TA_LEFT
            for cls in classes:
                if cls in ALIGN_MAP:
                    align = ALIGN_MAP[cls]
                    break

            # Build inner text
            inner_parts = []
            for child in tag.children:
                inner_parts.append(node_to_rl_text(child))
            inner_text = "".join(inner_parts).strip()

            if not inner_text or inner_text in ("&nbsp;", " "):
                story_out.append(Spacer(1, 4))
                continue

            # Heading styles
            if tag.name in ("h1", "h2", "h3"):
                sz = {"h1": 13, "h2": 11, "h3": 10}.get(tag.name, 10)
                story_out.append(Paragraph(inner_text, sty(f"hd_{tag.name}", sz, True, NAVY, align)))
            else:
                story_out.append(Paragraph(inner_text, sty("body", 9.5, False, None, align)))

    _quill_to_story(content, story)

    # ── Footer ─────────────────────────────────────────────────────────────
    if co:
        story.append(Spacer(1, 16))
        story.append(HRFlowable(width="100%", thickness=0.5, color=GREY))
        co_pan   = co.get("pan") or ""
        co_cin   = co.get("cin") or ""
        co_off   = co.get("registered_office") or ""
        co_email = co.get("email") or ""
        co_phone = co.get("phone") or ""
        lf = co.get("letterhead_footer") or " | ".join(
            p for p in [
                f"CIN: {co_cin}"     if co_cin   else "",
                f"PAN: {co_pan}"     if co_pan   else "",
                f"✉ {co_email}"     if co_email  else "",
                f"✆ {co_phone}"     if co_phone  else "",
                co_off,
            ] if p
        )
        story.append(Paragraph(lf, sty("ft", 7, False, GREY, TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    fname = f"{tmpl['name'].replace(' ', '_')[:40]}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)


@app.route("/api/documents/generate/docx", methods=["POST"])
@login_required
def gen_document_docx():
    """Generate a formatted Word (.docx) document from a template."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from bs4 import BeautifulSoup, NavigableString
        import copy
    except ImportError as ie:
        return jsonify({"error": f"Missing library: {ie}. Run: pip install python-docx beautifulsoup4"}), 500

    d = request.get_json(silent=True, force=True) or {}
    template_id = d.get("template_id")
    company_id  = d.get("company_id")
    if not template_id or not company_id:
        return jsonify({"error": "template_id and company_id required"}), 400

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM document_templates WHERE id=%s", (template_id,))
    tmpl = row(c.fetchone()); conn.close()
    if not tmpl: return jsonify({"error": "Template not found"}), 404

    ctx = build_context(company_id,
                        extra=d.get("extra_context", {}),
                        director_id=d.get("director_id"),
                        director2_id=d.get("director2_id"),
                        auditor_id=d.get("auditor_id"))
    content = generate_document(tmpl["template_body"], ctx)

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s", (company_id,)); co = row(c.fetchone()); conn.close()

    docx = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────
    for sec in docx.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    NAVY_RGB = RGBColor(0x0f, 0x2d, 0x5c)
    GREY_RGB = RGBColor(0x64, 0x74, 0x8b)
    BLACK    = RGBColor(0x1e, 0x29, 0x3b)

    def _add_hr(docx, color_hex="4472C4", width=6):
        """Add a horizontal rule paragraph."""
        p = docx.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    str(width))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Letterhead ────────────────────────────────────────────────────────
    if co:
        # Company name
        p = docx.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(co["name"].upper())
        run.bold      = True
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY_RGB

        # Address
        lh = co.get("letterhead_address") or co.get("registered_office", "")
        if lh:
            for lh_part in [x.strip() for x in lh.replace(" | ", "|").split("|") if x.strip()]:
                p2 = docx.add_paragraph(lh_part)
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(1)
                for run in p2.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = GREY_RGB

        # Email and Phone
        _em = co.get("email") or ""
        _ph = co.get("phone") or ""
        _contact = "  |  ".join(p for p in [
            (f"✉ {_em}" if _em else ""),
            (f"✆ {_ph}" if _ph else ""),
        ] if p)
        if _contact:
            cp = docx.add_paragraph(_contact)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after  = Pt(3)
            for run in cp.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = GREY_RGB

        _add_hr(docx, "1a56db", 12)  # thick blue line
        _add_hr(docx, "0f2d5c", 4)   # thin navy line

        p_sp = docx.add_paragraph()
        p_sp.paragraph_format.space_after = Pt(4)

    # ── Parse Quill HTML into Word paragraphs ──────────────────────────────
    ALIGN_MAP = {
        "ql-align-center":  WD_ALIGN_PARAGRAPH.CENTER,
        "ql-align-right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "ql-align-justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    def _apply_runs(para, node):
        """Recursively add runs from a BS4 node to a docx paragraph."""
        if isinstance(node, NavigableString):
            txt = str(node).replace("\u00a0", " ").strip()
            if txt:
                run = para.add_run(txt)
                run.font.size = Pt(10)
                run.font.color.rgb = BLACK
            return

        tag = node.name
        for child in node.children:
            _apply_runs(para, child)

        # Apply formatting retroactively to runs we just added
        new_runs = [r for r in para.runs]
        if tag in ("strong", "b"):
            for r in new_runs: r.bold = True
        if tag in ("em", "i"):
            for r in new_runs: r.italic = True
        if tag == "u":
            for r in new_runs: r.underline = True

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "li", "br"]):
        if tag.name == "br":
            docx.add_paragraph()
            continue

        # Detect alignment
        classes = tag.get("class") or []
        align = ALIGN_MAP.get(next((c for c in classes if c in ALIGN_MAP), ""), WD_ALIGN_PARAGRAPH.LEFT)

        # Collect text
        text = tag.get_text()
        if not text or not text.strip():
            p = docx.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            continue

        if tag.name in ("h1", "h2", "h3", "h4"):
            sz = {"h1": 14, "h2": 12, "h3": 11, "h4": 10}.get(tag.name, 11)
            p = docx.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(text.strip())
            run.bold = True
            run.font.size = Pt(sz)
            run.font.color.rgb = NAVY_RGB
        else:
            p = docx.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(14)
            for child in tag.children:
                _apply_runs(p, child)
            if not p.runs:
                p.add_run(text.strip())

    # ── Footer ────────────────────────────────────────────────────────────
    if co:
        _add_hr(docx, "94a3b8", 4)
        co_cin   = co.get("cin") or ""
        co_pan   = co.get("pan") or ""
        co_email = co.get("email") or ""
        co_phone = co.get("phone") or ""
        co_off   = co.get("registered_office") or ""
        ft = co.get("letterhead_footer") or " | ".join(p for p in [
            f"CIN: {co_cin}"  if co_cin   else "",
            f"PAN: {co_pan}"  if co_pan   else "",
            f"✉ {co_email}" if co_email  else "",
            f"✆ {co_phone}" if co_phone  else "",
            co_off,
        ] if p)
        fp = docx.add_paragraph(ft)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        for run in fp.runs:
            run.font.size = Pt(7)
            run.font.color.rgb = GREY_RGB

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    fname = f"{tmpl['name'].replace(' ', '_')[:40]}.docx"
    return send_file(buf,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True,
                     download_name=fname)


@app.route("/api/documents/doc/<doc_id>")
@login_required
def get_document(doc_id):
    """Fetch a single saved document for editing."""
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT d.*,t.name as template_name,t.category as template_category
                 FROM documents d
                 LEFT JOIN document_templates t ON d.template_id=t.id
                 WHERE d.id=%s""", (doc_id,))
    doc = row(c.fetchone()); conn.close()
    if not doc: return jsonify({"error": "Document not found"}), 404
    return jsonify(doc)

@app.route("/api/documents/doc/<doc_id>", methods=["PUT"])
@login_required
def update_document(doc_id):
    """Amend (update) a saved document's content and/or name."""
    d = request.get_json(silent=True, force=True) or {}
    fields = {}
    if "content"  in d: fields["content"]  = d["content"]
    if "doc_name" in d: fields["doc_name"]  = d["doc_name"]
    if not fields: return jsonify({"error": "Nothing to update"}), 400
    fields["updated_at"] = datetime.utcnow().isoformat()
    conn = get_db(); c = conn.cursor()
    # Ensure updated_at column exists (add migration if needed)
    try:
        c.execute(f"UPDATE documents SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",
                  list(fields.values()) + [doc_id])
        conn.commit()
    except Exception as ex:
        # updated_at column might not exist — retry without it
        if "updated_at" in str(ex):
            conn.rollback()
            fields.pop("updated_at", None)
            c.execute(f"UPDATE documents SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",
                      list(fields.values()) + [doc_id])
            conn.commit()
        else:
            conn.rollback(); conn.close()
            return jsonify({"error": str(ex)}), 400
    c.execute("SELECT * FROM documents WHERE id=%s", (doc_id,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/documents/doc/<doc_id>", methods=["DELETE"])
@login_required
def delete_document(doc_id):
    """Delete a saved document."""
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM documents WHERE id=%s", (doc_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/documents/<company_id>")
@login_required
def list_documents(company_id):
    conn=get_db(); c=conn.cursor()
    c.execute("""SELECT d.*,t.name as template_name FROM documents d
                 LEFT JOIN document_templates t ON d.template_id=t.id
                 WHERE d.company_id=%s ORDER BY d.created_at DESC""",(company_id,))
    return jsonify(rows(c.fetchall()))


# ══ REGISTER ENTRY CRUD ═══════════════════════════════════════════════════════
# Each register that needs data entry has full GET/POST/PUT/DELETE routes
# ─────────────────────────────────────────────────────────────────────────────

# ── Charges (CHG-1, Section 85) ──────────────────────────────────────────────
@app.route("/api/charges")
@login_required
def list_charges():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("""SELECT ch.*,co.name as company_name FROM charges ch
                     JOIN companies co ON ch.company_id=co.id WHERE ch.company_id=%s ORDER BY ch.date_of_creation DESC""",(cid,))
    else:
        c.execute("""SELECT ch.*,co.name as company_name FROM charges ch
                     JOIN companies co ON ch.company_id=co.id ORDER BY ch.date_of_creation DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/charges", methods=["POST"])
@login_required
def create_charge():
    if not can("company","create"): return jsonify({"error":"Insufficient permissions"}),403
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("charge_holder"): return jsonify({"error":"company_id and charge_holder required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO charges
        (id,company_id,charge_id,srn,date_of_creation,amount,charge_holder,assets_charged,
         charge_type,date_of_modification,date_of_satisfaction,status,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d.get("charge_id"),d.get("srn"),_dt(d.get("date_of_creation")),
         _num(d.get("amount")),d["charge_holder"],d.get("assets_charged"),
         d.get("charge_type","Hypothecation"),_dt(d.get("date_of_modification")),
         _dt(d.get("date_of_satisfaction")),d.get("status","Open"),d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM charges WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/charges/<rid>", methods=["PUT"])
@login_required
def update_charge(rid):
    if not can("company","update"): return jsonify({"error":"Insufficient permissions"}),403
    d = request.get_json(silent=True, force=True) or {}
    _fc = {k:d[k] for k in ["charge_id","srn","date_of_creation","amount","charge_holder",
              "assets_charged","charge_type","date_of_modification","date_of_satisfaction",
              "status","remarks"] if k in d}
    _date_chg = {"date_of_creation","date_of_modification","date_of_satisfaction"}
    fields = {k:(_dt(v) if k in _date_chg else (_num(v) if k=="amount" else v)) for k,v in _fc.items()}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE charges SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM charges WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/charges/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_charge(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM charges WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── Director Interests (MBP-1, Section 189) ──────────────────────────────────
@app.route("/api/director-interests")
@login_required
def list_director_interests():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("SELECT * FROM director_interests WHERE company_id=%s ORDER BY date_of_disclosure DESC",(cid,))
    else:
        c.execute("""SELECT di.*,co.name as company_name FROM director_interests di
                     JOIN companies co ON di.company_id=co.id ORDER BY di.date_of_disclosure DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/director-interests", methods=["POST"])
@login_required
def create_director_interest():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("director_name"): return jsonify({"error":"company_id and director_name required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO director_interests
        (id,company_id,director_id,director_name,din,entity_name,entity_type,
         nature_of_interest,date_of_disclosure,date_of_board_resolution,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d.get("director_id"),d["director_name"],d.get("din"),
         d.get("entity_name",""),d.get("entity_type","Company"),d.get("nature_of_interest"),
         d.get("date_of_disclosure"),d.get("date_of_board_resolution"),d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM director_interests WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/director-interests/<rid>", methods=["PUT"])
@login_required
def update_director_interest(rid):
    d = request.get_json(silent=True, force=True) or {}
    fields = {k:d[k] for k in ["director_name","din","entity_name","entity_type",
              "nature_of_interest","date_of_disclosure","date_of_board_resolution","remarks"] if k in d}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE director_interests SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM director_interests WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/director-interests/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_director_interest(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM director_interests WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── ESOP Grants (SH-6, Section 62) ───────────────────────────────────────────
@app.route("/api/esop-grants")
@login_required
def list_esop_grants():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("SELECT * FROM esop_grants WHERE company_id=%s ORDER BY grant_date DESC",(cid,))
    else:
        c.execute("""SELECT eg.*,co.name as company_name FROM esop_grants eg
                     JOIN companies co ON eg.company_id=co.id ORDER BY eg.grant_date DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/esop-grants", methods=["POST"])
@login_required
def create_esop_grant():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("employee_name"): return jsonify({"error":"company_id and employee_name required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO esop_grants
        (id,company_id,employee_name,designation,employee_id,grant_date,options_granted,
         exercise_price,vesting_date,vesting_period,options_exercised,options_lapsed,status,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d["employee_name"],d.get("designation"),d.get("employee_id"),
         _dt(d.get("grant_date")),_num(d.get("options_granted"),cast=int),_num(d.get("exercise_price")),
         _dt(d.get("vesting_date")),d.get("vesting_period"),_num(d.get("options_exercised"),cast=int),
         _num(d.get("options_lapsed"),cast=int),d.get("status","Active"),d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM esop_grants WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/esop-grants/<rid>", methods=["PUT"])
@login_required
def update_esop_grant(rid):
    d = request.get_json(silent=True, force=True) or {}
    _fe = {k:d[k] for k in ["employee_name","designation","employee_id","grant_date",
              "options_granted","exercise_price","vesting_date","vesting_period",
              "options_exercised","options_lapsed","status","remarks"] if k in d}
    _date_esop = {"grant_date","vesting_date"}
    _num_esop  = {"options_granted","exercise_price","options_exercised","options_lapsed"}
    fields = {k:(_dt(v) if k in _date_esop else (_num(v) if k in _num_esop else v)) for k,v in _fe.items()}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE esop_grants SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM esop_grants WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/esop-grants/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_esop_grant(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM esop_grants WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── Investments (MBP-3, Section 186) ─────────────────────────────────────────
@app.route("/api/investments")
@login_required
def list_investments():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("SELECT * FROM investments WHERE company_id=%s ORDER BY date_of_investment DESC",(cid,))
    else:
        c.execute("""SELECT i.*,co.name as company_name FROM investments i
                     JOIN companies co ON i.company_id=co.id ORDER BY i.date_of_investment DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/investments", methods=["POST"])
@login_required
def create_investment():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("investee_name"): return jsonify({"error":"company_id and investee_name required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO investments
        (id,company_id,investee_name,investee_type,investment_type,amount,date_of_investment,
         board_resolution_date,srn_mgb4,purpose,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d["investee_name"],d.get("investee_type","Company"),
         d.get("investment_type","Equity Shares"),_num(d.get("amount")),
         _dt(d.get("date_of_investment")),_dt(d.get("board_resolution_date")),d.get("srn_mgb4"),
         d.get("purpose"),d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM investments WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/investments/<rid>", methods=["PUT"])
@login_required
def update_investment(rid):
    d = request.get_json(silent=True, force=True) or {}
    _fi = {k:d[k] for k in ["investee_name","investee_type","investment_type","amount",
              "date_of_investment","board_resolution_date","srn_mgb4","purpose","remarks"] if k in d}
    fields = {k:(_dt(v) if k in ("date_of_investment","board_resolution_date") else (_num(v) if k=="amount" else v)) for k,v in _fi.items()}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE investments SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM investments WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/investments/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_investment(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM investments WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── Loans & Guarantees (MBP-2, Section 186) ───────────────────────────────────
@app.route("/api/loans-guarantees")
@login_required
def list_loans():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("SELECT * FROM loans_guarantees WHERE company_id=%s ORDER BY date_of_transaction DESC",(cid,))
    else:
        c.execute("""SELECT lg.*,co.name as company_name FROM loans_guarantees lg
                     JOIN companies co ON lg.company_id=co.id ORDER BY lg.date_of_transaction DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/loans-guarantees", methods=["POST"])
@login_required
def create_loan():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("party_name"): return jsonify({"error":"company_id and party_name required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO loans_guarantees
        (id,company_id,party_name,party_type,transaction_type,amount,date_of_transaction,
         rate_of_interest,repayment_date,security,board_resolution_date,outstanding_amount,status,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d["party_name"],d.get("party_type","Company"),
         d.get("transaction_type","Loan"),_num(d.get("amount")),_dt(d.get("date_of_transaction")),
         _num(d.get("rate_of_interest")),_dt(d.get("repayment_date")),d.get("security"),
         _dt(d.get("board_resolution_date")),_num(d.get("outstanding_amount")),
         d.get("status","Active"),d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM loans_guarantees WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/loans-guarantees/<rid>", methods=["PUT"])
@login_required
def update_loan(rid):
    d = request.get_json(silent=True, force=True) or {}
    _fl = {k:d[k] for k in ["party_name","party_type","transaction_type","amount",
              "date_of_transaction","rate_of_interest","repayment_date","security",
              "board_resolution_date","outstanding_amount","status","remarks"] if k in d}
    _date_fl = {"date_of_transaction","repayment_date","board_resolution_date"}
    _num_fl  = {"amount","rate_of_interest","outstanding_amount"}
    fields = {k:(_dt(v) if k in _date_fl else (_num(v) if k in _num_fl else v)) for k,v in _fl.items()}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE loans_guarantees SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM loans_guarantees WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/loans-guarantees/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_loan(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM loans_guarantees WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ── Related Party Transactions (Section 188) ──────────────────────────────────
@app.route("/api/related-party-transactions")
@login_required
def list_rpt():
    cid = request.args.get("company_id",""); conn = get_db(); c = conn.cursor()
    if cid:
        c.execute("SELECT * FROM related_party_transactions WHERE company_id=%s ORDER BY date_of_transaction DESC",(cid,))
    else:
        c.execute("""SELECT r.*,co.name as company_name FROM related_party_transactions r
                     JOIN companies co ON r.company_id=co.id ORDER BY r.date_of_transaction DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/related-party-transactions", methods=["POST"])
@login_required
def create_rpt():
    d = request.get_json(silent=True, force=True) or {}
    if not d.get("company_id") or not d.get("party_name"): return jsonify({"error":"company_id and party_name required"}),400
    rid = str(uuid.uuid4()); conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO related_party_transactions
        (id,company_id,party_name,relationship,nature_of_transaction,amount,date_of_transaction,
         date_of_board_approval,date_of_shareholders_approval,terms,justification,remarks,created_by)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (rid,d["company_id"],d["party_name"],d.get("relationship"),d.get("nature_of_transaction"),
         _num(d.get("amount")),_dt(d.get("date_of_transaction")),_dt(d.get("date_of_board_approval")),
         _dt(d.get("date_of_shareholders_approval")),d.get("terms"),d.get("justification"),
         d.get("remarks"),g.user_id))
    conn.commit()
    c.execute("SELECT * FROM related_party_transactions WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result),201

@app.route("/api/related-party-transactions/<rid>", methods=["PUT"])
@login_required
def update_rpt(rid):
    d = request.get_json(silent=True, force=True) or {}
    _fr = {k:d[k] for k in ["party_name","relationship","nature_of_transaction","amount",
              "date_of_transaction","date_of_board_approval","date_of_shareholders_approval",
              "terms","justification","remarks"] if k in d}
    _date_rpt = {"date_of_transaction","date_of_board_approval","date_of_shareholders_approval"}
    fields = {k:(_dt(v) if k in _date_rpt else (_num(v) if k=="amount" else v)) for k,v in _fr.items()}
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE related_party_transactions SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[rid])
    conn.commit()
    c.execute("SELECT * FROM related_party_transactions WHERE id=%s",(rid,)); result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/related-party-transactions/<rid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_rpt(rid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM related_party_transactions WHERE id=%s",(rid,)); conn.commit(); conn.close()
    return jsonify({"success":True})

# ══ ALERTS ═══════════════════════════════════════════════════════════════════
@app.route("/api/alerts")
@login_required
def list_alerts():
    uid  = g.user_id
    role = g.role
    cid  = request.args.get("company_id", "")
    sev  = request.args.get("severity",   "")
    atype= request.args.get("type",       "")
    conn = get_db(); c = conn.cursor()

    base = """SELECT a.*, co.name AS company_name
              FROM alerts a
              LEFT JOIN companies co ON a.company_id = co.id
              WHERE a.status = 'active'"""
    params = []

    if role == "superadmin":
        # ── Super Admin: all alerts ──────────────────────────────────────────
        pass

    elif role == "manager":
        # ── Manager: alerts for companies/entities where they manage tasks ───
        # Include: alerts on companies where they are task_manager on ANY task
        # Plus: alerts for entity_type=director/auditor/dsc where related task has task_manager=me
        base += """ AND (
            a.company_id IN (
                SELECT DISTINCT company_id FROM tasks WHERE task_manager = %s AND status NOT IN ('completed','cancelled')
            )
            OR EXISTS (
                SELECT 1 FROM tasks t WHERE t.task_manager = %s
                  AND t.status NOT IN ('completed','cancelled')
                  AND (
                    (a.entity_type = 'director'  AND t.module = 'director')  OR
                    (a.entity_type = 'auditor'   AND t.module = 'auditor')   OR
                    (a.entity_type = 'dsc'       AND t.module = 'dsc')       OR
                    (a.entity_type = 'filing'    AND t.module = 'filing')    OR
                    (a.entity_type = 'meeting'   AND t.module = 'meeting')   OR
                    a.entity_type  = 'company'
                  )
            )
        )"""
        params.extend([uid, uid])

    else:
        # ── Staff: only alerts for entities/companies where they are assigned ─
        base += """ AND (
            a.company_id IN (
                SELECT DISTINCT company_id FROM tasks WHERE assigned_to = %s AND status NOT IN ('completed','cancelled')
            )
            AND (
                EXISTS (
                    SELECT 1 FROM tasks t WHERE t.assigned_to = %s
                      AND t.company_id = a.company_id
                      AND t.status NOT IN ('completed','cancelled')
                      AND (
                        (a.entity_type = 'director'  AND t.module = 'director')  OR
                        (a.entity_type = 'auditor'   AND t.module = 'auditor')   OR
                        (a.entity_type = 'dsc'       AND t.module = 'dsc')       OR
                        (a.entity_type = 'filing'    AND t.module = 'filing')    OR
                        (a.entity_type = 'meeting'   AND t.module = 'meeting')   OR
                        a.entity_type  = 'company'
                      )
                )
            )
        )"""
        params.extend([uid, uid])

    # Common filters
    search = request.args.get("q", "").strip()
    if cid:   base += " AND a.company_id = %s"; params.append(cid)
    if sev:   base += " AND a.severity = %s";   params.append(sev)
    if atype: base += " AND a.entity_type = %s"; params.append(atype)
    if search:
        base += " AND (a.title ILIKE %s OR a.message ILIKE %s OR co.name ILIKE %s)"
        like = f"%{search}%"
        params.extend([like, like, like])

    base += " ORDER BY CASE a.severity WHEN 'critical' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END, a.due_date"
    c.execute(base, params)
    result = rows(c.fetchall())
    conn.close()
    return jsonify(result)

@app.route("/api/alerts/<aid>/dismiss", methods=["POST"])
@login_required
def dismiss_alert(aid):
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE alerts SET status='dismissed',resolved_at=NOW() WHERE id=%s",(aid,))
    conn.commit(); conn.close(); return jsonify({"success":True})

@app.route("/api/alerts/<aid>/resolve", methods=["POST"])
@login_required
def resolve_alert(aid):
    conn=get_db(); c=conn.cursor()
    c.execute("UPDATE alerts SET status='resolved',resolved_at=NOW() WHERE id=%s",(aid,))
    conn.commit(); conn.close(); return jsonify({"success":True})


# ══ BOARD MEETING COMPLIANCE ALERTS ═════════════════════════════════════════
# Rules (Companies Act 2013):
#   1. Every company must hold ≥1 Board Meeting per FY quarter
#      (Q1: Apr–Jun, Q2: Jul–Sep, Q3: Oct–Dec, Q4: Jan–Mar)
#   2. Gap between any two consecutive Board Meetings ≤ 120 days
#   3. Alert raised 15 days BEFORE the deadline
# Implementation notes:
#   • All SQL uses %s placeholders (ORM converts to ? for SQLite automatically)
#   • All NOW()/CURRENT_DATE uses Python date values passed as params
#   • No hardcoded datetime('now') or NOW() in SQL strings (breaks PostgreSQL)
# ═════════════════════════════════════════════════════════════════════════════

def _fy_quarters_bmc(today):
    """Return (label, start, end) for each quarter of the current Indian FY."""
    from datetime import date
    yr = today.year if today.month >= 4 else today.year - 1
    return [
        (f"Q1 FY{yr}-{str(yr+1)[-2:]}", date(yr,4,1),   date(yr,6,30)),
        (f"Q2 FY{yr}-{str(yr+1)[-2:]}", date(yr,7,1),   date(yr,9,30)),
        (f"Q3 FY{yr}-{str(yr+1)[-2:]}", date(yr,10,1),  date(yr,12,31)),
        (f"Q4 FY{yr}-{str(yr+1)[-2:]}", date(yr+1,1,1), date(yr+1,3,31)),
    ]


def _bmc_upsert(c, conn, company_id, tenant_id,
                alert_key, title, message, due_date, severity):
    """
    Insert or update a board_meeting_compliance alert.
    Uses SELECT→UPDATE/INSERT pattern (avoids INSERT OR IGNORE silently eating updates).
    All datetime params passed as Python values — never in SQL strings.
    """
    import uuid
    due_str = due_date.isoformat() if hasattr(due_date, 'isoformat') else str(due_date)
    now_str = _dt_now_str()

    # Check if active alert already exists for this key
    c.execute("""SELECT id FROM alerts
                 WHERE company_id=%s AND alert_type='board_meeting_compliance'
                   AND entity_id=%s AND status='active'""",
              (company_id, alert_key))
    existing = c.fetchone()

    if existing:
        aid = existing['id'] if isinstance(existing, dict) else existing[0]
        c.execute("""UPDATE alerts
                     SET title=%s, message=%s, due_date=%s, severity=%s, created_at=%s
                     WHERE id=%s""",
                  (title, message, due_str, severity, now_str, aid))
    else:
        c.execute("""INSERT INTO alerts
                     (id, company_id, entity_type, entity_id, alert_type,
                      title, message, due_date, severity, status, tenant_id, created_at)
                     VALUES (%s,%s,'meeting',%s,'board_meeting_compliance',
                             %s,%s,%s,%s,'active',%s,%s)""",
                  (str(uuid.uuid4()), company_id, alert_key,
                   title, message, due_str, severity, tenant_id, now_str))


def _bmc_clear_stale(c, company_id, active_keys):
    """Resolve board_meeting_compliance alerts whose keys are no longer needed."""
    now_str = _dt_now_str()
    if active_keys:
        ph = ','.join(['%s'] * len(active_keys))
        c.execute(f"""UPDATE alerts
                      SET status='resolved', resolved_at=%s
                      WHERE company_id=%s
                        AND alert_type='board_meeting_compliance'
                        AND status='active'
                        AND entity_id NOT IN ({ph})""",
                  [now_str, company_id] + list(active_keys))
    else:
        c.execute("""UPDATE alerts
                     SET status='resolved', resolved_at=%s
                     WHERE company_id=%s
                       AND alert_type='board_meeting_compliance'
                       AND status='active'""",
                  (now_str, company_id))


def _dt_now_str():
    """Return current datetime as ISO string — works for both SQLite and PostgreSQL."""
    from datetime import datetime
    return datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')


def sync_board_meeting_alerts(company_id=None, tenant_id=None):
    """
    Board meeting compliance alert engine.
    Safe for both SQLite (dev) and PostgreSQL (production).
    """
    from datetime import date, timedelta, datetime as dt
    today = date.today()
    ALERT_DAYS  = 15   # alert this many days before deadline
    GAP_DAYS    = 120  # max gap between consecutive board meetings

    conn = get_db()
    c    = conn.cursor()

    # ── Load companies ────────────────────────────────────────────────────────
    if company_id:
        c.execute("SELECT id, name, tenant_id FROM companies WHERE id=%s", (company_id,))
    elif tenant_id:
        c.execute("SELECT id, name, tenant_id FROM companies WHERE tenant_id=%s", (tenant_id,))
    else:
        c.execute("SELECT id, name, tenant_id FROM companies")
    companies = rows(c.fetchall())

    quarters = _fy_quarters_bmc(today)

    for co in companies:
        cid   = co['id']
        cname = co['name']
        ctid  = co.get('tenant_id') or tenant_id

        # ── Load board meetings (past + future) ───────────────────────────────
        c.execute("""SELECT meeting_date FROM meetings
                     WHERE company_id=%s
                       AND meeting_type='Board'
                       AND status IN ('held','completed','scheduled','minutes_approved')
                     ORDER BY meeting_date ASC""", (cid,))
        raw_dates = [r['meeting_date'] if isinstance(r, dict) else r[0]
                     for r in c.fetchall()]

        # Normalise to date objects
        mtg_dates = []
        for d in raw_dates:
            if isinstance(d, date):
                mtg_dates.append(d)
            else:
                try:
                    mtg_dates.append(dt.strptime(str(d)[:10], '%Y-%m-%d').date())
                except Exception:
                    pass

        active_keys = set()

        # ── RULE 1: Quarterly check ───────────────────────────────────────────
        for (qlabel, qstart, qend) in quarters:
            # Skip quarters that ended more than 30 days ago (no longer actionable)
            if qend < today - timedelta(days=30):
                continue

            held_in_q  = [d for d in mtg_dates if qstart <= d <= qend and d <= today]
            sched_in_q = [d for d in mtg_dates if qstart <= d <= qend and d > today]

            if held_in_q:
                # Requirement already met for this quarter
                continue

            # Only alert if we're within 15 days of quarter-end OR already in quarter
            alert_trigger = qend - timedelta(days=ALERT_DAYS)
            if not (today >= alert_trigger or today >= qstart):
                continue

            days_left = (qend - today).days
            key = f"quarterly_{cid}_{qlabel.replace(' ', '_')}"
            active_keys.add(key)

            if sched_in_q:
                sched_str = min(sched_in_q).strftime('%d %b %Y')
                sev   = 'medium'
                title = f"Board Meeting Scheduled — {qlabel} ({cname})"
                msg   = (f"A Board Meeting is scheduled for {sched_str} in {qlabel}. "
                         f"Quarter ends {qend.strftime('%d %b %Y')} ({days_left}d left).")
            else:
                sev   = ('critical' if days_left <= 7 else
                         'high'     if days_left <= 15 else 'medium')
                title = f"Board Meeting Pending — {qlabel} ({cname})"
                msg   = (f"No Board Meeting held or scheduled for {qlabel} "
                         f"({qstart.strftime('%d %b')} – {qend.strftime('%d %b %Y')}). "
                         f"Deadline: {qend.strftime('%d %b %Y')} ({days_left}d left).")

            _bmc_upsert(c, conn, cid, ctid, key, title, msg, qend, sev)

        # ── RULE 2: 120-day gap check ─────────────────────────────────────────
        past_dates = sorted([d for d in mtg_dates if d <= today])
        if past_dates:
            last_held    = past_dates[-1]
            gap_deadline = last_held + timedelta(days=GAP_DAYS)
            alert_trigger = gap_deadline - timedelta(days=ALERT_DAYS)

            if today >= alert_trigger:
                future_before = [d for d in mtg_dates if d > today and d <= gap_deadline]
                days_left = (gap_deadline - today).days
                key = f"gap120_{cid}_{last_held.isoformat()}"
                active_keys.add(key)

                if future_before:
                    sched_str = min(future_before).strftime('%d %b %Y')
                    sev   = 'medium'
                    title = f"120-Day Gap — Meeting Scheduled ({cname})"
                    msg   = (f"Last Board Meeting: {last_held.strftime('%d %b %Y')}. "
                             f"Next scheduled {sched_str} (within 120-day limit of "
                             f"{gap_deadline.strftime('%d %b %Y')}).")
                else:
                    sev   = ('critical' if days_left <= 7 else
                             'high'     if days_left <= 15 else 'medium')
                    title = f"120-Day Board Meeting Gap Alert ({cname})"
                    msg   = (f"Last Board Meeting: {last_held.strftime('%d %b %Y')}. "
                             f"Must hold next by {gap_deadline.strftime('%d %b %Y')} "
                             f"(120-day rule). {days_left} days remaining — no meeting scheduled.")

                _bmc_upsert(c, conn, cid, ctid, key, title, msg, gap_deadline, sev)

        # ── Resolve stale alerts ──────────────────────────────────────────────
        _bmc_clear_stale(c, cid, active_keys)

    conn.commit()
    conn.close()


@app.route("/api/alerts/sync-board-meetings", methods=["GET", "POST"])
@login_required
def sync_board_meeting_alerts_route():
    """Sync board meeting compliance alerts. Callable by superadmin or manager."""
    if g.role not in ("superadmin", "manager"):
        return jsonify({"error": "Forbidden"}), 403
    try:
        body = request.get_json(silent=True, force=True) or {}
        if not isinstance(body, dict):
            body = {}
        cid = body.get("company_id")
        tid = getattr(g, "tenant_id", None)
        sync_board_meeting_alerts(company_id=cid, tenant_id=tid)
        return jsonify({"success": True, "message": "Board meeting compliance alerts synced."})
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "detail": str(_ex)[:200] if app.debug else "Internal error"}), 500


# ══ TASKS ════════════════════════════════════════════════════════════════════
@app.route("/api/tasks")
@login_required
def list_tasks():
    uid  = g.user_id
    role = g.role
    cid    = request.args.get("company_id","")
    status = request.args.get("status","")
    # hierarchy_view: only honoured for superadmin (Task Leader)
    # Values: 'mine'(default) | 'manager' | 'staff' | 'all'
    hierarchy_view = request.args.get("hierarchy_view", "")

    conn=get_db(); c=conn.cursor()
    q="""SELECT t.*,
               co.name  AS company_name,
               ul.name  AS task_leader_name,
               um.name  AS task_manager_name,
               ua.name  AS assigned_to_name
               FROM tasks t
               LEFT JOIN companies co ON t.company_id=co.id
               LEFT JOIN users ul ON t.task_leader=ul.id
               LEFT JOIN users um ON t.task_manager=um.id
               LEFT JOIN users ua ON t.assigned_to=ua.id
               WHERE 1=1"""
    params=[]

    # Role-based visibility:
    # superadmin (Task Leader): by default sees only tasks where task_leader=me
    #   hierarchy_view='manager' → tasks assigned to any manager
    #   hierarchy_view='staff'   → tasks assigned to any staff (assigned_to)
    #   hierarchy_view='all'     → all tasks in the tenant
    # manager: sees tasks where task_manager=me
    # staff:   sees tasks where assigned_to=me
    if role == "superadmin":
        if hierarchy_view == "manager":
            q += " AND t.task_manager IS NOT NULL"
        elif hierarchy_view == "staff":
            q += " AND t.assigned_to IS NOT NULL"
        elif hierarchy_view == "all":
            pass  # no role filter — see everything
        else:
            # Default: my tasks as Task Leader
            q += " AND t.task_leader=%s"; params.append(uid)
    elif role == "manager":
        q += " AND t.task_manager=%s"; params.append(uid)
    else:
        q += " AND t.assigned_to=%s"; params.append(uid)

    if cid: q+=" AND t.company_id=%s"; params.append(cid)
    if status: q+=" AND t.status=%s"; params.append(status)
    # Order without NULLS LAST (SQLite compat): push NULLs last via CASE
    order_sql = " ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END, CASE WHEN t.due_date IS NULL THEN 1 ELSE 0 END, t.due_date"
    # Paginated path
    if request.args.get("page") or request.args.get("limit"):
        limit = request.args.get("limit", 0, type=int)
        page  = max(1, request.args.get("page", 1, type=int))
        count_sql = f"SELECT COUNT(*) FROM ({q}) AS _cnt"
        c.execute(count_sql, params if params else [])
        _r = c.fetchone()
        total = int(list(_r.values())[0] if isinstance(_r, dict) else _r[0]) if _r else 0
        if limit > 0:
            offset = (page-1)*limit
            c.execute(q + order_sql + f" LIMIT {limit} OFFSET {offset}", params if params else [])
        else:
            c.execute(q + order_sql, params if params else [])
        data = rows(c.fetchall())
        pages = (total + limit - 1) // limit if limit > 0 else 1
        conn.close()
        return jsonify({"data": data, "total": total, "page": page, "limit": limit, "pages": pages})
    # Default: plain array (backwards-compat)
    c.execute(q + order_sql, params if params else [])
    data = rows(c.fetchall())
    conn.close()
    return jsonify(data)

@app.route("/api/tasks", methods=["POST"])
@login_required
def create_task():
    d=request.get_json(silent=True, force=True) or {}; tid=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    _task_tid = g.tenant_id
    c.execute("""INSERT INTO tasks
        (id,company_id,title,description,assigned_to,due_date,priority,status,
         module,entity_id,created_by,task_leader,task_manager,team_members,
         billable,estimated_hrs,note,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (tid,_str(d.get("company_id")),d["title"],d.get("description"),_str(d.get("assigned_to")),
         _dt(d.get("due_date")),d.get("priority","medium"),d.get("status","pending"),
         d.get("module"),d.get("entity_id"),g.user_id,
         d.get("task_leader") or None,
         d.get("task_manager") or None,
         json.dumps(d.get("team_members") or []),
         1 if d.get("billable") else 0,
         float(d.get("estimated_hrs") or 0),
         d.get("note",""),
         _task_tid))
    conn.commit()
    c.execute("SELECT t.*,co.name as company_name FROM tasks t LEFT JOIN companies co ON t.company_id=co.id WHERE t.id=%s",(tid,))
    result=row(c.fetchone()); conn.close(); return jsonify(result),201

@app.route("/api/tasks/<tid>", methods=["PUT"])
@login_required
def update_task(tid):
    d=request.get_json(silent=True, force=True) or {}
    _rtask = {k:d[k] for k in ["title","description","assigned_to","due_date","priority",
                                "status","module","task_leader","task_manager",
                                "billable","estimated_hrs","actual_hrs","note"] if k in d}
    fields = {}
    for _k,_v in _rtask.items():
        if _k == "due_date": fields[_k] = _dt(_v)
        elif _k in ("assigned_to","task_leader","task_manager"): fields[_k] = _str(_v)
        elif _k in ("estimated_hrs","actual_hrs"): fields[_k] = _num(_v)
        else: fields[_k] = _v
    if "team_members" in d: fields["team_members"] = json.dumps(d["team_members"] or [])
    if d.get("status")=="completed": fields["completed_at"]=datetime.utcnow().isoformat()
    conn=get_db(); c=conn.cursor()
    c.execute(f"UPDATE tasks SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",list(fields.values())+[tid])
    conn.commit()
    c.execute("SELECT t.*,co.name as company_name FROM tasks t LEFT JOIN companies co ON t.company_id=co.id WHERE t.id=%s",(tid,))
    result=row(c.fetchone()); conn.close(); return jsonify(result)

@app.route("/api/tasks/<tid>", methods=["DELETE"])
@login_required
def delete_task(tid):
    conn=get_db(); c=conn.cursor(); c.execute("DELETE FROM tasks WHERE id=%s",(tid,)); conn.commit(); conn.close()
    return jsonify({"success":True})


# ══ EXCEL EXPORT ══════════════════════════════════════════════════════════════

@app.route("/api/export/excel", methods=["POST"])
@login_required
def export_excel():
    """
    Universal Excel export endpoint.
    Body: { module, company_id (optional), date_from, date_to, company_name }
    Returns a formatted .xlsx file.
    """
    d = request.get_json(silent=True, force=True) or {}
    module      = d.get("module", "")
    company_id  = d.get("company_id", "")
    date_from   = d.get("date_from", "")
    date_to     = d.get("date_to", "")
    co_name     = d.get("company_name", "All Companies")

    conn = get_db(); c = conn.cursor()

    # ── Date filter helper ────────────────────────────────────────────────────
    def date_filter(col, params):
        clause = ""
        if date_from:
            clause += f" AND {col} >= %s"
            params.append(date_from)
        if date_to:
            clause += f" AND {col} <= %s"
            params.append(date_to)
        return clause

    def co_filter(col, params):
        if company_id:
            params.append(company_id)
            return f" AND {col} = %s"
        return ""

    # ── Fetch data per module ─────────────────────────────────────────────────
    sheets = {}

    if module in ("companies", "all"):
        params = []
        q = "SELECT name,cin,company_type,pan,tan,gstin,email,phone,incorporation_date,registered_office,authorized_capital,paid_up_capital,business_activity,roc,status,created_at FROM companies WHERE 1=1"
        q += date_filter("created_at", params)
        c.execute(q, params)
        sheets["Companies"] = {
            "headers": ["Company Name","CIN","Type","PAN","TAN","GSTIN","Email","Phone","Incorporation Date","Registered Office","Auth Capital","Paid-up Capital","Business Activity","ROC","Status","Created"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("directors", "all"):
        params = []
        q = """SELECT d.name AS director_name,d.din,d.pan,d.designation,
                      co.name AS company_name,d.email,d.mobile,
                      d.date_of_appointment,d.date_of_cessation,
                      k.kyc_status,k.last_kyc_date,k.next_due_date,
                      CASE d.is_active WHEN 1 THEN 'Active' ELSE 'Inactive' END AS status
               FROM directors d
               LEFT JOIN director_kyc k ON d.id=k.director_id
               LEFT JOIN companies co ON d.company_id=co.id WHERE 1=1"""
        q += co_filter("d.company_id", params)
        q += date_filter("d.date_of_appointment", params)
        q += " ORDER BY co.name,d.name"
        c.execute(q, params)
        sheets["Directors"] = {
            "headers": ["Name","DIN","PAN","Designation","Company","Email","Mobile","Date Appointed","Date Ceased","KYC Status","Last KYC","KYC Due Date","Status"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("auditors", "all"):
        params = []
        q = """SELECT a.name AS auditor_name,a.firm_name,a.membership_no,a.frn,a.pan,
                      co.name AS company_name,a.nature_of_appointment,a.appointment_type,
                      a.start_date,a.end_date,a.srn_adt1,a.email,
                      CASE a.is_active WHEN 1 THEN 'Active' ELSE 'Inactive' END AS status
               FROM auditors a LEFT JOIN companies co ON a.company_id=co.id WHERE 1=1"""
        q += co_filter("a.company_id", params)
        q += date_filter("a.start_date", params)
        q += " ORDER BY co.name,a.start_date"
        c.execute(q, params)
        sheets["Auditors"] = {
            "headers": ["Auditor Name","Firm Name","Membership No","FRN","PAN","Company","Nature","Appointment Type","Start Date","End Date","ADT-1 SRN","Email","Status"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("shareholders", "all"):
        params = []
        q = """SELECT s.name AS shareholder_name,s.folio_no,s.pan,
                      co.name AS company_name,s.share_class,
                      s.shares_held,s.face_value,s.email,s.mobile,s.date_of_entry,
                      CASE s.is_active WHEN 1 THEN 'Active' ELSE 'Inactive' END AS status
               FROM shareholders s LEFT JOIN companies co ON s.company_id=co.id WHERE 1=1"""
        q += co_filter("s.company_id", params)
        q += date_filter("s.date_of_entry", params)
        q += " ORDER BY co.name,s.folio_no"
        c.execute(q, params)
        sheets["Shareholders"] = {
            "headers": ["Name","Folio No","PAN","Company","Share Class","Shares Held","Face Value","Email","Mobile","Date of Entry","Status"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("dsc", "all"):
        params = []
        q = """SELECT d.holder_name,d.holder_type,d.dsc_class,d.issued_by,
                      d.token_type,d.valid_from,d.valid_to,
                      d.custody_status,d.custody_date,co.name,d.notes
               FROM dsc_records d LEFT JOIN companies co ON d.company_id=co.id WHERE d.is_active=1"""
        q += co_filter("d.company_id", params)
        q += date_filter("d.valid_from", params)
        q += " ORDER BY d.valid_to"
        c.execute(q, params)
        sheets["DSC Records"] = {
            "headers": ["Holder Name","Holder Type","Class","Issued By","Token","Valid From","Valid To","Custody Status","Custody Date","Company","Notes"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("meetings", "all"):
        params = []
        q = """SELECT m.meeting_type,m.meeting_no,m.meeting_date,m.meeting_time,
                      m.venue,co.name,m.status,
                      CASE m.minutes_drafted WHEN 1 THEN 'Yes' ELSE 'No' END,
                      m.agenda
               FROM meetings m LEFT JOIN companies co ON m.company_id=co.id WHERE 1=1"""
        q += co_filter("m.company_id", params)
        q += date_filter("m.meeting_date", params)
        q += " ORDER BY m.meeting_date DESC"
        c.execute(q, params)
        sheets["Meetings"] = {
            "headers": ["Type","Meeting No","Date","Time","Venue","Company","Status","Minutes Done","Agenda"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("alerts", "all"):
        params = []
        q = """SELECT a.title,a.entity_type,a.alert_type,a.severity,a.due_date,
                      a.status,co.name,a.message,a.created_at
               FROM alerts a LEFT JOIN companies co ON a.company_id=co.id WHERE 1=1"""
        q += co_filter("a.company_id", params)
        q += date_filter("a.due_date", params)
        q += " ORDER BY CASE a.severity WHEN 'critical' THEN 1 WHEN 'high' THEN 2 ELSE 3 END,a.due_date"
        c.execute(q, params)
        sheets["Alerts"] = {
            "headers": ["Title","Entity Type","Alert Type","Severity","Due Date","Status","Company","Message","Created"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("tasks", "all"):
        params = []
        q = """SELECT t.title,t.priority,t.status,t.due_date,t.module,
                      co.name AS company_name,u.name AS assigned_to_name,
                      t.description,
                      TO_CHAR(t.created_at,'YYYY-MM-DD') AS created_at,
                      CASE WHEN t.completed_at IS NULL THEN NULL
                           ELSE TO_CHAR(t.completed_at,'YYYY-MM-DD') END AS completed_at
               FROM tasks t
               LEFT JOIN companies co ON t.company_id=co.id
               LEFT JOIN users u ON t.assigned_to=u.id WHERE 1=1"""
        q += co_filter("t.company_id", params)
        q += date_filter("t.due_date", params)
        q += " ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END,t.due_date"
        c.execute(q, params)
        sheets["Tasks"] = {
            "headers": ["Title","Priority","Status","Due Date","Module","Company","Assigned To","Description","Created","Completed"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("charges", "all"):
        params = []
        q = """SELECT ch.charge_id,ch.charge_type,ch.charge_holder,ch.assets_charged,
                      ch.amount,ch.date_of_creation,ch.date_of_modification,
                      ch.date_of_satisfaction,ch.status,ch.srn,co.name,ch.remarks
               FROM charges ch LEFT JOIN companies co ON ch.company_id=co.id WHERE 1=1"""
        q += co_filter("ch.company_id", params)
        q += date_filter("ch.date_of_creation", params)
        q += " ORDER BY ch.date_of_creation"
        c.execute(q, params)
        sheets["Charges"] = {
            "headers": ["Charge ID","Type","Charge Holder","Assets Charged","Amount","Date Created","Date Modified","Date Satisfied","Status","SRN","Company","Remarks"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("investments", "all"):
        params = []
        q = """SELECT i.investee_name,i.investee_type,i.investment_type,i.amount,
                      i.date_of_investment,i.board_resolution_date,i.srn_mgb4,
                      i.purpose,co.name,i.remarks
               FROM investments i LEFT JOIN companies co ON i.company_id=co.id WHERE 1=1"""
        q += co_filter("i.company_id", params)
        q += date_filter("i.date_of_investment", params)
        q += " ORDER BY i.date_of_investment"
        c.execute(q, params)
        sheets["Investments"] = {
            "headers": ["Investee Name","Type","Investment Type","Amount","Date","Board Resolution Date","MGT-14 SRN","Purpose","Company","Remarks"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("loans", "all"):
        params = []
        q = """SELECT lg.party_name,lg.party_type,lg.transaction_type,lg.amount,
                      lg.date_of_transaction,lg.rate_of_interest,lg.repayment_date,
                      lg.outstanding_amount,lg.security,lg.status,
                      lg.board_resolution_date,co.name,lg.remarks
               FROM loans_guarantees lg LEFT JOIN companies co ON lg.company_id=co.id WHERE 1=1"""
        q += co_filter("lg.company_id", params)
        q += date_filter("lg.date_of_transaction", params)
        q += " ORDER BY lg.date_of_transaction"
        c.execute(q, params)
        sheets["Loans & Guarantees"] = {
            "headers": ["Party Name","Party Type","Transaction Type","Amount","Date","Interest %","Repayment Date","Outstanding","Security","Status","Board Resolution","Company","Remarks"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("rpt", "all"):
        params = []
        q = """SELECT r.party_name,r.relationship,r.nature_of_transaction,r.amount,
                      r.date_of_transaction,r.date_of_board_approval,
                      r.date_of_shareholders_approval,r.terms,co.name,r.remarks
               FROM related_party_transactions r LEFT JOIN companies co ON r.company_id=co.id WHERE 1=1"""
        q += co_filter("r.company_id", params)
        q += date_filter("r.date_of_transaction", params)
        q += " ORDER BY r.date_of_transaction"
        c.execute(q, params)
        sheets["Related Party Txns"] = {
            "headers": ["Party Name","Relationship","Nature","Amount","Date","Board Approval","Shareholder Approval","Terms","Company","Remarks"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    if module in ("kyc", "all"):
        params = []
        q = """SELECT d.name AS director_name,d.din,d.pan,
                      co.name AS company_name,d.designation,
                      k.last_kyc_date,k.next_due_date,k.kyc_status,d.email,d.mobile
               FROM directors d
               LEFT JOIN director_kyc k ON d.id=k.director_id
               LEFT JOIN companies co ON d.company_id=co.id
               WHERE d.is_active=1"""
        q += co_filter("d.company_id", params)
        q += " ORDER BY k.kyc_status DESC,k.next_due_date"
        c.execute(q, params)
        sheets["Director KYC"] = {
            "headers": ["Director Name","DIN","PAN","Company","Designation","Last KYC Date","KYC Due Date","KYC Status","Email","Mobile"],
            "rows": [(list(r.values()) if isinstance(r,dict) else list(r)) for r in c.fetchall()]
        }

    conn.close()

    if not sheets:
        return jsonify({"error": f"Unknown module: {module}"}), 400

    # ── Build Excel workbook ──────────────────────────────────────────────────
    buf = io.BytesIO()
    wb  = xlsxwriter.Workbook(buf, {"in_memory": True, "remove_timezone": True})

    # Formats
    fmt_title  = wb.add_format({"bold":True,"font_size":14,"font_color":"#0f2044","border":0})
    fmt_meta   = wb.add_format({"italic":True,"font_size":9,"font_color":"#5a6a88","border":0})
    fmt_hdr    = wb.add_format({"bold":True,"bg_color":"#0f2044","font_color":"#ffffff",
                                 "font_size":10,"border":1,"border_color":"#0f2044",
                                 "text_wrap":True,"valign":"vcenter","align":"center"})
    fmt_row_a  = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "bg_color":"#f8fafc","valign":"vcenter"})
    fmt_row_b  = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "bg_color":"#ffffff","valign":"vcenter"})
    fmt_num    = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "num_format":"#,##0.00","align":"right"})
    fmt_date   = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "align":"center"})
    fmt_red    = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "font_color":"#ef4444","bold":True})
    fmt_green  = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "font_color":"#10b981","bold":True})
    fmt_orange = wb.add_format({"font_size":10,"border":1,"border_color":"#e2e8f0",
                                 "font_color":"#f59e0b","bold":True})

    numeric_keywords = {"amount","capital","value","shares","rate","outstanding","granted","exercised","lapsed","price"}
    date_keywords    = {"date","from","to","appointed","ceased","entry","created","modified","satisfied","due","expiry","valid"}
    status_values    = {"active","inactive","open","satisfied","pending","overdue","critical","high","medium","low","filed","compliant"}

    period_str = ""
    if date_from or date_to:
        period_str = f"Period: {date_from or 'Start'} to {date_to or 'Today'}"

    for sheet_name, data in sheets.items():
        ws = wb.add_worksheet(sheet_name[:31])
        headers = data["headers"]
        rows    = data["rows"]

        # Title rows
        ws.merge_range(0, 0, 0, max(len(headers)-1,1), f"TAXLY-CMS — {sheet_name.upper()} REPORT", fmt_title)
        meta_parts = [f"Company: {co_name}"]
        if period_str: meta_parts.append(period_str)
        meta_parts.append(f"Generated: {date.today().strftime('%d %b %Y')}")
        ws.merge_range(1, 0, 1, max(len(headers)-1,1), "   |   ".join(meta_parts), fmt_meta)
        ws.merge_range(2, 0, 2, max(len(headers)-1,1), f"Total Records: {len(rows)}", fmt_meta)

        ws.set_row(0, 22); ws.set_row(1, 14); ws.set_row(2, 14); ws.set_row(3, 20)

        # Headers
        for ci, hdr in enumerate(headers):
            ws.write(3, ci, hdr, fmt_hdr)

        # Auto column widths
        col_widths = [max(len(str(h)), 8) for h in headers]

        # Data rows
        for ri, row_data in enumerate(rows):
            row_fmt = fmt_row_a if ri % 2 == 0 else fmt_row_b
            for ci, val in enumerate(row_data):
                val_str = str(val) if val is not None else ""
                hdr_lower = headers[ci].lower() if ci < len(headers) else ""

                # Choose format
                if val is None or val == "":
                    ws.write(ri+4, ci, "—", row_fmt)
                elif any(kw in hdr_lower for kw in numeric_keywords):
                    try:
                        ws.write_number(ri+4, ci, float(val), fmt_num)
                    except: ws.write(ri+4, ci, val_str, row_fmt)
                elif any(kw in hdr_lower for kw in date_keywords) and val_str and len(val_str) >= 8:
                    ws.write(ri+4, ci, _ds(val) or val_str[:10], fmt_date)
                elif val_str.lower() in {"critical","expired","overdue","inactive","fail"}:
                    ws.write(ri+4, ci, val_str, fmt_red)
                elif val_str.lower() in {"active","valid","filed","compliant","completed","open"}:
                    ws.write(ri+4, ci, val_str, fmt_green)
                elif val_str.lower() in {"high","expiring_soon","pending","due_soon","in_progress"}:
                    ws.write(ri+4, ci, val_str, fmt_orange)
                else:
                    ws.write(ri+4, ci, val_str, row_fmt)

                col_widths[ci] = max(col_widths[ci], min(len(val_str), 40))

        # Set column widths
        for ci, w in enumerate(col_widths):
            ws.set_column(ci, ci, w + 2)

        # Freeze header rows
        ws.freeze_panes(4, 0)
        ws.autofilter(3, 0, 3, len(headers)-1)

    # Summary sheet for "all" exports
    if module == "all":
        ws_sum = wb.add_worksheet("Summary")
        ws_sum.set_column(0, 0, 28); ws_sum.set_column(1, 1, 12)
        ws_sum.merge_range(0, 0, 0, 1, "TAXLY-CMS — FULL EXPORT SUMMARY", fmt_title)
        ws_sum.merge_range(1, 0, 1, 1, f"Company: {co_name}   |   {period_str}   |   {date.today().strftime('%d %b %Y')}", fmt_meta)
        ws_sum.set_row(0, 22); ws_sum.set_row(1, 14); ws_sum.set_row(3, 18)
        ws_sum.write(3, 0, "Module", fmt_hdr); ws_sum.write(3, 1, "Records", fmt_hdr)
        for ri, (sname, sdata) in enumerate(sheets.items()):
            sum_fmt = fmt_row_a if ri%2==0 else fmt_row_b
            ws_sum.write(ri+4, 0, sname, sum_fmt)
            ws_sum.write_number(ri+4, 1, len(sdata["rows"]), fmt_num)

    wb.close()
    buf.seek(0)

    # Filename
    safe_co = co_name.replace(" ","_").replace("/","_")[:25]
    safe_mod = module.title()
    fname = f"TaxlyCMS_{safe_mod}_{safe_co}_{date.today().strftime('%Y%m%d')}.xlsx"

    return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                     as_attachment=True, download_name=fname)


# ══ BULK UPLOAD ═══════════════════════════════════════════════════════════════

# Module → (sheet_name, required_cols, optional_cols, insert_fn)
BULK_TEMPLATES = {
    "companies": {
        "sheet": "Companies",
        "headers": ["Company Name*","CIN","Company Type","PAN","TAN","GSTIN","Email","Phone",
                    "Incorporation Date (YYYY-MM-DD)","Registered Office","Authorised Capital","Paid-up Capital",
                    "Business Activity","ROC"],
        "sample": [["Innovate Solutions Pvt Ltd","U72200MH2018PTC309876","Private Limited","AABCI1234P",
                    "MUMA12345A","27AABCI1234P1Z5","cs@company.in","022-12345678",
                    "2018-03-15","12 Nariman Point, Mumbai","5000000","1000000","Software Development","RoC-Mumbai"]],
    },
    "directors": {
        "sheet": "Directors",
        "headers": ["Company CIN*","Director Name*","DIN","PAN","Designation","Email","Mobile",
                    "Date of Appointment (YYYY-MM-DD)","Address","MCA User ID","MCA Password"],
        "sample": [["U72200MH2018PTC309876","Arjun Mehta","00123456","ABCPM1234D",
                    "Managing Director","arjun@company.in","9876543210","2018-03-15",
                    "123 MG Road, Mumbai","arjun_mca","password123"]],
    },
    "auditors": {
        "sheet": "Auditors",
        "headers": ["Company CIN*","Auditor Name*","Firm Name","Membership No","FRN","PAN","Email",
                    "Nature of Appointment","Appointment Type","Start Date (YYYY-MM-DD)",
                    "End Date (YYYY-MM-DD)","ADT-1 SRN"],
        "sample": [["U72200MH2018PTC309876","CA Ramesh Gupta","Gupta & Associates","123456",
                    "012345W","ABCPG3456G","ca@firm.in",
                    "Subsequent Auditor","AGM Appointment","2024-09-30","2025-09-29","ADT1202312345"]],
    },
    "shareholders": {
        "sheet": "Shareholders",
        "headers": ["Company CIN*","Holder Name*","Folio No","PAN","Email","Mobile",
                    "Share Class","Shares Held","Face Value","Date of Entry (YYYY-MM-DD)",
                    "Address","MCA User ID","MCA Password"],
        "sample": [["U72200MH2018PTC309876","Arjun Mehta","F0001","ABCPM1234D",
                    "arjun@company.in","9876543210","Equity","50000","10","2018-03-15",
                    "123 MG Road, Mumbai","arjun_mca","password123"]],
    },
    "directors_kyc": {
        "sheet": "Director KYC",
        "headers": ["Director DIN*","Last KYC Date (YYYY-MM-DD)*"],
        "sample": [["00123456","2024-08-15"],["00234567","2024-09-10"]],
    },
    "meetings": {
        "sheet": "Meetings",
        "headers": ["Company CIN*","Meeting Type*","Meeting No","Date (YYYY-MM-DD)*","Time (HH:MM)",
                    "Venue","Status","Agenda"],
        "sample": [["U72200MH2018PTC309876","Board","BM-2025-01","2025-04-15","11:00",
                    "Registered Office, Mumbai","scheduled",
                    "1. Financial review\n2. Director KYC status"]],
    },
    "tasks": {
        "sheet": "Tasks",
        "headers": ["Title*","Company CIN","Priority","Status","Due Date (YYYY-MM-DD)","Module","Description"],
        "sample": [["File ADT-1","U72200MH2018PTC309876","high","pending","2025-09-30","auditor",
                    "File ADT-1 for statutory auditor appointment"]],
    },
    "charges": {
        "sheet": "Charges",
        "headers": ["Company CIN*","Charge Holder*","Charge Type","Amount","Date of Creation (YYYY-MM-DD)",
                    "Assets Charged","Status","Charge ID","SRN"],
        "sample": [["U72200MH2018PTC309876","State Bank of India","Hypothecation","10000000",
                    "2024-01-15","Plant & Machinery","Open","CHG-001","CHG120234567"]],
    },
    "investments": {
        "sheet": "Investments",
        "headers": ["Company CIN*","Investee Name*","Investee Type","Investment Type","Amount",
                    "Date of Investment (YYYY-MM-DD)","Board Resolution Date (YYYY-MM-DD)","Purpose"],
        "sample": [["U72200MH2018PTC309876","TechVentures Ltd","Company","Equity Shares",
                    "500000","2024-01-15","2024-01-10","Strategic investment"]],
    },
}

@app.route("/api/bulk/template/<module>")
@login_required
def bulk_template(module):
    """Download pre-formatted Excel template for bulk upload."""
    tmpl = BULK_TEMPLATES.get(module)
    if not tmpl: return jsonify({"error": f"No template for module: {module}"}), 400

    buf = io.BytesIO()
    wb  = xlsxwriter.Workbook(buf, {"in_memory": True})

    # Formats
    fmt_title   = wb.add_format({"bold":True,"font_size":14,"font_color":"#0f2044"})
    fmt_info    = wb.add_format({"italic":True,"font_size":9,"font_color":"#5a6a88"})
    fmt_hdr     = wb.add_format({"bold":True,"bg_color":"#0f2044","font_color":"#ffffff",
                                  "font_size":10,"border":1,"border_color":"#0f2044",
                                  "text_wrap":True,"valign":"vcenter","align":"center"})
    fmt_hdr_req = wb.add_format({"bold":True,"bg_color":"#ef4444","font_color":"#ffffff",
                                  "font_size":10,"border":1,"border_color":"#c53030",
                                  "text_wrap":True,"valign":"vcenter","align":"center"})
    fmt_sample  = wb.add_format({"font_size":10,"bg_color":"#f0fdf4","border":1,
                                  "border_color":"#bbf7d0","font_color":"#166534"})
    fmt_input   = wb.add_format({"font_size":10,"bg_color":"#fafafa","border":1,
                                  "border_color":"#e2e8f0"})

    ws = wb.add_worksheet(tmpl["sheet"])
    headers = tmpl["headers"]

    # Title
    ws.merge_range(0,0,0,len(headers)-1, f"TAXLY-CMS — {tmpl['sheet'].upper()} BULK UPLOAD TEMPLATE", fmt_title)
    ws.merge_range(1,0,1,len(headers)-1,
        "▶ Red headers = Required  |  Grey = Optional  |  Green rows = Sample data (delete before upload)", fmt_info)
    ws.merge_range(2,0,2,len(headers)-1,
        "• Date format: YYYY-MM-DD  |  Do not change column order  |  Upload via Bulk Upload button", fmt_info)
    ws.set_row(0, 22); ws.set_row(1, 14); ws.set_row(2, 14); ws.set_row(3, 22)

    # Headers
    for ci, hdr in enumerate(headers):
        is_req = hdr.endswith("*")
        label  = hdr.rstrip("*")
        ws.write(3, ci, label, fmt_hdr_req if is_req else fmt_hdr)
        ws.set_column(ci, ci, max(len(label)+4, 14))

    # Sample rows
    for ri, row_data in enumerate(tmpl["sample"]):
        for ci, val in enumerate(row_data):
            ws.write(ri+4, ci, val, fmt_sample)

    # Blank input rows (10 rows ready for data)
    for ri in range(len(tmpl["sample"]), len(tmpl["sample"])+15):
        for ci in range(len(headers)):
            ws.write(ri+4, ci, "", fmt_input)

    ws.freeze_panes(4, 0)
    wb.close()
    buf.seek(0)

    fname = f"TaxlyCMS_Upload_Template_{module.title()}.xlsx"
    return send_file(buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True, download_name=fname)


@app.route("/api/bulk/upload/<module>", methods=["POST"])
@login_required
def bulk_upload(module):
    """Process bulk upload Excel file."""
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename.endswith((".xlsx", ".xls")):
        return jsonify({"error": "Only .xlsx files supported"}), 400

    try:
        wb = openpyxl.load_workbook(f, data_only=True)
        ws = wb.active
        rows_data = list(ws.iter_rows(min_row=5, values_only=True))  # Skip 4 header rows
    except Exception as e:
        return jsonify({"error": f"Cannot read file: {str(e)}"}), 400

    conn = get_db(); c = conn.cursor()
    inserted = 0; skipped = 0; errors = []

    def safe(v, default=""):
        if v is None: return default
        s = str(v).strip()
        return s if s and s.lower() not in ("none","null","nan") else default

    def get_co_id(cin):
        c.execute("SELECT id FROM companies WHERE cin=%s", (cin,))
        r = c.fetchone()
        return (list(r.values())[0] if isinstance(r,dict) else r[0]) if r else None

    try:
        for ri, row in enumerate(rows_data, start=5):
            if not row or all(v is None or str(v).strip() == "" for v in row):
                continue  # skip empty rows

            try:
                if module == "companies":
                    name = safe(row[0])
                    if not name: skipped += 1; continue
                    c.execute("""INSERT INTO companies
                        (id,name,cin,company_type,pan,tan,gstin,email,phone,
                         incorporation_date,registered_office,authorized_capital,
                         paid_up_capital,business_activity,roc,created_by,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), name, _str(safe(row[1])),
                         safe(row[2],"Private Limited"),
                         _str((safe(row[3])).upper()), _str((safe(row[4])).upper()),
                         safe(row[5]), safe(row[6]), safe(row[7]),
                         _dt(safe(row[8])), safe(row[9]),
                         _num(safe(row[10])), _num(safe(row[11])),
                         safe(row[12]), safe(row[13]), g.user_id, g.tenant_id))
                    if c.rowcount: inserted += 1
                    else: skipped += 1

                elif module == "directors":
                    cin  = safe(row[0]); name = safe(row[1])
                    if not cin or not name: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    did = str(uuid.uuid4())
                    c.execute("""INSERT INTO directors
                        (id,company_id,name,din,pan,designation,email,mobile,
                         date_of_appointment,address,mca_user_id,mca_password,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (did, co_id, name, _str(safe(row[2])), _str((safe(row[3])).upper()),
                         safe(row[4],"Director"), safe(row[5]), safe(row[6]),
                         _dt(safe(row[7])), safe(row[8]), safe(row[9]), safe(row[10]), g.tenant_id))
                    nd = _kyc_due()
                    c.execute("INSERT INTO director_kyc (id,director_id,next_due_date,kyc_status,tenant_id) VALUES (%s,%s,%s,%s,%s)",
                              (str(uuid.uuid4()), did, nd, "pending", g.tenant_id))
                    inserted += 1

                elif module == "directors_kyc":
                    din = safe(row[0]); kyc_date = safe(row[1])
                    if not din or not kyc_date: skipped += 1; continue
                    c.execute("SELECT id FROM directors WHERE din=%s", (din,))
                    dr = c.fetchone()
                    if not dr:
                        errors.append(f"Row {ri}: DIN '{din}' not found"); skipped += 1; continue
                    dr_id = dr['id'] if isinstance(dr, dict) else dr[0]
                    c.execute("""UPDATE director_kyc SET last_kyc_date=%s,kyc_status='filed',
                                 next_due_date=%s,updated_at=NOW()
                                 WHERE director_id=%s""",
                              (_dt(kyc_date), _kyc_due(), dr_id))
                    if c.rowcount: inserted += 1
                    else: skipped += 1

                elif module == "shareholders":
                    cin  = safe(row[0]); name = safe(row[1])
                    if not cin or not name: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    c.execute("""INSERT INTO shareholders
                        (id,company_id,name,folio_no,pan,email,mobile,share_class,
                         shares_held,face_value,date_of_entry,address,mca_user_id,mca_password,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), co_id, name, safe(row[2]),
                         _str((safe(row[3])).upper()), safe(row[4]), safe(row[5]),
                         safe(row[6],"Equity"), _num(safe(row[7]),cast=int),
                         _num(safe(row[8]),default=10), _dt(safe(row[9])), safe(row[10]),
                         safe(row[11]), safe(row[12]), g.tenant_id))
                    inserted += 1

                elif module == "auditors":
                    cin  = safe(row[0]); name = safe(row[1])
                    if not cin or not name: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    c.execute("""INSERT INTO auditors
                        (id,company_id,name,firm_name,membership_no,frn,pan,email,
                         nature_of_appointment,appointment_type,start_date,end_date,srn_adt1)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), co_id, name, safe(row[2]), safe(row[3]),
                         safe(row[4]), _str((safe(row[5])).upper()), safe(row[6]),
                         safe(row[7],"Subsequent Auditor"), safe(row[8],"AGM Appointment"),
                         _dt(safe(row[9])), _dt(safe(row[10])), safe(row[11])))
                    inserted += 1

                elif module == "meetings":
                    cin  = safe(row[0]); mtype = safe(row[1]); mdate = safe(row[3])
                    if not cin or not mtype or not mdate: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    c.execute("""INSERT INTO meetings
                        (id,company_id,meeting_type,meeting_no,meeting_date,meeting_time,
                         venue,status,agenda,created_by,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), co_id, mtype, safe(row[2]), _dt(mdate),
                         safe(row[4],""), safe(row[5],""), safe(row[6],"scheduled"),
                         safe(row[7]), g.user_id, g.tenant_id))
                    inserted += 1

                elif module == "tasks":
                    title = safe(row[0])
                    if not title: skipped += 1; continue
                    cin   = safe(row[1])
                    co_id = get_co_id(cin) if cin else None
                    c.execute("""INSERT INTO tasks
                        (id,company_id,title,priority,status,due_date,module,description,created_by,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), _str(co_id), title,
                         safe(row[2],"medium"), safe(row[3],"pending"),
                         _dt(safe(row[4])), safe(row[5]), safe(row[6]), g.user_id, g.tenant_id))
                    inserted += 1

                elif module == "charges":
                    cin    = safe(row[0]); holder = safe(row[1])
                    if not cin or not holder: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    c.execute("""INSERT INTO charges
                        (id,company_id,charge_holder,charge_type,amount,date_of_creation,
                         assets_charged,status,charge_id,srn,created_by,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), co_id, holder,
                         safe(row[2],"Hypothecation"), _num(safe(row[3])),
                         _dt(safe(row[4])), safe(row[5]), safe(row[6],"Open"),
                         safe(row[7]), safe(row[8]), g.user_id, g.tenant_id))
                    inserted += 1

                elif module == "investments":
                    cin  = safe(row[0]); name = safe(row[1])
                    if not cin or not name: skipped += 1; continue
                    co_id = get_co_id(cin)
                    if not co_id:
                        errors.append(f"Row {ri}: Company CIN '{cin}' not found"); skipped += 1; continue
                    c.execute("""INSERT INTO investments
                        (id,company_id,investee_name,investee_type,investment_type,amount,
                         date_of_investment,board_resolution_date,purpose,created_by,tenant_id)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (str(uuid.uuid4()), co_id, name,
                         safe(row[2],"Company"), safe(row[3],"Equity Shares"),
                         _num(safe(row[4])), _dt(safe(row[5])),
                         _dt(safe(row[6])), safe(row[7]), g.user_id, g.tenant_id))
                    inserted += 1

                else:
                    conn.close()
                    return jsonify({"error": f"Upload not supported for module: {module}"}), 400

            except Exception as row_err:
                errors.append(f"Row {ri}: {str(row_err)}")
                skipped += 1

        conn.commit()
    except Exception as e:
        conn.rollback()
        conn.close()
        return jsonify({"error": str(e)}), 500

    conn.close()
    return jsonify({
        "success": True,
        "inserted": inserted,
        "skipped": skipped,
        "errors": errors[:20],  # cap error list
        "message": f"✓ {inserted} record(s) uploaded successfully. {skipped} skipped."
    })


# ══ SMART DOCUMENT GENERATION WITH PERIOD RANGE ══════════════════════════════

@app.route("/api/documents/generate", methods=["POST"])
@login_required
def generate_doc():
    """Generate document with period-range-aware smart placeholders."""
    d = request.get_json(silent=True, force=True) or {}
    tid   = d.get("template_id")
    cid   = d.get("company_id")
    d_id  = d.get("director_id")
    d2_id = d.get("director2_id")
    a_id  = d.get("auditor_id")
    extra = d.get("extra_context") or {}
    date_from = d.get("date_from","")
    date_to   = d.get("date_to","")
    rich_mode = d.get("rich_mode", False)  # True = return HTML, False = plain text

    if not tid or not cid:
        return jsonify({"error":"template_id and company_id required"}),400

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT template_body,name,category FROM document_templates WHERE id=%s", (tid,))
    tpl = c.fetchone()
    if not tpl: conn.close(); return jsonify({"error":"Template not found"}),404

    body   = (tpl['template_body'] if isinstance(tpl,dict) else tpl[0]) or ""
    tname  = tpl['name']     if isinstance(tpl,dict) else tpl[1]
    tcat   = tpl['category'] if isinstance(tpl,dict) else tpl[2]

    # ── Build base CRM context ─────────────────────────────────────────────
    from compliance import build_context
    try:
        ctx = build_context(cid, extra=extra,
                            director_id=d_id, director2_id=d2_id, auditor_id=a_id)
    except Exception as e:
        conn.close(); return jsonify({"error": str(e)}), 400

    # ── Smart dynamic placeholders with period range ───────────────────────
    # board_meetings_list / board_meetings_numbered / bm_dates_only
    mtg_params = [cid]
    mtg_sql = """SELECT meeting_type,meeting_no,meeting_date,meeting_time,venue,status
                 FROM meetings WHERE company_id=%s"""
    if date_from: mtg_sql += " AND meeting_date>=%s"; mtg_params.append(date_from)
    if date_to:   mtg_sql += " AND meeting_date<=%s"; mtg_params.append(date_to)
    mtg_sql += " ORDER BY meeting_date"
    c.execute(mtg_sql, mtg_params)
    meetings = rows(c.fetchall())

    # Build meeting smart placeholders
    bm_list = [m for m in meetings if 'board' in m['meeting_type'].lower()]
    agm_list = [m for m in meetings if 'agm' in m['meeting_type'].lower() or 'annual' in m['meeting_type'].lower()]
    all_list = meetings

    def fmt_meeting(i, m):
        raw = m.get('meeting_date','')
        try:
            d_obj = _to_date(raw)
            d_str = d_obj.strftime('%d %B %Y') if d_obj else (str(raw) if raw else '')
        except: d_str = str(raw) if raw else ''
        return f"  BM-{i}: {d_str}" + (f" at {m['meeting_time']}" if m.get('meeting_time') else "") + (f", {m['venue']}" if m.get('venue') else "")

    NL = chr(10)
    ctx["board_meetings_list"]     = NL.join(fmt_meeting(i+1, m) for i,m in enumerate(bm_list)) or "[No board meetings in this period]"
    ctx["board_meetings_numbered"] = NL.join(f"{i+1}. {_ds(m.get('meeting_date')) or ''}" for i,m in enumerate(bm_list)) or "[No board meetings]"
    ctx["board_meetings_count"]    = str(len(bm_list))
    ctx["all_meetings_list"]       = NL.join(fmt_meeting(i+1, m) for i,m in enumerate(all_list)) or "[No meetings in this period]"
    ctx["agm_meetings_list"]       = NL.join(fmt_meeting(i+1, m) for i,m in enumerate(agm_list)) or "[No AGMs in this period]"

    # Individual BM slots — unlimited, up to however many exist
    for i, m in enumerate(bm_list, 1):
        try:
            d_obj = _to_date(m.get('meeting_date',''))
            d_str = d_obj.strftime('%d %B %Y') if d_obj else str(m.get('meeting_date',''))
        except: d_str = str(m.get('meeting_date',''))
        ctx[f"board_meeting_{i}"] = d_str
        ctx[f"bm_{i}_date"]       = d_str
        ctx[f"bm_{i}_venue"]      = m.get('venue','')
        ctx[f"bm_{i}_time"]       = m.get('meeting_time','')

    # DSC records in period
    dsc_params = [cid]
    dsc_sql = "SELECT holder_name,holder_type,valid_to FROM dsc_records WHERE company_id=%s AND is_active=1"
    if date_to: dsc_sql += " AND valid_to>=%s"; dsc_params.append(date_to)
    c.execute(dsc_sql, dsc_params)
    dsc_rows = c.fetchall()
    ctx["dsc_list"] = NL.join(
        f"  {(_rv(r,'holder_name') or _rv(r,0))} ({(_rv(r,'dsc_class') or _rv(r,1))}) — valid to {(_rv(r,'valid_to') or _rv(r,2))}"
        for r in dsc_rows) or "[No DSC records]"

    # Charges
    chrg_params = [cid]
    chrg_sql = "SELECT charge_holder,charge_type,amount,date_of_creation FROM charges WHERE company_id=%s"
    if date_from: chrg_sql += " AND date_of_creation>=%s"; chrg_params.append(date_from)
    if date_to:   chrg_sql += " AND date_of_creation<=%s"; chrg_params.append(date_to)
    c.execute(chrg_sql, chrg_params)
    charges = c.fetchall()
    ctx["charges_list"] = NL.join(
        f"  {(_rv(r,'charge_holder') or _rv(r,0))} — {(_rv(r,'charge_type') or _rv(r,1))} — ₹{float(_rv(r,'amount') or _rv(r,2) or 0):,.0f} (created {(_rv(r,'date_of_creation') or _rv(r,3))})"
        for r in charges) or "[No charges in this period]"

    # Director KYC status
    c.execute("""SELECT d.name,k.kyc_status,k.next_due_date
                 FROM directors d LEFT JOIN director_kyc k ON d.id=k.director_id
                 WHERE d.company_id=%s AND d.is_active=1""", (cid,))
    kyc_rows = c.fetchall()
    ctx["director_kyc_list"] = NL.join(
        f"  {(_rv(r,'name') or _rv(r,0))}: {(_rv(r,'kyc_status') or _rv(r,1) or 'pending')} (due {(_rv(r,'next_due_date') or _rv(r,2) or 'N/A')})"
        for r in kyc_rows) or "[No directors]"

    ctx["period_from"] = date_from or ctx.get("financial_year","")
    ctx["period_to"]   = date_to   or ctx.get("year_ended_on","")

    conn.close()

    # ── Resolve all {{placeholders}} in body ──────────────────────────────
    import re as _re
    def resolve(body, ctx):
        def replace(m):
            key = m.group(1).strip()
            val = ctx.get(key)
            if val is not None:
                return str(val)
            return f"[{key} — TO BE FILLED]"
        return _re.sub(r'\{\{([^}]+)\}\}', replace, body)

    content = resolve(body, ctx)

    # Save to documents table
    c2 = get_db(); c3 = c2.cursor()
    c3.execute("""INSERT INTO documents
        (id,company_id,template_id,doc_name,doc_type,content,generated_by,date_from,date_to,tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
        (str(uuid.uuid4()), cid, tid, tname, tcat, content,
         g.user_id, date_from, date_to, g.tenant_id))
    c2.commit(); c2.close()

    return jsonify({"content": content, "context": {k:v for k,v in ctx.items() if isinstance(v,str) and len(v)<200}})


# (duplicate PDF route removed — see gen_document_pdf above)


@app.route("/api/custom-placeholders", methods=["GET"])
@login_required
def list_custom_placeholders():
    """List all manually defined custom placeholders."""
    conn = get_db(); c = conn.cursor()
    try:
        c.execute("SELECT * FROM custom_placeholders ORDER BY name")
        result = rows(c.fetchall())
    except:
        result = []
    conn.close()
    return jsonify(result)

@app.route("/api/custom-placeholders", methods=["POST"])
@login_required
def create_custom_placeholder():
    """Create a custom placeholder."""
    d = request.get_json(silent=True, force=True) or {}
    name  = d.get("name","").strip().lower().replace(" ","_")
    label = d.get("label","").strip()
    desc  = d.get("description","").strip()
    ph_type = d.get("ph_type","text")
    if not name or not label:
        return jsonify({"error":"name and label required"}),400
    conn = get_db(); c = conn.cursor()
    try:
        c.execute("""INSERT INTO custom_placeholders (id,name,label,description,ph_type,created_by)
                     VALUES (%s,%s,%s,%s,%s,%s)""",
                  (str(uuid.uuid4()), name, label, desc, ph_type, g.user_id))
        conn.commit()
    except Exception as e:
        conn.close(); return jsonify({"error":str(e)}),400
    conn.close()
    return jsonify({"success":True,"name":name})

@app.route("/api/custom-placeholders/<pid>", methods=["DELETE"])
@login_required
def delete_custom_placeholder(pid):
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM custom_placeholders WHERE id=%s", (pid,))
    conn.commit(); conn.close()
    return jsonify({"success":True})


# ══ USER PERMISSION MANAGEMENT ═══════════════════════════════════════════════

@app.route("/api/permissions/schema")
@login_required
def get_perm_schema():
    """Return all modules and their actions for UI rendering."""
    return jsonify({
        "modules": ALL_MODULES,
        "defaults": DEFAULT_PERMISSIONS,
    })

@app.route("/api/users/<uid>/permissions")
@login_required
def get_user_permissions(uid):
    """Get all permission overrides for a specific user."""
    if g.role != "superadmin":
        return jsonify({"error": "Only Super Admin can view permissions"}), 403
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM users WHERE id=%s", (uid,))
    user = row(c.fetchone())
    if not user: conn.close(); return jsonify({"error": "User not found"}), 404
    c.execute("""SELECT module, action, granted, granted_by, granted_at, note
                 FROM user_permissions WHERE user_id=%s ORDER BY module, action""", (uid,))
    overrides = rows(c.fetchall())
    conn.close()
    return jsonify({
        "user": user,
        "overrides": overrides,
        "role": user["role"],
    })

@app.route("/api/users/<uid>/permissions", methods=["POST"])
@login_required
def set_user_permissions(uid):
    """Set/update permission overrides for a user (superadmin only)."""
    if g.role != "superadmin":
        return jsonify({"error": "Only Super Admin can modify permissions"}), 403
    d = request.get_json(silent=True, force=True) or {}
    # d = { "permissions": [ {"module": "company", "action": "delete", "granted": true/false, "note": ""}, ... ] }
    permissions = d.get("permissions", [])
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT id, role FROM users WHERE id=%s", (uid,))
    user = c.fetchone()
    if not user: conn.close(); return jsonify({"error": "User not found"}), 404
    # Cannot set permissions for superadmin
    _urole = user.get('role') if isinstance(user,dict) else user[1]
    if _urole == "superadmin":
        conn.close(); return jsonify({"error": "Cannot restrict Super Admin permissions"}), 400
    for perm in permissions:
        module  = perm.get("module","").strip()
        action  = perm.get("action","").strip()
        granted = 1 if perm.get("granted", True) else 0
        note    = perm.get("note", "")
        if not module or not action: continue
        c.execute("""INSERT INTO user_permissions (id, user_id, module, action, granted, granted_by, note)
                     VALUES (%s, %s, %s, %s, %s, %s, %s)
                     ON CONFLICT(user_id, module, action)
                     DO UPDATE SET granted=excluded.granted, granted_by=excluded.granted_by,
                                   note=excluded.note, granted_at=CURRENT_TIMESTAMP""",
                  (str(uuid.uuid4()), uid, module, action, granted, g.user_id, note))
    conn.commit(); conn.close()
    return jsonify({"success": True, "updated": len(permissions)})

@app.route("/api/users/<uid>/permissions/reset", methods=["POST"])
@login_required
def reset_user_permissions(uid):
    """Remove ALL overrides for a user (revert to role defaults)."""
    if g.role != "superadmin":
        return jsonify({"error": "Only Super Admin can reset permissions"}), 403
    d = request.get_json(silent=True, force=True) or {}
    module = d.get("module")  # optional: reset only one module
    conn = get_db(); c = conn.cursor()
    if module:
        c.execute("DELETE FROM user_permissions WHERE user_id=%s AND module=%s", (uid, module))
    else:
        c.execute("DELETE FROM user_permissions WHERE user_id=%s", (uid,))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/permission-presets", methods=["GET"])
@login_required
def list_presets():
    """List all saved permission presets."""
    conn = get_db(); c = conn.cursor()
    try:
        c.execute("SELECT * FROM permission_presets ORDER BY name")
        result = rows(c.fetchall())
    except: result = []
    conn.close()
    return jsonify(result)

@app.route("/api/permission-presets", methods=["POST"])
@login_required
def create_preset():
    """Save current permission set as a named preset."""
    if g.role != "superadmin":
        return jsonify({"error": "Only Super Admin can create presets"}), 403
    d = request.get_json(silent=True, force=True) or {}
    name = d.get("name","").strip()
    if not name: return jsonify({"error": "Preset name required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO permission_presets (id, name, description, permissions, created_by)
                 VALUES (%s,%s,%s,%s,%s)""",
              (str(uuid.uuid4()), name, d.get("description",""),
               __import__("json").dumps(d.get("permissions",[])), g.user_id))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/permission-presets/<pid>", methods=["DELETE"])
@login_required
def delete_preset(pid):
    if g.role != "superadmin": return jsonify({"error":"Forbidden"}), 403
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM permission_presets WHERE id=%s", (pid,))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/users/<uid>/permissions/apply-preset/<pid>", methods=["POST"])
@login_required
def apply_preset(uid, pid):
    """Apply a named preset to a user."""
    if g.role != "superadmin": return jsonify({"error":"Forbidden"}), 403
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT permissions FROM permission_presets WHERE id=%s", (pid,))
    preset = c.fetchone()
    if not preset: conn.close(); return jsonify({"error":"Preset not found"}),404
    import json as _json
    _preset_perms = preset.get('permissions') if isinstance(preset,dict) else preset[0]
    permissions = _json.loads(_preset_perms) if _preset_perms else []
    for perm in permissions:
        module = perm.get("module",""); action = perm.get("action","")
        granted = 1 if perm.get("granted", True) else 0
        if not module or not action: continue
        c.execute("""INSERT INTO user_permissions (id,user_id,module,action,granted,granted_by,note)
                     VALUES (%s,%s,%s,%s,%s,%s,%s)
                     ON CONFLICT(user_id,module,action)
                     DO UPDATE SET granted=excluded.granted, granted_by=excluded.granted_by,
                                   granted_at=CURRENT_TIMESTAMP""",
                  (str(uuid.uuid4()),uid,module,action,granted,g.user_id,f"Applied preset: {pid}"))
    conn.commit(); conn.close()
    return jsonify({"success":True,"applied":len(permissions)})


# ══ PERSONAL DASHBOARD ════════════════════════════════════════════════════════

@app.route("/api/my-dashboard")
@login_required
def my_dashboard():
    """Role-aware personal dashboard.
    
    Role mapping:
      superadmin → Task Leader   (task_leader = me)
      manager    → Task Manager  (task_manager = me)
      staff      → Assignee      (assigned_to = me)
    
    Each panel shows OTHER roles' names for context.
    """
    uid  = g.user_id
    role = g.role
    conn = get_db(); c = conn.cursor()
    from datetime import date, timedelta
    today = date.today()
    week  = today + timedelta(days=7)

    # Column filter based on role — whitelisted to prevent injection
    MY_COL_MAP = {"superadmin": "task_leader", "manager": "task_manager"}
    my_col = MY_COL_MAP.get(role, "assigned_to")  # safe whitelist

    # Task summary counts (only my tasks per role)
    c.execute(f"""SELECT status, COUNT(*) as cnt FROM tasks WHERE {my_col}=%s GROUP BY status""", (uid,))
    task_counts = {}
    for _r in c.fetchall():
        if isinstance(_r, dict):
            task_counts[_r.get('status', list(_r.values())[0])] = _r.get('cnt', list(_r.values())[1])
        else:
            task_counts[_r[0]] = _r[1]

    c.execute(f"""SELECT COUNT(*) FROM tasks WHERE {my_col}=%s
                  AND due_date=%s AND status NOT IN ('completed','cancelled')""",
              (uid, today))
    due_today = _count(c)

    c.execute(f"""SELECT COUNT(*) FROM tasks WHERE {my_col}=%s
                  AND due_date BETWEEN %s AND %s AND status NOT IN ('completed','cancelled')""",
              (uid, today, week))
    due_week = _count(c)

    c.execute(f"""SELECT COUNT(*) FROM tasks WHERE {my_col}=%s
                  AND due_date < %s AND status NOT IN ('completed','cancelled')""",
              (uid, today))
    overdue = _count(c)

    # Full task list — only tasks where I play MY role
    # Always join all three roles so frontend can show "who else is on this task"
    task_sql = f"""
        SELECT t.*,
               co.name  AS company_name,
               ul.name  AS task_leader_name,
               um.name  AS task_manager_name,
               ua.name  AS assigned_to_name,
               ul.id    AS leader_id,
               um.id    AS manager_id,
               ua.id    AS assignee_id
        FROM tasks t
        LEFT JOIN companies co ON t.company_id = co.id
        LEFT JOIN users ul ON t.task_leader  = ul.id
        LEFT JOIN users um ON t.task_manager = um.id
        LEFT JOIN users ua ON t.assigned_to  = ua.id
        WHERE t.{my_col} = %s
        ORDER BY
            CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2
                            WHEN 'medium'   THEN 3 ELSE 4 END,
            t.due_date
    """
    c.execute(task_sql, (uid,))
    my_tasks = rows(c.fetchall())

    # ── Role-wise team summary (superadmin + manager see this) ──────────────
    tid = g.tenant_id
    role_summary = []
    if role in ("superadmin", "manager"):
        if tid:
            c.execute("""
                SELECT u.id, u.name, u.role,
                       COUNT(CASE WHEN t.status='pending'     THEN 1 END) AS pending,
                       COUNT(CASE WHEN t.status='in_progress' THEN 1 END) AS in_progress,
                       COUNT(CASE WHEN t.status='completed'   THEN 1 END) AS completed,
                       COUNT(CASE WHEN t.due_date < CURRENT_DATE AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS overdue
                FROM users u
                LEFT JOIN tasks t ON (
                    (u.role='superadmin' AND t.task_leader=u.id) OR
                    (u.role='manager'    AND t.task_manager=u.id) OR
                    (u.role='staff'      AND t.assigned_to=u.id)
                )
                WHERE u.is_active=1
                  AND u.is_platform_admin=0
                  AND (u.tenant_id=%s OR u.tenant_id IS NULL)
                GROUP BY u.id
                ORDER BY CASE u.role WHEN 'superadmin' THEN 1 WHEN 'manager' THEN 2 ELSE 3 END, u.name
            """, (tid,))
        else:
            c.execute("""
                SELECT u.id, u.name, u.role,
                       COUNT(CASE WHEN t.status='pending'     THEN 1 END) AS pending,
                       COUNT(CASE WHEN t.status='in_progress' THEN 1 END) AS in_progress,
                       COUNT(CASE WHEN t.status='completed'   THEN 1 END) AS completed,
                       COUNT(CASE WHEN t.due_date < CURRENT_DATE AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS overdue
                FROM users u
                LEFT JOIN tasks t ON (
                    (u.role='superadmin' AND t.task_leader=u.id) OR
                    (u.role='manager'    AND t.task_manager=u.id) OR
                    (u.role='staff'      AND t.assigned_to=u.id)
                )
                WHERE u.is_active=1 AND u.is_platform_admin=0
                GROUP BY u.id
                ORDER BY CASE u.role WHEN 'superadmin' THEN 1 WHEN 'manager' THEN 2 ELSE 3 END, u.name
            """)
        role_summary = rows(c.fetchall())

    conn.close()
    return jsonify({
        "my_role":    role,
        "my_col":     my_col,
        "summary": {
            "pending":     task_counts.get("pending", 0),
            "in_progress": task_counts.get("in_progress", 0),
            "completed":   task_counts.get("completed", 0),
            "due_today":   due_today,
            "due_week":    due_week,
            "overdue":     overdue,
        },
        "my_tasks":     my_tasks,
        "role_summary": role_summary,
    })


# ══ ROLE-WISE TASK WIDGET ═════════════════════════════════════════════════════

@app.route("/api/tasks/rolewise-summary")
@login_required
def tasks_rolewise_summary():
    """
    Returns task counts for the logged-in user broken down by role:
      leader_cnt   = tasks where task_leader  = me  (not done)
      manager_cnt  = tasks where task_manager = me  (not done)
      assignee_cnt = tasks where assigned_to  = me  (not done)
      total        = distinct tasks where I appear in any role
    """
    uid  = g.user_id
    role = g.role
    _tid = getattr(g, 'tenant_id', None)
    empty = {"uid": uid, "my_role": role,
             "totals": {"leader":0,"manager":0,"assignee":0,"total":0},
             "modules": [], "persons": [], "debug": ""}
    try:
        conn = get_db(); c = conn.cursor()
        EXCL = ('completed', 'cancelled')

        # ── Per-module breakdown ──────────────────────────────────────────────
        c.execute("""
            SELECT
                COALESCE(module, 'general')                                         AS module,
                COUNT(CASE WHEN task_leader  = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS leader_cnt,
                COUNT(CASE WHEN task_manager = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS manager_cnt,
                COUNT(CASE WHEN assigned_to  = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS assignee_cnt,
                COUNT(CASE WHEN (task_leader = %s OR task_manager = %s OR assigned_to = %s)
                               AND status NOT IN ('completed','cancelled') THEN 1 END) AS total
            FROM tasks
            WHERE (task_leader = %s OR task_manager = %s OR assigned_to = %s)
              AND status NOT IN ('completed','cancelled')
            GROUP BY COALESCE(module, 'general')
            ORDER BY total DESC
        """, (uid, uid, uid, uid, uid, uid, uid, uid, uid))

        # ── Parse module rows ─────────────────────────────────────────────────
        modules = []
        for _r in c.fetchall():
            if isinstance(_r, dict):
                _mod = _r.get('module') or 'general'
                _l   = int(_r.get('leader_cnt')   or 0)
                _m   = int(_r.get('manager_cnt')  or 0)
                _a   = int(_r.get('assignee_cnt') or 0)
                _t   = int(_r.get('total')        or 0)
            else:
                _mod = _r[0] or 'general'
                _l   = int(_r[1] or 0)
                _m   = int(_r[2] or 0)
                _a   = int(_r[3] or 0)
                _t   = int(_r[4] or 0)
            modules.append({"module":_mod,"leader":_l,"manager":_m,"assignee":_a,"total":_t})

        # ── Top-level totals ──────────────────────────────────────────────────
        c.execute("""
            SELECT
                COUNT(CASE WHEN task_leader  = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS leader_cnt,
                COUNT(CASE WHEN task_manager = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS manager_cnt,
                COUNT(CASE WHEN assigned_to  = %s AND status NOT IN ('completed','cancelled') THEN 1 END) AS assignee_cnt,
                COUNT(CASE WHEN (task_leader = %s OR task_manager = %s OR assigned_to = %s)
                               AND status NOT IN ('completed','cancelled') THEN 1 END) AS total_cnt
            FROM tasks
        """, (uid, uid, uid, uid, uid, uid))
        _tr = c.fetchone()
        if _tr:
            if isinstance(_tr, dict):
                totals = {
                    "leader":   int(_tr.get('leader_cnt')   or 0),
                    "manager":  int(_tr.get('manager_cnt')  or 0),
                    "assignee": int(_tr.get('assignee_cnt') or 0),
                    "total":    int(_tr.get('total_cnt')    or 0),
                }
            else:
                totals = {
                    "leader":   int(_tr[0] or 0),
                    "manager":  int(_tr[1] or 0),
                    "assignee": int(_tr[2] or 0),
                    "total":    int(_tr[3] or 0),
                }
        else:
            totals = {"leader":0,"manager":0,"assignee":0,"total":0}

        # ── Person strip: all active users + their task counts ────────────────
        # Shows team members for context (leader sees full team, others see peers)
        q_persons = """
            SELECT u.id, u.name, u.role,
                COUNT(CASE WHEN t.task_leader=u.id   AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS as_leader,
                COUNT(CASE WHEN t.task_manager=u.id  AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS as_manager,
                COUNT(CASE WHEN t.assigned_to=u.id   AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS as_assignee,
                COUNT(CASE WHEN (t.task_leader=u.id OR t.task_manager=u.id OR t.assigned_to=u.id)
                               AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS pending,
                COUNT(CASE WHEN (t.task_leader=u.id OR t.task_manager=u.id OR t.assigned_to=u.id)
                               AND t.status='completed' THEN 1 END) AS completed
            FROM users u
            LEFT JOIN tasks t ON (t.task_leader=u.id OR t.task_manager=u.id OR t.assigned_to=u.id)
            WHERE u.is_active=1 AND u.is_platform_admin=0
        """
        if _tid:
            c.execute(q_persons + " AND (u.tenant_id=%s OR u.tenant_id IS NULL) GROUP BY u.id,u.name,u.role ORDER BY CASE u.role WHEN 'superadmin' THEN 1 WHEN 'manager' THEN 2 ELSE 3 END, pending DESC", (_tid,))
        else:
            c.execute(q_persons + " GROUP BY u.id,u.name,u.role ORDER BY CASE u.role WHEN 'superadmin' THEN 1 WHEN 'manager' THEN 2 ELSE 3 END, pending DESC")
        persons = rows(c.fetchall())

        conn.close()
        return jsonify({
            "uid":     uid,
            "my_role": role,
            "totals":  totals,
            "modules": modules,
            "persons": persons,
        })
    except Exception as _ex:
        import traceback as _tb
        _tb_str = _tb.format_exc()
        app.logger.error(f"rolewise-summary error: {_ex}\n{_tb_str}")
        empty["debug"] = str(_ex)
        return jsonify(empty)

@app.route("/api/tasks/by-role-module")
@login_required
def tasks_by_role_module():
    """Return tasks filtered by module + my role for the drill-down page."""
    uid    = g.user_id
    module = request.args.get("module","")
    role   = request.args.get("role","")    # leader | manager | assignee | all
    status = request.args.get("status","")

    conn = get_db(); c = conn.cursor()
    sql = """
        SELECT t.*,
               co.name AS company_name,
               ul.name AS task_leader_name,
               um.name AS task_manager_name,
               ua.name AS assigned_to_name
        FROM tasks t
        LEFT JOIN companies co ON t.company_id=co.id
        LEFT JOIN users ul ON t.task_leader=ul.id
        LEFT JOIN users um ON t.task_manager=um.id
        LEFT JOIN users ua ON t.assigned_to=ua.id
        WHERE 1=1
    """
    params = []
    if role == "leader":   sql += " AND t.task_leader=%s";  params.append(uid)
    elif role == "manager": sql += " AND t.task_manager=%s"; params.append(uid)
    elif role == "assignee":sql += " AND t.assigned_to=%s";  params.append(uid)
    else:                  sql += " AND (t.task_leader=%s OR t.task_manager=%s OR t.assigned_to=%s)"; params.extend([uid,uid,uid])
    if module and module != 'all': sql += " AND COALESCE(t.module,'general')=%s"; params.append(module)
    if status: sql += " AND t.status=%s"; params.append(status)
    _tsk_q = request.args.get("q","").strip()
    if _tsk_q:
        sql += " AND (t.title ILIKE %s OR co.name ILIKE %s OR t.description ILIKE %s)"
        _like = f"%{_tsk_q}%"
        params.extend([_like, _like, _like])
    sql += " ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2 WHEN 'medium' THEN 3 ELSE 4 END, t.due_date"

    c.execute(sql, params)
    tasks = rows(c.fetchall())
    conn.close()
    return jsonify(tasks)


@app.route("/api/tasks/subordinate-view")
@login_required  
def tasks_subordinate_view():
    """
    Superior view: I can see ALL tasks in my chain of command.
    - If I am leader (superadmin): see all tasks I lead + their managers + assignees
    - If I am manager: see all tasks I manage + their assignees
    - Also returns subordinate pending summary grouped by person
    """
    uid  = g.user_id
    role = g.role
    conn = get_db(); c = conn.cursor()

    if role == "superadmin":
        # I am leader — see all tasks where task_leader = me
        c.execute("""
            SELECT t.*,
                   co.name AS company_name,
                   ul.name AS task_leader_name,
                   um.name AS task_manager_name, um.id AS manager_id,
                   ua.name AS assigned_to_name,  ua.id AS assignee_id
            FROM tasks t
            LEFT JOIN companies co ON t.company_id=co.id
            LEFT JOIN users ul ON t.task_leader=ul.id
            LEFT JOIN users um ON t.task_manager=um.id
            LEFT JOIN users ua ON t.assigned_to=ua.id
            WHERE t.task_leader=%s
            ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2
                                     WHEN 'medium' THEN 3 ELSE 4 END, t.due_date
        """, (uid,))
    elif role == "manager":
        # I am manager — see all tasks where task_manager = me
        c.execute("""
            SELECT t.*,
                   co.name AS company_name,
                   ul.name AS task_leader_name, ul.id AS leader_id,
                   um.name AS task_manager_name,
                   ua.name AS assigned_to_name,  ua.id AS assignee_id
            FROM tasks t
            LEFT JOIN companies co ON t.company_id=co.id
            LEFT JOIN users ul ON t.task_leader=ul.id
            LEFT JOIN users um ON t.task_manager=um.id
            LEFT JOIN users ua ON t.assigned_to=ua.id
            WHERE t.task_manager=%s
            ORDER BY CASE t.priority WHEN 'critical' THEN 1 WHEN 'high' THEN 2
                                     WHEN 'medium' THEN 3 ELSE 4 END, t.due_date
        """, (uid,))
    else:
        conn.close()
        return jsonify({"tasks": [], "subordinates": [], "my_role": role})

    tasks = rows(c.fetchall())

    # Subordinate summary — people below me and their pending counts
    if role == "superadmin":
        # Subordinates = all managers + assignees on my tasks
        c.execute("""
            SELECT u.id, u.name, u.role,
                   COUNT(CASE WHEN t.status NOT IN ('completed','cancelled') THEN 1 END) AS pending,
                   COUNT(CASE WHEN t.status='completed' THEN 1 END) AS completed,
                   COUNT(CASE WHEN t.due_date < CURRENT_DATE AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS overdue,
                   COUNT(*) AS total
            FROM users u
            JOIN tasks t ON (t.task_manager=u.id OR t.assigned_to=u.id)
            WHERE t.task_leader=%s AND u.id != %s AND u.is_active=1
            GROUP BY u.id
            ORDER BY pending DESC
        """, (uid, uid))
    else:  # manager
        # Subordinates = assignees on tasks I manage
        c.execute("""
            SELECT u.id, u.name, u.role,
                   COUNT(CASE WHEN t.status NOT IN ('completed','cancelled') THEN 1 END) AS pending,
                   COUNT(CASE WHEN t.status='completed' THEN 1 END) AS completed,
                   COUNT(CASE WHEN t.due_date < CURRENT_DATE AND t.status NOT IN ('completed','cancelled') THEN 1 END) AS overdue,
                   COUNT(*) AS total
            FROM users u
            JOIN tasks t ON t.assigned_to=u.id
            WHERE t.task_manager=%s AND u.id != %s AND u.is_active=1
            GROUP BY u.id
            ORDER BY pending DESC
        """, (uid, uid))

    subordinates = rows(c.fetchall())

    conn.close()
    return jsonify({
        "my_role":     role,
        "tasks":       tasks,
        "subordinates": subordinates,
        "totals": {
            "total":     len(tasks),
            "pending":   sum(1 for t in tasks if t["status"] == "pending"),
            "in_progress": sum(1 for t in tasks if t["status"] == "in_progress"),
            "completed": sum(1 for t in tasks if t["status"] == "completed"),
            "overdue":   sum(1 for t in tasks if t.get("due_date","") and t.get("due_date","") < str(date.today()) and t["status"] not in ("completed","cancelled")),
        }
    })


# ══ PLATFORM ADMIN — TENANT & PLAN MANAGEMENT ════════════════════════════════

@app.route("/api/platform/plans")
@platform_admin_required
def list_plans():
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT * FROM plans ORDER BY price_monthly")
    return jsonify(rows(c.fetchall()))

@app.route("/api/platform/plans", methods=["POST"])
@platform_admin_required
def create_plan():
    d=request.get_json(silent=True,force=True) or {}
    if not d.get("name"): return jsonify({"error":"Name required"}),400
    pid=str(uuid.uuid4()); conn=get_db(); c=conn.cursor()
    c.execute("""INSERT INTO plans (id,name,description,price_monthly,price_annual,
                max_users,max_companies,max_documents,features,is_active)
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,1)""",
             (pid,d["name"],d.get("description",""),
              float(d.get("price_monthly",0)),float(d.get("price_annual",0)),
              int(d.get("max_users",5)),int(d.get("max_companies",5)),
              int(d.get("max_documents",100)),
              __import__("json").dumps(d.get("features",[]))))
    conn.commit(); c.execute("SELECT * FROM plans WHERE id=%s", (pid,))
    result=row(c.fetchone()); conn.close(); return jsonify(result),201

@app.route("/api/platform/plans/<pid>", methods=["PUT"])
@platform_admin_required
def update_plan(pid):
    d=request.get_json(silent=True,force=True) or {}
    conn=get_db(); c=conn.cursor()
    fields={k:d[k] for k in ["name","description","price_monthly","price_annual",
                               "max_users","max_companies","max_documents","is_active"] if k in d}
    if "features" in d: fields["features"]=__import__("json").dumps(d["features"])
    if not fields: return jsonify({"error":"Nothing to update"}),400
    c.execute(f"UPDATE plans SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",
              list(fields.values())+[pid])
    conn.commit(); c.execute("SELECT * FROM plans WHERE id=%s", (pid,))
    result=row(c.fetchone()); conn.close(); return jsonify(result)

@app.route("/api/platform/plans/<pid>", methods=["DELETE"])
@platform_admin_required
def delete_plan(pid):
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT COUNT(*) FROM tenants WHERE plan_id=%s", (pid,))
    if _count(c): conn.close(); return jsonify({"error":"Plan in use by tenants"}),409
    c.execute("DELETE FROM plans WHERE id=%s", (pid,))
    conn.commit(); conn.close(); return jsonify({"success":True})

@app.route("/api/platform/tenants")
@platform_admin_required
def list_tenants():
    conn=get_db(); c=conn.cursor()
    c.execute("""SELECT t.*, p.name as plan_name, p.price_monthly,
                 (SELECT COUNT(*) FROM users u WHERE u.tenant_id=t.id AND u.is_active=1) as user_count,
                 (SELECT COUNT(*) FROM companies co WHERE co.tenant_id=t.id) as company_count
                 FROM tenants t LEFT JOIN plans p ON t.plan_id=p.id
                 ORDER BY t.created_at DESC""")
    return jsonify(rows(c.fetchall()))

@app.route("/api/platform/tenants", methods=["POST"])
@platform_admin_required
def create_tenant():
    d=request.get_json(silent=True,force=True) or {}
    if not d.get("name"):  return jsonify({"error":"Name required"}),400
    if not d.get("email"): return jsonify({"error":"Email required"}),400
    if not d.get("admin_name"): return jsonify({"error":"Admin name required"}),400
    if not d.get("admin_password"): return jsonify({"error":"Admin password required"}),400

    conn=get_db(); c=conn.cursor()
    # Get plan limits
    plan_id=d.get("plan_id")
    max_u=5; max_co=5
    if plan_id:
        c.execute("SELECT max_users,max_companies FROM plans WHERE id=%s", (plan_id,))
        pl=c.fetchone()
        if pl: max_u,max_co=(pl['max_users'] if isinstance(pl,dict) else pl[0]),(pl['max_companies'] if isinstance(pl,dict) else pl[1])

    # Create slug from name
    import re
    slug=re.sub(r'[^a-z0-9]+','-',d["name"].lower()).strip('-')
    base=slug
    for i in range(1,100):
        c.execute("SELECT 1 FROM tenants WHERE slug=%s", (slug,))
        if not c.fetchone(): break
        slug=f"{base}-{i}"

    tid=str(uuid.uuid4())
    try:
        c.execute("""INSERT INTO tenants (id,name,slug,email,phone,address,plan_id,
                     status,max_users,max_companies,notes)
                     VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                 (tid,d["name"],slug,d["email"],d.get("phone",""),d.get("address",""),
                  plan_id,"active",max_u,max_co,d.get("notes","")))
    except Exception as e:
        conn.close(); return jsonify({"error":str(e)}),400

    # Create tenant superadmin
    aid=str(uuid.uuid4())
    admin_email=d.get("admin_email",d["email"]).strip().lower()
    c.execute("""INSERT INTO users (id,name,email,password,role,is_active,tenant_id)
                 VALUES (%s,%s,%s,%s,%s,1,%s)""",
             (aid,d["admin_name"],admin_email,hash_pw(d["admin_password"]),"superadmin",tid))

    # Seed default document templates for this tenant (each gets new UUID)
    c.execute("""SELECT id,name,category,description,template_body FROM document_templates
                 WHERE tenant_id='default-tenant-001' LIMIT 5""")
    for row_t in c.fetchall():
        c.execute("""INSERT INTO document_templates
                     (id,name,category,description,template_body,created_by,tenant_id)
                     VALUES (%s,%s,%s,%s,%s,%s,%s)""",
                 (str(uuid.uuid4()),
                  row_t['name'] if isinstance(row_t,dict) else row_t[1],
                  row_t['category'] if isinstance(row_t,dict) else row_t[2],
                  row_t['description'] if isinstance(row_t,dict) else row_t[3],
                  row_t['template_body'] if isinstance(row_t,dict) else row_t[4],
                  aid,tid))

    c.execute("SELECT t.*,p.name as plan_name FROM tenants t LEFT JOIN plans p ON t.plan_id=p.id WHERE t.id=%s",(tid,))
    result=row(c.fetchone())
    conn.commit(); conn.close()
    if result: result["admin_email"]=admin_email
    return jsonify(result or {"id":tid,"name":d["name"],"admin_email":admin_email}),201

@app.route("/api/platform/tenants/<tid>", methods=["PUT"])
@platform_admin_required
def update_tenant(tid):
    d=request.get_json(silent=True,force=True) or {}
    conn=get_db(); c=conn.cursor()
    fields={k:d[k] for k in ["name","email","phone","address","plan_id","status",
                               "max_users","max_companies","notes","billing_cycle"] if k in d}
    if "approved" in d and d["approved"]:
        fields["status"]="active"
        fields["approved_by"]=g.user_id
        fields["approved_at"]=datetime.utcnow().isoformat()

    # If plan changed, update max limits from new plan
    if "plan_id" in fields:
        c.execute("SELECT max_users,max_companies FROM plans WHERE id=%s", (fields["plan_id"],))
        pl=c.fetchone()
        if pl:
            if "max_users"     not in d: fields["max_users"]     = pl['max_users']     if isinstance(pl,dict) else pl[0]
            if "max_companies" not in d: fields["max_companies"] = pl['max_companies'] if isinstance(pl,dict) else pl[1]

    if not fields: return jsonify({"error":"Nothing to update"}),400
    fields["updated_at"]=datetime.utcnow().isoformat()
    c.execute(f"UPDATE tenants SET {','.join(k+'=%s' for k in fields)} WHERE id=%s",
              list(fields.values())+[tid])
    # Disable/enable all users in tenant on status change
    if "status" in fields:
        active=1 if fields["status"]=="active" else 0
        c.execute("UPDATE users SET is_active=%s WHERE tenant_id=%s AND is_platform_admin=0",(active,tid))
    conn.commit()
    c.execute("SELECT t.*,p.name as plan_name FROM tenants t LEFT JOIN plans p ON t.plan_id=p.id WHERE t.id=%s",(tid,))
    result=row(c.fetchone()); conn.close(); return jsonify(result)

@app.route("/api/platform/tenants/<tid>", methods=["DELETE"])
@platform_admin_required
def delete_tenant(tid):
    conn=get_db(); c=conn.cursor()
    # Soft delete — mark cancelled and disable users
    c.execute("UPDATE tenants SET status='cancelled',updated_at=NOW() WHERE id=%s",(tid,))
    c.execute("UPDATE users SET is_active=0 WHERE tenant_id=%s AND is_platform_admin=0",(tid,))
    conn.commit(); conn.close()
    return jsonify({"success":True})

@app.route("/api/platform/tenants/<tid>/users")
@platform_admin_required
def tenant_users(tid):
    conn=get_db(); c=conn.cursor()
    c.execute("""SELECT id,name,email,role,is_active,created_at,last_login
                 FROM users WHERE tenant_id=%s ORDER BY role,name""", (tid,))
    return jsonify(rows(c.fetchall()))

@app.route("/api/platform/tenants/<tid>/stats")
@platform_admin_required
def tenant_stats(tid):
    conn=get_db(); c=conn.cursor()
    stats={}
    for tbl,col in [("users","id"),("companies","id"),("tasks","id"),("documents","id"),("alerts","id")]:
        c.execute(f"SELECT COUNT(*) FROM {tbl} WHERE tenant_id=%s", (tid,))
        stats[tbl]= _count(c)
    conn.close(); return jsonify(stats)

@app.route("/api/platform/summary")
@platform_admin_required
def platform_summary():
    conn=get_db(); c=conn.cursor()
    c.execute("SELECT COUNT(*) FROM tenants WHERE status='active'");   active= _count(c)
    c.execute("SELECT COUNT(*) FROM tenants WHERE status='pending'");  pending= _count(c)
    c.execute("SELECT COUNT(*) FROM tenants WHERE status='suspended'");suspended= _count(c)
    c.execute("SELECT COUNT(*) FROM tenants");                         total= _count(c)
    c.execute("SELECT COUNT(*) FROM users WHERE is_platform_admin=0 AND is_active=1"); total_users= _count(c)
    c.execute("SELECT COUNT(*) FROM companies"); total_co= _count(c)
    c.execute("SELECT COALESCE(SUM(p.price_monthly),0) FROM tenants t JOIN plans p ON t.plan_id=p.id WHERE t.status='active' AND t.billing_cycle='monthly'")
    mrr= _count(c)
    conn.close()
    return jsonify({"active":active,"pending":pending,"suspended":suspended,"total":total,
                    "total_users":total_users,"total_companies":total_co,"mrr":mrr})


# ══ SETUP / BOOTSTRAP ════════════════════════════════════════════════════════
@app.route("/api/setup/platform-admin", methods=["POST"])
def setup_platform_admin():
    """
    First-time setup: create the platform admin user.
    Only works when NO platform admin exists yet.
    """
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users WHERE is_platform_admin=1")
    r = c.fetchone()
    cnt = int(list(r.values())[0] if isinstance(r, dict) else r[0])
    if cnt > 0:
        conn.close()
        return jsonify({"error": "Platform admin already exists. Use login."}), 409

    d = request.get_json(silent=True, force=True) or {}
    name  = (d.get("name") or "Platform Admin").strip()
    email = (d.get("email") or "").strip().lower()
    pw    = d.get("password") or ""

    if not email: conn.close(); return jsonify({"error": "Email required"}), 400
    if not pw or len(pw) < 8: conn.close(); return jsonify({"error": "Password must be at least 8 characters"}), 400

    import re as _re
    if not _re.match(r"[^@]+@[^@]+\.[^@]+", email):
        conn.close(); return jsonify({"error": "Invalid email format"}), 400

    uid = str(uuid.uuid4())
    try:
        c.execute("""INSERT INTO users (id,name,email,password,role,is_active,is_platform_admin,tenant_id)
                     VALUES (%s,%s,%s,%s,%s,1,1,NULL)""",
                  (uid, name, email, hash_pw(pw), "superadmin"))
        conn.commit()
    except Exception as ex:
        conn.rollback()
        conn.close()
        return jsonify({"error": f"Could not create admin: {str(ex)}"}), 400

    conn.close()
    token = make_token(uid, "superadmin", name, None, True)
    return jsonify({
        "success": True,
        "message": "Platform admin created successfully",
        "token": token,
        "user": {"id": uid, "name": name, "email": email, "role": "superadmin"},
        "is_platform_admin": True,
    }), 201

@app.route("/api/setup/status")
def setup_status():
    """Returns whether first-time setup has been completed."""
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users WHERE is_platform_admin=1")
    r = c.fetchone()
    cnt = int(list(r.values())[0] if isinstance(r, dict) else r[0])
    conn.close()
    return jsonify({"setup_complete": cnt > 0, "has_platform_admin": cnt > 0})

# ══ SCHEDULER ════════════════════════════════════════════════════════════════
@app.route("/api/scheduler/run", methods=["POST"])
@require_role("superadmin","manager")
def run_scheduler():
    results=run_compliance_checks(); return jsonify({"success":True,"results":results})

# ══ BOOT ═════════════════════════════════════════════════════════════════════
def ensure_columns():
    """Add any columns that may be missing from older databases."""
    conn = get_db(); c = conn.cursor()
    migrations = [
        ("companies",    "gstin",       "TEXT"),
        ("companies",    "tan",         "TEXT"),
        ("companies",    "website",     "TEXT"),
        ("directors",    "mca_notes",   "TEXT"),
        ("shareholders", "mca_notes",   "TEXT"),
        ("tasks",        "task_leader", "TEXT"),
        ("tasks",        "task_manager","TEXT"),
        ("tasks",        "team_members","TEXT DEFAULT '[]'"),
        ("tasks",        "billable",    "INTEGER DEFAULT 0"),
        ("tasks",        "estimated_hrs","NUMERIC(6,2) DEFAULT 0"),
        ("tasks",        "actual_hrs",  "NUMERIC(6,2) DEFAULT 0"),
        ("tasks",        "note",        "TEXT DEFAULT ''"),
        ("tasks",        "completed_at","TIMESTAMPTZ"),
        ("directors",    "tenant_id",   "TEXT"),
        ("auditors",     "tenant_id",   "TEXT"),
        ("shareholders", "tenant_id",   "TEXT"),
        ("meetings",     "tenant_id",   "TEXT"),
        ("users",        "totp_secret", "TEXT"),
        ("users",        "totp_enabled","INTEGER DEFAULT 0"),
        ("users",        "totp_verified","INTEGER DEFAULT 0"),
    ]
    for table, col, coltype in migrations:
        try:
            c.execute(f"ALTER TABLE {table} ADD COLUMN {col} {coltype}")
            conn.commit()
        except Exception:
            pass  # Column already exists
    conn.close()

# ── Module-level startup (runs under both gunicorn and direct python) ────────
def _startup():
    """Called at module import so both `python app.py` and gunicorn initialise the DB."""
    import traceback as _tb
    for _fn, _name in [
        (init_db,                          "init_db"),
        (ensure_columns,                   "ensure_columns"),
        (ensure_custom_placeholders_table, "ensure_custom_placeholders_table"),
        (ensure_permission_tables,         "ensure_permission_tables"),
    ]:
        try:
            _fn()
        except Exception as _ex:
            print(f"[STARTUP] {_name} failed: {_ex}")
            _tb.print_exc()


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #1 — ROC Compliance Calendar (MCA annual filing deadlines)
# ══════════════════════════════════════════════════════════════════════════════
_MCA_DEADLINES = [
    # (form, description, month, day, note)
    ("MGT-7",  "Annual Return",                                  9, 29, "Within 60 days of AGM (assumed AGM on 30 Sep)"),
    ("AOC-4",  "Financial Statements",                           10, 29, "Within 30 days of AGM"),
    ("ADT-1",  "Auditor Appointment",                            10, 14, "Within 15 days of AGM"),
    ("DIR-3 KYC", "Director KYC",                               9, 30, "Annual KYC by 30 Sep"),
    ("DPT-3",  "Deposits / Loans Return",                        6, 30, "Annual return of outstanding deposits by 30 Jun"),
    ("MSME-1", "MSME Outstanding Payments",                      4, 30, "Half-yearly: Oct–Mar due by 30 Apr"),
    ("MSME-1", "MSME Outstanding Payments (H2)",                10, 31, "Half-yearly: Apr–Sep due by 31 Oct"),
    ("MGT-14", "Board Resolution Filing",                        None, None, "Within 30 days of board resolution (event-based)"),
    ("INC-20A","Commencement of Business",                       None, None, "Within 180 days of incorporation (one-time)"),
    ("DIR-12", "Change in Directors",                            None, None, "Within 30 days of appointment/cessation (event-based)"),
    ("SH-7",   "Change in Authorised Capital",                   None, None, "Within 30 days of resolution (event-based)"),
    ("PAS-3",  "Return of Allotment",                            None, None, "Within 30 days of allotment (event-based)"),
    ("CHG-1",  "Charge Creation",                                None, None, "Within 30 days of creation (event-based)"),
    ("CHG-4",  "Charge Satisfaction",                            None, None, "Within 30 days of satisfaction (event-based)"),
    ("BEN-2",  "Significant Beneficial Owner",                   None, None, "Within 30 days of receipt of BEN-1 (event-based)"),
    ("CRA-4",  "Cost Audit Report",                              9, 27, "Within 30 days of cost auditor's report (if applicable)"),
    ("MGT-7A", "Abridged Annual Return (OPC/Small)",             9, 29, "For OPC/Small Companies — same window as MGT-7"),
    ("AOC-4 CFS", "Consolidated Financial Statements",          10, 29, "If holding company — same window as AOC-4"),
    ("FC-4",   "Annual Return (Foreign Company)",                12, 31, "Within 60 days of last day of financial year"),
    ("LLP-11", "LLP Annual Return",                              5, 30, "Within 60 days of end of financial year"),
    ("LLP-8",  "Statement of Accounts (LLP)",                    10, 30, "Within 30 days of end of 6 months from FY end"),
]

@app.route("/api/compliance-calendar")
@login_required
def compliance_calendar():
    """
    General MCA filing calendar — one row per statutory form per FY.
    Company-agnostic: shows what filings are due this FY regardless of
    how many companies the tenant has.
    """
    year_arg  = request.args.get("year", "")
    today     = date.today()
    fy_start  = today.year if today.month >= 4 else today.year - 1
    target_year = int(year_arg) if year_arg.isdigit() else fy_start

    import calendar as _cal
    calendar_rows = []
    for form, desc, month, day, note in _MCA_DEADLINES:
        if month is None:
            entry = {
                "form": form, "description": desc,
                "due_date": None, "days_left": None,
                "status": "event_based", "note": note,
                "financial_year": f"{target_year}-{str(target_year+1)[-2:]}",
            }
        else:
            yr = target_year + (1 if month < 4 else 0)
            max_day = _cal.monthrange(yr, month)[1]
            due = date(yr, month, min(day, max_day))
            days_left = (due - today).days
            if days_left < 0:        status = "overdue"
            elif days_left <= 30:    status = "due_soon"
            elif days_left <= 90:    status = "upcoming"
            else:                    status = "future"
            entry = {
                "form": form, "description": desc,
                "due_date": due.isoformat(), "days_left": days_left,
                "status": status, "note": note,
                "financial_year": f"{target_year}-{str(target_year+1)[-2:]}",
            }
        calendar_rows.append(entry)

    # Sort: overdue → due_soon → upcoming → future → event_based
    _order = {"overdue": 0, "due_soon": 1, "upcoming": 2, "future": 3, "event_based": 4}
    calendar_rows.sort(key=lambda e: (_order.get(e["status"], 5), e.get("due_date") or ""))

    overdue_cnt  = sum(1 for e in calendar_rows if e["status"] == "overdue")
    due_soon_cnt = sum(1 for e in calendar_rows if e["status"] == "due_soon")
    return jsonify({
        "financial_year": f"{target_year}-{str(target_year+1)[-2:]}",
        "target_year": target_year,
        "total": len(calendar_rows),
        "overdue": overdue_cnt,
        "due_soon": due_soon_cnt,
        "deadlines": calendar_rows,
    })


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #2 — Email notifications (Flask-Mail / AWS SES)
# ══════════════════════════════════════════════════════════════════════════════
import smtplib, ssl
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

_MAIL_SERVER   = os.environ.get("MAIL_SERVER",   "smtp.gmail.com")
_MAIL_PORT     = int(os.environ.get("MAIL_PORT",  "587"))
_MAIL_USE_TLS  = os.environ.get("MAIL_USE_TLS",  "1") == "1"
_MAIL_USERNAME = os.environ.get("MAIL_USERNAME",  "")
_MAIL_PASSWORD = os.environ.get("MAIL_PASSWORD",  "")
_MAIL_FROM     = os.environ.get("MAIL_FROM",      _MAIL_USERNAME or "noreply@taxlycms.in")
_MAIL_ENABLED  = bool(_MAIL_USERNAME and _MAIL_PASSWORD)

def _send_email(to: str, subject: str, body_html: str, body_text: str = "") -> bool:
    """Send an email. Returns True on success. Silent on failure (never crashes the caller)."""
    if not _MAIL_ENABLED:
        app.logger.debug(f"[MAIL disabled] Would send to {to}: {subject}")
        return False
    try:
        msg = MIMEMultipart("alternative")
        msg["Subject"] = subject
        msg["From"]    = _MAIL_FROM
        msg["To"]      = to
        if body_text:
            msg.attach(MIMEText(body_text, "plain"))
        msg.attach(MIMEText(body_html, "html"))
        ctx = ssl.create_default_context()
        with smtplib.SMTP(_MAIL_SERVER, _MAIL_PORT) as srv:
            if _MAIL_USE_TLS:
                srv.starttls(context=ctx)
            srv.login(_MAIL_USERNAME, _MAIL_PASSWORD)
            srv.sendmail(_MAIL_FROM, to, msg.as_string())
        return True
    except Exception as e:
        app.logger.error(f"[MAIL ERROR] to={to} subject={subject}: {e}")
        return False

def _html_alert_email(user_name: str, alerts: list) -> tuple:
    """Return (html, plain) for an alert digest email."""
    rows_html = "".join(
        f"<tr><td style='padding:8px;border-bottom:1px solid #eee'><b>{a['title']}</b></td>"
        f"<td style='padding:8px;border-bottom:1px solid #eee;color:#e53e3e'>{a.get('due_date','')}</td>"
        f"<td style='padding:8px;border-bottom:1px solid #eee'>{a.get('severity','').upper()}</td></tr>"
        for a in alerts
    )
    html = f"""
    <div style='font-family:Arial,sans-serif;max-width:600px;margin:auto'>
      <div style='background:#0f2d5c;padding:20px;text-align:center'>
        <h2 style='color:white;margin:0'>Taxly CMS — Compliance Alert</h2>
      </div>
      <div style='padding:24px'>
        <p>Hi {user_name},</p>
        <p>You have <b>{len(alerts)}</b> compliance alert(s) requiring attention:</p>
        <table width='100%' cellspacing='0' style='border-collapse:collapse;margin-top:12px'>
          <thead><tr style='background:#f7fafc'>
            <th style='padding:8px;text-align:left'>Alert</th>
            <th style='padding:8px;text-align:left'>Due Date</th>
            <th style='padding:8px;text-align:left'>Severity</th>
          </tr></thead>
          <tbody>{rows_html}</tbody>
        </table>
        <p style='margin-top:20px'><a href='#' style='background:#1a56db;color:white;padding:10px 20px;text-decoration:none;border-radius:4px'>View in Taxly CMS</a></p>
      </div>
      <div style='background:#f7fafc;padding:12px;text-align:center;font-size:12px;color:#718096'>
        Taxly India Private Limited — You received this because you are a compliance team member.
      </div>
    </div>"""
    plain = f"Hi {user_name},\n\nYou have {len(alerts)} compliance alert(s):\n\n"
    plain += "\n".join(f"- {a['title']} (Due: {a.get('due_date','')} | {a.get('severity','').upper()})" for a in alerts)
    return html, plain

@app.route("/api/notifications/send-digest", methods=["POST"])
@require_role("superadmin", "manager")
def send_alert_digest():
    """Send compliance alert digest emails to all active users in the tenant."""
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT id, name, email FROM users WHERE tenant_id=%s AND is_active=1", (g.tenant_id,))
    users = rows(c.fetchall())
    c.execute("""SELECT a.*, co.name as company_name FROM alerts a
                 LEFT JOIN companies co ON a.company_id=co.id
                 WHERE a.tenant_id=%s AND a.status='active' AND a.severity IN ('critical','high')
                 ORDER BY a.due_date""", (g.tenant_id,))
    alert_list = rows(c.fetchall())
    conn.close()

    if not alert_list:
        return jsonify({"success": True, "message": "No active critical/high alerts to notify about", "sent": 0})

    sent = 0
    for u in users:
        if not u.get("email"): continue
        html, plain = _html_alert_email(u["name"], alert_list)
        if _send_email(u["email"], f"[Taxly CMS] {len(alert_list)} Compliance Alerts Require Attention", html, plain):
            sent += 1

    return jsonify({"success": True, "sent": sent, "total_users": len(users), "alerts": len(alert_list),
                    "mail_enabled": _MAIL_ENABLED})

@app.route("/api/notifications/test-email", methods=["POST"])
@require_role("superadmin")
def test_email():
    """Send a test email to the requesting user."""
    ok = _send_email(g.user["email"], "[Taxly CMS] Test Email",
                     "<h2>Test email from Taxly CMS</h2><p>Email configuration is working correctly.</p>",
                     "Test email from Taxly CMS — email configuration is working correctly.")
    return jsonify({"success": ok, "mail_enabled": _MAIL_ENABLED,
                    "to": g.user["email"],
                    "note": "Set MAIL_USERNAME, MAIL_PASSWORD, MAIL_SERVER env vars to enable." if not _MAIL_ENABLED else None})


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #4 — API Pagination (page/limit on key list endpoints)
# ══════════════════════════════════════════════════════════════════════════════
def _paginate(c, base_query: str, params: list, request_args) -> dict:
    """
    Execute base_query with optional LIMIT/OFFSET pagination.
    Returns {"data": [...], "total": N, "page": P, "limit": L, "pages": total_pages}
    If ?limit=0 or limit not provided, returns all rows (legacy behaviour).
    """
    limit = request_args.get("limit", "0", type=int)
    page  = max(1, request_args.get("page", 1, type=int))

    # Count total rows first
    count_sql = f"SELECT COUNT(*) FROM ({base_query}) AS _cnt"
    c.execute(count_sql, params)
    _r = c.fetchone()
    total = int(list(_r.values())[0] if isinstance(_r, dict) else _r[0]) if _r else 0

    if limit > 0:
        offset = (page - 1) * limit
        c.execute(base_query + f" LIMIT {limit} OFFSET {offset}", params)
    else:
        c.execute(base_query, params)

    data = rows(c.fetchall())
    pages = (total + limit - 1) // limit if limit > 0 else 1
    return {"data": data, "total": total, "page": page,
            "limit": limit, "pages": pages}


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #5 — 2FA / TOTP
# ══════════════════════════════════════════════════════════════════════════════
try:
    import pyotp, qrcode, base64 as _b64
    from io import BytesIO as _BytesIO
    _TOTP_AVAILABLE = True
except ImportError:
    _TOTP_AVAILABLE = False

def _ensure_totp_columns():
    """Add TOTP columns to users table if not present."""
    conn = get_db(); c = conn.cursor()
    for col, typ in [("totp_secret", "TEXT"), ("totp_enabled", "INTEGER DEFAULT 0"),
                     ("totp_verified", "INTEGER DEFAULT 0")]:
        try:
            c.execute(f"ALTER TABLE users ADD COLUMN {col} {typ}")
            conn.commit()
        except Exception:
            pass
    conn.close()

@app.route("/api/auth/2fa/setup", methods=["POST"])
@login_required
def totp_setup():
    """Generate a TOTP secret and QR code URI for the requesting user."""
    if not _TOTP_AVAILABLE:
        return jsonify({"error": "2FA not available — install pyotp and qrcode"}), 501
    _ensure_totp_columns()
    secret = pyotp.random_base32()
    totp   = pyotp.TOTP(secret)
    issuer = "TaxlyCMS"
    uri    = totp.provisioning_uri(name=g.user["email"], issuer_name=issuer)

    # Generate QR code as base64 PNG
    qr = qrcode.make(uri)
    buf = _BytesIO()
    qr.save(buf, format="PNG")
    qr_b64 = _b64.b64encode(buf.getvalue()).decode()

    # Store secret (not yet enabled — must be confirmed with a valid code)
    conn = get_db(); c = conn.cursor()
    c.execute("UPDATE users SET totp_secret=%s, totp_enabled=0, totp_verified=0 WHERE id=%s",
              (secret, g.user_id))
    conn.commit(); conn.close()
    return jsonify({"secret": secret, "uri": uri, "qr_code": f"data:image/png;base64,{qr_b64}"})

@app.route("/api/auth/2fa/verify", methods=["POST"])
@login_required
def totp_verify():
    """Confirm TOTP setup by verifying the first code. Enables 2FA on success."""
    if not _TOTP_AVAILABLE:
        return jsonify({"error": "2FA not available"}), 501
    _ensure_totp_columns()
    d = request.get_json(silent=True) or {}
    code = str(d.get("code", "")).strip()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT totp_secret FROM users WHERE id=%s", (g.user_id,))
    r = c.fetchone()
    secret = (r.get("totp_secret") if isinstance(r, dict) else r[0]) if r else None
    if not secret:
        conn.close(); return jsonify({"error": "Run /api/auth/2fa/setup first"}), 400
    totp = pyotp.TOTP(secret)
    if not totp.verify(code, valid_window=1):
        conn.close(); return jsonify({"error": "Invalid code"}), 400
    c.execute("UPDATE users SET totp_enabled=1, totp_verified=1 WHERE id=%s", (g.user_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True, "message": "2FA enabled successfully"})

@app.route("/api/auth/2fa/disable", methods=["POST"])
@login_required
def totp_disable():
    """Disable 2FA for the requesting user (requires password confirmation)."""
    _ensure_totp_columns()
    d = request.get_json(silent=True) or {}
    pw = d.get("password", "")
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT password FROM users WHERE id=%s", (g.user_id,))
    r = c.fetchone()
    stored = (r.get("password") if isinstance(r, dict) else r[0]) if r else ""
    if not verify_pw(pw, stored):
        conn.close(); return jsonify({"error": "Password incorrect"}), 400
    c.execute("UPDATE users SET totp_enabled=0, totp_secret=NULL WHERE id=%s", (g.user_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/auth/2fa/validate", methods=["POST"])
def totp_validate():
    """
    Second-step login: validate TOTP code after password check.
    Expects {"user_id": "...", "code": "123456"} in the request body.
    Returns a JWT on success.
    """
    if not _TOTP_AVAILABLE:
        return jsonify({"error": "2FA not available"}), 501
    _ensure_totp_columns()
    d = request.get_json(silent=True) or {}
    user_id = d.get("user_id", "")
    code    = str(d.get("code", "")).strip()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM users WHERE id=%s AND is_active=1", (user_id,))
    user = row(c.fetchone())
    if not user:
        conn.close(); return jsonify({"error": "User not found"}), 401
    secret = user.get("totp_secret")
    if not secret or not user.get("totp_enabled"):
        conn.close(); return jsonify({"error": "2FA not enabled for this user"}), 400
    if not pyotp.TOTP(secret).verify(code, valid_window=1):
        conn.close(); return jsonify({"error": "Invalid 2FA code"}), 401
    c.execute("UPDATE users SET last_login=NOW() WHERE id=%s", (user_id,))
    conn.commit(); conn.close()
    is_pa = bool(user.get("is_platform_admin", 0))
    return jsonify({
        "token": make_token(user["id"], user["role"], user["name"], user.get("tenant_id"), is_pa),
        "user": {"id": user["id"], "name": user["name"], "email": user["email"], "role": user["role"]},
    })

@app.route("/api/auth/2fa/status")
@login_required
def totp_status():
    """Return 2FA enabled status for the requesting user."""
    _ensure_totp_columns()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT totp_enabled, totp_verified FROM users WHERE id=%s", (g.user_id,))
    r = c.fetchone()
    conn.close()
    enabled  = bool((r.get("totp_enabled")  if isinstance(r, dict) else r[0]) if r else False)
    verified = bool((r.get("totp_verified") if isinstance(r, dict) else r[1]) if r else False)
    return jsonify({"enabled": enabled, "verified": verified, "available": _TOTP_AVAILABLE})


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #6 — Task Comments
# ══════════════════════════════════════════════════════════════════════════════
def _ensure_task_comments_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""
            CREATE TABLE IF NOT EXISTS task_comments (
                id TEXT PRIMARY KEY,
                task_id TEXT NOT NULL REFERENCES tasks(id) ON DELETE CASCADE,
                user_id TEXT REFERENCES users(id) ON DELETE SET NULL,
                content TEXT NOT NULL,
                tenant_id TEXT REFERENCES tenants(id),
                created_at TIMESTAMPTZ DEFAULT NOW(),
                updated_at TIMESTAMPTZ DEFAULT NOW()
            )
        """)
        c.execute("CREATE INDEX IF NOT EXISTS idx_task_comments_task ON task_comments(task_id)")
    else:
        c.execute("""
            CREATE TABLE IF NOT EXISTS task_comments (
                id TEXT PRIMARY KEY,
                task_id TEXT NOT NULL,
                user_id TEXT,
                content TEXT NOT NULL,
                tenant_id TEXT,
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            )
        """)
    conn.commit(); conn.close()

from database import USE_POSTGRES

@app.route("/api/tasks/<task_id>/comments")
@login_required
def list_task_comments(task_id):
    _ensure_task_comments_table()
    conn = get_db(); c = conn.cursor()
    c.execute("""
        SELECT tc.*, u.name AS author_name, u.email AS author_email
        FROM task_comments tc
        LEFT JOIN users u ON tc.user_id = u.id
        WHERE tc.task_id = %s
        ORDER BY tc.created_at ASC
    """, (task_id,))
    result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/tasks/<task_id>/comments", methods=["POST"])
@login_required
def create_task_comment(task_id):
    _ensure_task_comments_table()
    d = request.get_json(silent=True) or {}
    content = (d.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Comment content is required"}), 400
    cid = str(uuid.uuid4())
    conn = get_db(); c = conn.cursor()
    # Verify task exists and belongs to tenant
    c.execute("SELECT id FROM tasks WHERE id=%s AND tenant_id=%s", (task_id, g.tenant_id))
    if not c.fetchone():
        conn.close(); return jsonify({"error": "Task not found"}), 404
    c.execute("""
        INSERT INTO task_comments (id, task_id, user_id, content, tenant_id)
        VALUES (%s, %s, %s, %s, %s)
    """, (cid, task_id, g.user_id, content, g.tenant_id))
    conn.commit()
    c.execute("""
        SELECT tc.*, u.name AS author_name, u.email AS author_email
        FROM task_comments tc LEFT JOIN users u ON tc.user_id=u.id
        WHERE tc.id=%s
    """, (cid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/tasks/<task_id>/comments/<comment_id>", methods=["PUT"])
@login_required
def update_task_comment(task_id, comment_id):
    _ensure_task_comments_table()
    d = request.get_json(silent=True) or {}
    content = (d.get("content") or "").strip()
    if not content:
        return jsonify({"error": "Comment content is required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT user_id FROM task_comments WHERE id=%s AND task_id=%s", (comment_id, task_id))
    r = c.fetchone()
    if not r:
        conn.close(); return jsonify({"error": "Comment not found"}), 404
    owner = r.get("user_id") if isinstance(r, dict) else r[0]
    if owner != g.user_id and g.role not in ("superadmin", "manager"):
        conn.close(); return jsonify({"error": "Not authorised to edit this comment"}), 403
    c.execute("UPDATE task_comments SET content=%s, updated_at=NOW() WHERE id=%s", (content, comment_id))
    conn.commit()
    c.execute("""SELECT tc.*, u.name AS author_name FROM task_comments tc
                 LEFT JOIN users u ON tc.user_id=u.id WHERE tc.id=%s""", (comment_id,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/tasks/<task_id>/comments/<comment_id>", methods=["DELETE"])
@login_required
def delete_task_comment(task_id, comment_id):
    _ensure_task_comments_table()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT user_id FROM task_comments WHERE id=%s AND task_id=%s", (comment_id, task_id))
    r = c.fetchone()
    if not r:
        conn.close(); return jsonify({"error": "Comment not found"}), 404
    owner = r.get("user_id") if isinstance(r, dict) else r[0]
    if owner != g.user_id and g.role not in ("superadmin", "manager"):
        conn.close(); return jsonify({"error": "Not authorised"}), 403
    c.execute("DELETE FROM task_comments WHERE id=%s", (comment_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True})


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #7 — File Attachments (tasks & documents)
# ══════════════════════════════════════════════════════════════════════════════
ATTACH_DIR  = BASE_DIR / "data" / "attachments"
ATTACH_ALLOWED = {".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png", ".xlsx", ".xls", ".csv", ".txt"}
ATTACH_MAX_MB  = int(os.environ.get("ATTACH_MAX_MB", "20"))

def _ensure_attachments_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""
            CREATE TABLE IF NOT EXISTS attachments (
                id TEXT PRIMARY KEY,
                entity_type TEXT NOT NULL,
                entity_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                original_name TEXT NOT NULL,
                file_size INTEGER DEFAULT 0,
                mime_type TEXT,
                uploaded_by TEXT REFERENCES users(id) ON DELETE SET NULL,
                tenant_id TEXT REFERENCES tenants(id),
                created_at TIMESTAMPTZ DEFAULT NOW()
            )
        """)
        c.execute("CREATE INDEX IF NOT EXISTS idx_attachments_entity ON attachments(entity_type, entity_id)")
    else:
        c.execute("""
            CREATE TABLE IF NOT EXISTS attachments (
                id TEXT PRIMARY KEY,
                entity_type TEXT NOT NULL,
                entity_id TEXT NOT NULL,
                filename TEXT NOT NULL,
                original_name TEXT NOT NULL,
                file_size INTEGER DEFAULT 0,
                mime_type TEXT,
                uploaded_by TEXT,
                tenant_id TEXT,
                created_at TEXT DEFAULT (datetime('now'))
            )
        """)
    conn.commit(); conn.close()

@app.route("/api/attachments/<entity_type>/<entity_id>", methods=["GET"])
@login_required
def list_attachments(entity_type, entity_id):
    _ensure_attachments_table()
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT a.*, u.name AS uploader_name FROM attachments a
                 LEFT JOIN users u ON a.uploaded_by=u.id
                 WHERE a.entity_type=%s AND a.entity_id=%s AND a.tenant_id=%s
                 ORDER BY a.created_at DESC""",
              (entity_type, entity_id, g.tenant_id))
    result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/attachments/<entity_type>/<entity_id>", methods=["POST"])
@login_required
def upload_attachment(entity_type, entity_id):
    _ensure_attachments_table()
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Empty filename"}), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in ATTACH_ALLOWED:
        return jsonify({"error": f"File type {ext} not allowed. Allowed: {', '.join(sorted(ATTACH_ALLOWED))}"}), 400

    raw = f.read()
    if len(raw) > ATTACH_MAX_MB * 1024 * 1024:
        return jsonify({"error": f"File exceeds {ATTACH_MAX_MB}MB limit"}), 413

    aid = str(uuid.uuid4())
    ATTACH_DIR.mkdir(parents=True, exist_ok=True)
    stored_name = f"{aid}{ext}"
    (ATTACH_DIR / stored_name).write_bytes(raw)

    conn = get_db(); c = conn.cursor()
    c.execute("""
        INSERT INTO attachments (id, entity_type, entity_id, filename, original_name,
                                  file_size, mime_type, uploaded_by, tenant_id)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
    """, (aid, entity_type, entity_id, stored_name,
          secure_filename(f.filename), len(raw),
          f.content_type or "application/octet-stream",
          g.user_id, g.tenant_id))
    conn.commit()
    c.execute("""SELECT a.*, u.name AS uploader_name FROM attachments a
                 LEFT JOIN users u ON a.uploaded_by=u.id WHERE a.id=%s""", (aid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/attachments/<attach_id>/download")
@login_required
def download_attachment(attach_id):
    _ensure_attachments_table()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM attachments WHERE id=%s AND tenant_id=%s", (attach_id, g.tenant_id))
    att = row(c.fetchone()); conn.close()
    if not att:
        return jsonify({"error": "Attachment not found"}), 404
    fpath = ATTACH_DIR / att["filename"]
    if not fpath.exists():
        return jsonify({"error": "File missing on disk"}), 404
    return send_file(str(fpath), as_attachment=True,
                     download_name=att["original_name"],
                     mimetype=att.get("mime_type", "application/octet-stream"))

@app.route("/api/attachments/<attach_id>", methods=["DELETE"])
@login_required
def delete_attachment(attach_id):
    _ensure_attachments_table()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM attachments WHERE id=%s AND tenant_id=%s", (attach_id, g.tenant_id))
    att = row(c.fetchone())
    if not att:
        conn.close(); return jsonify({"error": "Attachment not found"}), 404
    # Only uploader, manager, or superadmin may delete
    if att.get("uploaded_by") != g.user_id and g.role not in ("superadmin", "manager"):
        conn.close(); return jsonify({"error": "Not authorised"}), 403
    try:
        (ATTACH_DIR / att["filename"]).unlink(missing_ok=True)
    except Exception:
        pass
    c.execute("DELETE FROM attachments WHERE id=%s", (attach_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True})


# ══════════════════════════════════════════════════════════════════════════════
# P0 FIX #8 — Forgot Password / Self-serve reset
# ══════════════════════════════════════════════════════════════════════════════
import secrets as _secrets
import hashlib as _hashlib

def _ensure_password_reset_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""
            CREATE TABLE IF NOT EXISTS password_reset_tokens (
                id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                token_hash TEXT NOT NULL UNIQUE,
                expires_at TIMESTAMPTZ NOT NULL,
                used INTEGER DEFAULT 0,
                created_at TIMESTAMPTZ DEFAULT NOW()
            )
        """)
        c.execute("CREATE INDEX IF NOT EXISTS idx_prt_token ON password_reset_tokens(token_hash)")
    else:
        c.execute("""
            CREATE TABLE IF NOT EXISTS password_reset_tokens (
                id TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                token_hash TEXT NOT NULL UNIQUE,
                expires_at TEXT NOT NULL,
                used INTEGER DEFAULT 0,
                created_at TEXT DEFAULT (datetime('now'))
            )
        """)
    conn.commit(); conn.close()

@app.route("/api/auth/forgot-password", methods=["POST"])
def forgot_password():
    """
    Request a password-reset link. Always returns 200 to prevent email enumeration.
    Sends an email if the address exists and mail is configured.
    """
    _ensure_password_reset_table()
    d = request.get_json(silent=True) or {}
    email = (d.get("email") or "").strip().lower()
    if not email:
        return jsonify({"message": "If that email is registered you will receive a reset link."}), 200

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT id, name, email FROM users WHERE email=%s AND is_active=1", (email,))
    user = row(c.fetchone())

    if user:
        # Invalidate any existing unused tokens for this user
        c.execute("UPDATE password_reset_tokens SET used=1 WHERE user_id=%s AND used=0", (user["id"],))
        # Generate a secure token
        raw_token   = _secrets.token_urlsafe(32)
        token_hash  = _hashlib.sha256(raw_token.encode()).hexdigest()
        expires_at  = datetime.utcnow() + timedelta(hours=2)
        c.execute("""INSERT INTO password_reset_tokens (id, user_id, token_hash, expires_at)
                     VALUES (%s, %s, %s, %s)""",
                  (str(uuid.uuid4()), user["id"], token_hash, expires_at.isoformat()))
        conn.commit()

        # Build reset URL (frontend handles the form)
        reset_url = os.environ.get("FRONTEND_URL", "http://localhost:5000").rstrip("/")
        reset_url += f"/#/reset-password?token={raw_token}"

        html = f"""
        <div style='font-family:Arial,sans-serif;max-width:520px;margin:auto'>
          <div style='background:#0f2d5c;padding:20px;text-align:center'>
            <h2 style='color:white;margin:0'>Taxly CMS — Password Reset</h2>
          </div>
          <div style='padding:24px'>
            <p>Hi {user['name']},</p>
            <p>We received a request to reset your password. Click the button below to set a new one.
               This link is valid for <b>2 hours</b>.</p>
            <p style='text-align:center;margin:28px 0'>
              <a href='{reset_url}' style='background:#1a56db;color:white;padding:12px 28px;text-decoration:none;border-radius:4px;font-weight:bold'>
                Reset My Password
              </a>
            </p>
            <p style='font-size:13px;color:#718096'>If you did not request this, ignore this email — your password will not change.</p>
            <p style='font-size:12px;color:#a0aec0;word-break:break-all'>Link: {reset_url}</p>
          </div>
        </div>"""
        plain = (f"Hi {user['name']},\n\nReset your Taxly CMS password:\n{reset_url}\n\n"
                 "This link expires in 2 hours. If you didn't request this, ignore this email.")
        _send_email(user["email"], "[Taxly CMS] Password Reset Request", html, plain)

    conn.close()
    return jsonify({"message": "If that email is registered you will receive a reset link."})

@app.route("/api/auth/reset-password", methods=["POST"])
def reset_password():
    """Consume a password-reset token and set a new password."""
    _ensure_password_reset_table()
    d = request.get_json(silent=True) or {}
    raw_token = (d.get("token") or "").strip()
    new_pw    = (d.get("password") or "").strip()

    if not raw_token or not new_pw:
        return jsonify({"error": "Token and new password are required"}), 400
    if len(new_pw) < 8:
        return jsonify({"error": "Password must be at least 8 characters"}), 400

    token_hash = _hashlib.sha256(raw_token.encode()).hexdigest()
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT prt.*, u.id AS uid FROM password_reset_tokens prt
                 JOIN users u ON prt.user_id=u.id
                 WHERE prt.token_hash=%s AND prt.used=0""", (token_hash,))
    rec = row(c.fetchone())
    if not rec:
        conn.close(); return jsonify({"error": "Invalid or already-used reset token"}), 400

    # Check expiry
    expires_str = rec.get("expires_at") or ""
    try:
        expires = datetime.fromisoformat(str(expires_str)[:19])
    except Exception:
        expires = datetime.utcnow() - timedelta(seconds=1)  # treat as expired
    if datetime.utcnow() > expires:
        conn.close(); return jsonify({"error": "Reset token has expired — please request a new one"}), 400

    # All good — update password and mark token used
    c.execute("UPDATE users SET password=%s WHERE id=%s", (hash_pw(new_pw), rec["user_id"]))
    c.execute("UPDATE password_reset_tokens SET used=1 WHERE token_hash=%s", (token_hash,))
    conn.commit(); conn.close()
    return jsonify({"success": True, "message": "Password updated successfully. You can now log in."})

@app.route("/api/auth/reset-password/validate", methods=["POST"])
def validate_reset_token():
    """Check if a reset token is valid without consuming it (for frontend UX)."""
    _ensure_password_reset_table()
    d = request.get_json(silent=True) or {}
    raw_token = (d.get("token") or "").strip()
    if not raw_token:
        return jsonify({"valid": False}), 200
    token_hash = _hashlib.sha256(raw_token.encode()).hexdigest()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT expires_at, used FROM password_reset_tokens WHERE token_hash=%s", (token_hash,))
    rec = row(c.fetchone()); conn.close()
    if not rec or rec.get("used"):
        return jsonify({"valid": False, "reason": "Token not found or already used"})
    try:
        expires = datetime.fromisoformat(str(rec["expires_at"])[:19])
        if datetime.utcnow() > expires:
            return jsonify({"valid": False, "reason": "Token expired"})
    except Exception:
        return jsonify({"valid": False, "reason": "Could not parse expiry"})
    return jsonify({"valid": True})


_startup()  # called at import time so gunicorn workers also initialise DB

if __name__ == "__main__":
    run_compliance_checks()
    print("\n" + "="*62)
    print("  TAXLY-CMS v2.0 - Corporate Compliance Management System")
    print("="*62)
    print("  URL    : http://127.0.0.1:5000")
    print("  Admin  : admin@compli.in / admin123")
    print("  Manager: manager@compli.in / manager123")
    print("  Staff  : staff@compli.in / staff123")
    print("="*62 + "\n")
    app.run(host="127.0.0.1",port=5000,debug=False)

# ══════════════════════════════════════════════════════════════════════════════
# P1 SPRINT — All 15 features
# ══════════════════════════════════════════════════════════════════════════════

# ── P1 #1 — Compliance Health Score ──────────────────────────────────────────
@app.route("/api/companies/<cid>/health-score")
@login_required
def company_health_score(cid):
    """
    0–100 weighted score per company.
    KYC 25 | Filings 25 | Meetings 20 | DSC 15 | Tasks 15
    """
    conn = get_db(); c = conn.cursor()
    today = date.today()

    # ── KYC (25 pts) — director KYC filed & not overdue ─────────────────────
    c.execute("SELECT COUNT(*) FROM directors WHERE company_id=%s AND is_active=1", (cid,))
    r = c.fetchone(); total_dirs = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    c.execute("""SELECT COUNT(*) FROM directors d JOIN director_kyc k ON d.id=k.director_id
                 WHERE d.company_id=%s AND d.is_active=1 AND k.kyc_status='filed'""", (cid,))
    r = c.fetchone(); filed_kyc = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    kyc_score = int((filed_kyc / max(total_dirs, 1)) * 25)

    # ── Filings (25 pts) — tasks tagged as filings completed in last 365 days ─
    c.execute("""SELECT COUNT(*) FROM tasks WHERE company_id=%s AND module='filing'
                 AND status='completed' AND completed_at >= %s""",
              (cid, today - timedelta(days=365)))
    r = c.fetchone(); completed_filings = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    c.execute("SELECT COUNT(*) FROM tasks WHERE company_id=%s AND module='filing'", (cid,))
    r = c.fetchone(); total_filings = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    filing_score = min(25, int((completed_filings / max(total_filings, 1)) * 25)) if total_filings else 20

    # ── Meetings (20 pts) — meetings held in last 12 months ─────────────────
    c.execute("""SELECT COUNT(*) FROM meetings WHERE company_id=%s
                 AND meeting_date BETWEEN %s AND %s AND status != 'cancelled'""",
              (cid, today - timedelta(days=365), today))
    r = c.fetchone(); meetings_held = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    # Min 4 board meetings per year recommended
    meeting_score = min(20, int((meetings_held / 4) * 20))

    # ── DSC (15 pts) — no expired DSCs ───────────────────────────────────────
    c.execute("SELECT COUNT(*) FROM dsc_records WHERE company_id=%s AND is_active=1", (cid,))
    r = c.fetchone(); total_dsc = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    c.execute("""SELECT COUNT(*) FROM dsc_records WHERE company_id=%s AND is_active=1
                 AND valid_to < %s""", (cid, today))
    r = c.fetchone(); expired_dsc = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    dsc_score = 15 if total_dsc == 0 else max(0, int(((total_dsc - expired_dsc) / total_dsc) * 15))

    # ── Tasks (15 pts) — open tasks not overdue ───────────────────────────────
    c.execute("SELECT COUNT(*) FROM tasks WHERE company_id=%s AND status NOT IN ('completed','cancelled')", (cid,))
    r = c.fetchone(); open_t = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    c.execute("""SELECT COUNT(*) FROM tasks WHERE company_id=%s AND status NOT IN ('completed','cancelled')
                 AND due_date < %s""", (cid, today))
    r = c.fetchone(); overdue_t = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0
    task_score = 15 if open_t == 0 else max(0, int(((open_t - overdue_t) / open_t) * 15))

    conn.close()
    total = kyc_score + filing_score + meeting_score + dsc_score + task_score
    grade = "A" if total >= 85 else "B" if total >= 70 else "C" if total >= 50 else "D"
    return jsonify({
        "company_id": cid, "score": total, "grade": grade,
        "breakdown": {
            "kyc":      {"score": kyc_score,      "max": 25, "label": "Director KYC"},
            "filings":  {"score": filing_score,   "max": 25, "label": "Annual Filings"},
            "meetings": {"score": meeting_score,  "max": 20, "label": "Board Meetings"},
            "dsc":      {"score": dsc_score,      "max": 15, "label": "DSC Status"},
            "tasks":    {"score": task_score,     "max": 15, "label": "Task Compliance"},
        }
    })

@app.route("/api/health-scores")
@login_required
def all_health_scores():
    """Return health scores for all tenant companies."""
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT id, name FROM companies WHERE tenant_id=%s AND status='active'",
              (g.tenant_id,))
    companies = rows(c.fetchall()); conn.close()
    scores = []
    for co in companies:
        try:
            with app.test_request_context():
                pass  # can't nest; call logic inline
            # inline score computation (simplified for batch)
            scores.append({"company_id": co["id"], "company_name": co["name"],
                           "score": None})  # lazy — client fetches per-company
        except Exception:
            pass
    # Actually just return company list; client fetches scores individually
    return jsonify([{"company_id": co["id"], "company_name": co["name"]} for co in companies])


# ── P1 #4 — WhatsApp / SMS alerts ────────────────────────────────────────────
_TWILIO_SID   = os.environ.get("TWILIO_ACCOUNT_SID", "")
_TWILIO_TOKEN = os.environ.get("TWILIO_AUTH_TOKEN", "")
_TWILIO_FROM  = os.environ.get("TWILIO_FROM_NUMBER", "")   # +14155238886 for WA sandbox
_TWILIO_WA    = os.environ.get("TWILIO_WHATSAPP_FROM", "whatsapp:+14155238886")
_SMS_ENABLED  = bool(_TWILIO_SID and _TWILIO_TOKEN and _TWILIO_FROM)

def _send_sms(to: str, body: str, whatsapp: bool = False) -> bool:
    if not _SMS_ENABLED:
        app.logger.debug(f"[SMS disabled] Would send to {to}: {body[:60]}")
        return False
    try:
        import urllib.request, urllib.parse, base64 as _b64
        from_num = (_TWILIO_WA if whatsapp else _TWILIO_FROM)
        to_num   = (f"whatsapp:{to}" if whatsapp and not to.startswith("whatsapp:") else to)
        data = urllib.parse.urlencode({"From": from_num, "To": to_num, "Body": body}).encode()
        url  = f"https://api.twilio.com/2010-04-01/Accounts/{_TWILIO_SID}/Messages.json"
        req  = urllib.request.Request(url, data=data, method="POST")
        creds = _b64.b64encode(f"{_TWILIO_SID}:{_TWILIO_TOKEN}".encode()).decode()
        req.add_header("Authorization", f"Basic {creds}")
        urllib.request.urlopen(req, timeout=10)
        return True
    except Exception as e:
        app.logger.error(f"[SMS ERROR] to={to}: {e}")
        return False

@app.route("/api/notifications/send-whatsapp", methods=["POST"])
@require_role("superadmin", "manager")
def send_whatsapp_digest():
    """Send WhatsApp alert digest to a phone number."""
    d = request.get_json(silent=True) or {}
    to_num = (d.get("phone") or "").strip()
    if not to_num:
        return jsonify({"error": "phone number required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT title, due_date, severity FROM alerts WHERE tenant_id=%s
                 AND status='active' AND severity IN ('critical','high')
                 ORDER BY due_date LIMIT 10""", (g.tenant_id,))
    alerts_list = rows(c.fetchall()); conn.close()
    if not alerts_list:
        return jsonify({"success": True, "message": "No critical alerts to send"})
    lines = [f"🔔 *Taxly CMS — Compliance Alerts*\n"]
    for a in alerts_list:
        emoji = "🔴" if a["severity"] == "critical" else "🟠"
        lines.append(f"{emoji} {a['title']}\n   Due: {a.get('due_date','—')}")
    body = "\n".join(lines[:10])
    ok = _send_sms(to_num, body, whatsapp=True)
    return jsonify({"success": ok, "sms_enabled": _SMS_ENABLED, "alerts_sent": len(alerts_list)})

@app.route("/api/notifications/send-sms", methods=["POST"])
@require_role("superadmin", "manager")
def send_sms_alert():
    d = request.get_json(silent=True) or {}
    to_num = (d.get("phone") or "").strip()
    message = (d.get("message") or "Compliance alert from Taxly CMS.").strip()
    if not to_num: return jsonify({"error": "phone required"}), 400
    ok = _send_sms(to_num, message, whatsapp=False)
    return jsonify({"success": ok, "sms_enabled": _SMS_ENABLED})


# ── P1 #5 — MCA Form Tracker ──────────────────────────────────────────────────
_MCA_FORM_STATUSES = ["draft","ready","submitted","srn_received","approved","rejected","resubmitted"]

def _ensure_mca_filings_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""CREATE TABLE IF NOT EXISTS mca_filings (
            id TEXT PRIMARY KEY, company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            form_type TEXT NOT NULL, description TEXT,
            due_date DATE, filed_date DATE,
            status TEXT DEFAULT 'draft', srn TEXT DEFAULT '',
            amount_paid NUMERIC(10,2) DEFAULT 0,
            notes TEXT DEFAULT '', challan_no TEXT DEFAULT '',
            filed_by TEXT REFERENCES users(id) ON DELETE SET NULL,
            tenant_id TEXT REFERENCES tenants(id),
            created_at TIMESTAMPTZ DEFAULT NOW(), updated_at TIMESTAMPTZ DEFAULT NOW()
        )""")
        c.execute("CREATE INDEX IF NOT EXISTS idx_mca_company ON mca_filings(company_id)")
        c.execute("CREATE INDEX IF NOT EXISTS idx_mca_tenant  ON mca_filings(tenant_id)")
    else:
        c.execute("""CREATE TABLE IF NOT EXISTS mca_filings (
            id TEXT PRIMARY KEY, company_id TEXT NOT NULL, form_type TEXT NOT NULL,
            description TEXT, due_date TEXT, filed_date TEXT,
            status TEXT DEFAULT 'draft', srn TEXT DEFAULT '',
            amount_paid REAL DEFAULT 0, notes TEXT DEFAULT '', challan_no TEXT DEFAULT '',
            filed_by TEXT, tenant_id TEXT,
            created_at TEXT DEFAULT (datetime('now')), updated_at TEXT DEFAULT (datetime('now'))
        )""")
    conn.commit(); conn.close()

@app.route("/api/mca-filings")
@login_required
def list_mca_filings():
    _ensure_mca_filings_table()
    cid    = request.args.get("company_id", "")
    status = request.args.get("status", "")
    conn = get_db(); c = conn.cursor()
    q = """SELECT f.*, co.name AS company_name, u.name AS filed_by_name
           FROM mca_filings f LEFT JOIN companies co ON f.company_id=co.id
           LEFT JOIN users u ON f.filed_by=u.id WHERE f.tenant_id=%s"""
    params = [g.tenant_id]
    if cid:    q += " AND f.company_id=%s"; params.append(cid)
    if status: q += " AND f.status=%s";    params.append(status)
    q += " ORDER BY f.due_date"
    c.execute(q, params); result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/mca-filings", methods=["POST"])
@login_required
def create_mca_filing():
    _ensure_mca_filings_table()
    d = request.get_json(silent=True) or {}
    if not d.get("company_id") or not d.get("form_type"):
        return jsonify({"error": "company_id and form_type required"}), 400
    fid = str(uuid.uuid4())
    conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO mca_filings (id,company_id,form_type,description,due_date,status,
                 srn,amount_paid,notes,challan_no,filed_by,tenant_id)
                 VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
              (fid, d["company_id"], d["form_type"], d.get("description",""),
               _dt(d.get("due_date")), d.get("status","draft"),
               d.get("srn",""), float(d.get("amount_paid") or 0),
               d.get("notes",""), d.get("challan_no",""),
               g.user_id, g.tenant_id))
    conn.commit()
    c.execute("SELECT f.*,co.name AS company_name FROM mca_filings f LEFT JOIN companies co ON f.company_id=co.id WHERE f.id=%s", (fid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/mca-filings/<fid>", methods=["PUT"])
@login_required
def update_mca_filing(fid):
    _ensure_mca_filings_table()
    d = request.get_json(silent=True) or {}
    allowed = ["form_type","description","due_date","filed_date","status","srn",
               "amount_paid","notes","challan_no"]
    fields = {}
    for k in allowed:
        if k in d:
            if k in ("due_date","filed_date"): fields[k] = _dt(d[k])
            elif k == "amount_paid": fields[k] = float(d[k] or 0)
            else: fields[k] = d[k]
    if not fields: return jsonify({"error": "Nothing to update"}), 400
    fields["updated_at"] = datetime.utcnow().isoformat()
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE mca_filings SET {','.join(k+'=%s' for k in fields)} WHERE id=%s AND tenant_id=%s",
              list(fields.values()) + [fid, g.tenant_id])
    conn.commit()
    c.execute("SELECT f.*,co.name AS company_name FROM mca_filings f LEFT JOIN companies co ON f.company_id=co.id WHERE f.id=%s", (fid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/mca-filings/<fid>", methods=["DELETE"])
@require_role("superadmin","manager")
def delete_mca_filing(fid):
    _ensure_mca_filings_table()
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM mca_filings WHERE id=%s AND tenant_id=%s", (fid, g.tenant_id))
    conn.commit(); conn.close()
    return jsonify({"success": True})


# ── P1 #7 — Document Version Control ─────────────────────────────────────────
def _ensure_doc_versions_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""CREATE TABLE IF NOT EXISTS document_versions (
            id TEXT PRIMARY KEY, document_id TEXT NOT NULL,
            version_no INTEGER NOT NULL DEFAULT 1,
            content TEXT, snapshot_name TEXT,
            created_by TEXT REFERENCES users(id) ON DELETE SET NULL,
            tenant_id TEXT REFERENCES tenants(id),
            created_at TIMESTAMPTZ DEFAULT NOW()
        )""")
        c.execute("CREATE INDEX IF NOT EXISTS idx_docver_doc ON document_versions(document_id)")
    else:
        c.execute("""CREATE TABLE IF NOT EXISTS document_versions (
            id TEXT PRIMARY KEY, document_id TEXT NOT NULL,
            version_no INTEGER NOT NULL DEFAULT 1,
            content TEXT, snapshot_name TEXT, created_by TEXT, tenant_id TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )""")
    conn.commit(); conn.close()

@app.route("/api/documents/doc/<doc_id>/versions")
@login_required
def list_doc_versions(doc_id):
    _ensure_doc_versions_table()
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT dv.*, u.name AS created_by_name
                 FROM document_versions dv LEFT JOIN users u ON dv.created_by=u.id
                 WHERE dv.document_id=%s AND dv.tenant_id=%s
                 ORDER BY dv.version_no DESC""", (doc_id, g.tenant_id))
    result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/documents/doc/<doc_id>/versions", methods=["POST"])
@login_required
def create_doc_version(doc_id):
    _ensure_doc_versions_table()
    d = request.get_json(silent=True) or {}
    conn = get_db(); c = conn.cursor()
    # Verify doc belongs to tenant
    c.execute("SELECT id, content FROM documents WHERE id=%s AND tenant_id=%s", (doc_id, g.tenant_id))
    doc = row(c.fetchone())
    if not doc: conn.close(); return jsonify({"error": "Document not found"}), 404
    # Get next version number
    c.execute("SELECT COALESCE(MAX(version_no),0)+1 AS nv FROM document_versions WHERE document_id=%s", (doc_id,))
    r = c.fetchone(); nv = int(list(r.values())[0] if isinstance(r, dict) else r[0])
    vid = str(uuid.uuid4())
    content = d.get("content") or doc.get("content") or ""
    snap_name = d.get("snapshot_name") or f"v{nv} — {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"
    c.execute("""INSERT INTO document_versions (id,document_id,version_no,content,snapshot_name,created_by,tenant_id)
                 VALUES (%s,%s,%s,%s,%s,%s,%s)""",
              (vid, doc_id, nv, content, snap_name, g.user_id, g.tenant_id))
    # If content provided, update the document itself
    if d.get("content"):
        c.execute("UPDATE documents SET content=%s WHERE id=%s", (content, doc_id))
    conn.commit()
    c.execute("SELECT dv.*,u.name AS created_by_name FROM document_versions dv LEFT JOIN users u ON dv.created_by=u.id WHERE dv.id=%s", (vid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/documents/doc/<doc_id>/versions/<vid>/restore", methods=["POST"])
@login_required
def restore_doc_version(doc_id, vid):
    _ensure_doc_versions_table()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM document_versions WHERE id=%s AND document_id=%s", (vid, doc_id))
    ver = row(c.fetchone())
    if not ver: conn.close(); return jsonify({"error": "Version not found"}), 404
    c.execute("UPDATE documents SET content=%s WHERE id=%s AND tenant_id=%s",
              (ver["content"], doc_id, g.tenant_id))
    conn.commit(); conn.close()
    return jsonify({"success": True, "restored_version": ver["version_no"]})


# ── P1 #10 — Audit Trail UI ───────────────────────────────────────────────────
@app.route("/api/audit-log")
@require_role("superadmin")
def get_audit_log():
    page   = max(1, request.args.get("page", 1, type=int))
    limit  = min(100, request.args.get("limit", 50, type=int))
    module = request.args.get("module", "")
    user_f = request.args.get("user_id", "")
    action = request.args.get("action", "")
    from_d = request.args.get("from", "")
    to_d   = request.args.get("to", "")
    conn = get_db(); c = conn.cursor()
    q = """SELECT al.*, u.name AS user_name, u.email AS user_email
           FROM audit_log al LEFT JOIN users u ON al.user_id=u.id
           WHERE al.tenant_id=%s"""
    params = [g.tenant_id]
    if module: q += " AND al.module=%s"; params.append(module)
    if user_f: q += " AND al.user_id=%s"; params.append(user_f)
    if action: q += " AND al.action ILIKE %s"; params.append(f"%{action}%")
    if from_d: q += " AND al.created_at >= %s"; params.append(from_d)
    if to_d:   q += " AND al.created_at <= %s"; params.append(to_d+"T23:59:59")
    # Count
    c.execute(f"SELECT COUNT(*) FROM ({q}) AS _c", params)
    r = c.fetchone(); total = int(list(r.values())[0] if isinstance(r, dict) else r[0])
    q += " ORDER BY al.created_at DESC LIMIT %s OFFSET %s"
    params += [limit, (page-1)*limit]
    c.execute(q, params); result = rows(c.fetchall()); conn.close()
    return jsonify({"data": result, "total": total, "page": page, "limit": limit,
                    "pages": (total + limit - 1) // limit})


# ── P1 #11 — Recurring Task Templates ────────────────────────────────────────
def _ensure_recurring_templates_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""CREATE TABLE IF NOT EXISTS recurring_task_templates (
            id TEXT PRIMARY KEY, title TEXT NOT NULL, description TEXT DEFAULT '',
            module TEXT, priority TEXT DEFAULT 'medium',
            recurrence TEXT DEFAULT 'annual',  -- annual | quarterly | monthly
            due_month INTEGER, due_day INTEGER, -- for annual: month & day; quarterly: day of quarter
            estimated_hrs NUMERIC(6,2) DEFAULT 0,
            auto_assign_role TEXT DEFAULT 'staff',
            is_active INTEGER DEFAULT 1,
            tenant_id TEXT REFERENCES tenants(id), created_by TEXT,
            created_at TIMESTAMPTZ DEFAULT NOW()
        )""")
    else:
        c.execute("""CREATE TABLE IF NOT EXISTS recurring_task_templates (
            id TEXT PRIMARY KEY, title TEXT NOT NULL, description TEXT DEFAULT '',
            module TEXT, priority TEXT DEFAULT 'medium',
            recurrence TEXT DEFAULT 'annual',
            due_month INTEGER, due_day INTEGER,
            estimated_hrs REAL DEFAULT 0, auto_assign_role TEXT DEFAULT 'staff',
            is_active INTEGER DEFAULT 1, tenant_id TEXT, created_by TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )""")
    conn.commit(); conn.close()

@app.route("/api/recurring-templates")
@login_required
def list_recurring_templates():
    _ensure_recurring_templates_table()
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM recurring_task_templates WHERE tenant_id=%s ORDER BY title",
              (g.tenant_id,))
    result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/recurring-templates", methods=["POST"])
@require_role("superadmin","manager")
def create_recurring_template():
    _ensure_recurring_templates_table()
    d = request.get_json(silent=True) or {}
    if not d.get("title"): return jsonify({"error": "title required"}), 400
    tid = str(uuid.uuid4())
    conn = get_db(); c = conn.cursor()
    c.execute("""INSERT INTO recurring_task_templates
                 (id,title,description,module,priority,recurrence,due_month,due_day,
                  estimated_hrs,auto_assign_role,is_active,tenant_id,created_by)
                 VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,1,%s,%s)""",
              (tid, d["title"], d.get("description",""), d.get("module"),
               d.get("priority","medium"), d.get("recurrence","annual"),
               d.get("due_month"), d.get("due_day"),
               float(d.get("estimated_hrs") or 0), d.get("auto_assign_role","staff"),
               g.tenant_id, g.user_id))
    conn.commit()
    c.execute("SELECT * FROM recurring_task_templates WHERE id=%s", (tid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result), 201

@app.route("/api/recurring-templates/<tid>", methods=["PUT"])
@require_role("superadmin","manager")
def update_recurring_template(tid):
    _ensure_recurring_templates_table()
    d = request.get_json(silent=True) or {}
    allowed = ["title","description","module","priority","recurrence","due_month",
               "due_day","estimated_hrs","auto_assign_role","is_active"]
    fields = {k: d[k] for k in allowed if k in d}
    if not fields: return jsonify({"error": "Nothing to update"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute(f"UPDATE recurring_task_templates SET {','.join(k+'=%s' for k in fields)} WHERE id=%s AND tenant_id=%s",
              list(fields.values()) + [tid, g.tenant_id])
    conn.commit()
    c.execute("SELECT * FROM recurring_task_templates WHERE id=%s", (tid,))
    result = row(c.fetchone()); conn.close()
    return jsonify(result)

@app.route("/api/recurring-templates/<tid>", methods=["DELETE"])
@require_role("superadmin")
def delete_recurring_template(tid):
    _ensure_recurring_templates_table()
    conn = get_db(); c = conn.cursor()
    c.execute("DELETE FROM recurring_task_templates WHERE id=%s AND tenant_id=%s", (tid, g.tenant_id))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/recurring-templates/<tid>/spawn", methods=["POST"])
@require_role("superadmin","manager")
def spawn_tasks_from_template(tid):
    """
    Create tasks for active companies from this template.

    Body (JSON):
      due_date          : str  — shared due date for all companies (optional)
      task_leader       : str  — global fallback leader user id
      task_manager      : str  — global fallback manager user id
      assigned_to       : str  — global fallback assignee user id
      company_assignments: list of {
            company_id  : str,
            assigned_to : str | null,
            task_manager: str | null,
            task_leader : str | null,
            due_date    : str | null   (overrides global due_date for this co)
        }
    Per-company assignments take priority over the global fallbacks.
    """
    _ensure_recurring_templates_table()
    d = request.get_json(silent=True) or {}
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM recurring_task_templates WHERE id=%s AND tenant_id=%s", (tid, g.tenant_id))
    tmpl = row(c.fetchone())
    if not tmpl: conn.close(); return jsonify({"error": "Template not found"}), 404

    c.execute("SELECT id, name FROM companies WHERE tenant_id=%s AND status='active'", (g.tenant_id,))
    all_active = {(r["id"] if isinstance(r,dict) else r[0]): (r["name"] if isinstance(r,dict) else r[1])
                  for r in c.fetchall()}

    # Build per-company assignment map from request
    co_map = {}  # company_id -> {assigned_to, task_manager, task_leader, due_date}
    for entry in (d.get("company_assignments") or []):
        cid = entry.get("company_id")
        if cid: co_map[cid] = entry

    # Global fallbacks
    global_due = d.get("due_date") or (
        f"{date.today().year}-{str(tmpl.get('due_month',12)).zfill(2)}-{str(tmpl.get('due_day',31)).zfill(2)}"
        if tmpl.get("due_month") else None
    )
    global_leader  = d.get("task_leader")  or None
    global_manager = d.get("task_manager") or None
    global_assign  = d.get("assigned_to")  or None

    created = 0; skipped = []
    for co_id in all_active:
        ov = co_map.get(co_id, {})
        # Skip only if explicitly excluded (future feature) — for now create for all
        task_id = str(uuid.uuid4())
        c.execute("""INSERT INTO tasks (id,company_id,title,description,priority,status,
                     due_date,module,estimated_hrs,created_by,tenant_id,
                     assigned_to,task_manager,task_leader)
                     VALUES (%s,%s,%s,%s,%s,'pending',%s,%s,%s,%s,%s,%s,%s,%s)""",
                  (task_id, co_id, tmpl["title"], tmpl.get("description",""),
                   tmpl.get("priority","medium"),
                   _dt(ov.get("due_date") or global_due),
                   tmpl.get("module"),
                   float(tmpl.get("estimated_hrs") or 0),
                   g.user_id, g.tenant_id,
                   ov.get("assigned_to")  or global_assign  or None,
                   ov.get("task_manager") or global_manager or None,
                   ov.get("task_leader")  or global_leader  or None))
        created += 1

    conn.commit(); conn.close()
    return jsonify({"success": True, "tasks_created": created, "skipped": skipped})


# ── P1 #13 — Session Management ──────────────────────────────────────────────
def _ensure_sessions_table():
    conn = get_db(); c = conn.cursor()
    if USE_POSTGRES:
        c.execute("""CREATE TABLE IF NOT EXISTS user_sessions (
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
            token_hash TEXT NOT NULL UNIQUE,
            device_info TEXT DEFAULT '', ip_address TEXT DEFAULT '',
            user_agent TEXT DEFAULT '',
            created_at TIMESTAMPTZ DEFAULT NOW(),
            last_seen TIMESTAMPTZ DEFAULT NOW(),
            expires_at TIMESTAMPTZ,
            is_active INTEGER DEFAULT 1
        )""")
        c.execute("CREATE INDEX IF NOT EXISTS idx_sessions_user ON user_sessions(user_id)")
        c.execute("CREATE INDEX IF NOT EXISTS idx_sessions_hash ON user_sessions(token_hash)")
    else:
        c.execute("""CREATE TABLE IF NOT EXISTS user_sessions (
            id TEXT PRIMARY KEY, user_id TEXT NOT NULL,
            token_hash TEXT NOT NULL UNIQUE,
            device_info TEXT DEFAULT '', ip_address TEXT DEFAULT '',
            user_agent TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now')),
            last_seen TEXT DEFAULT (datetime('now')),
            expires_at TEXT, is_active INTEGER DEFAULT 1
        )""")
    conn.commit(); conn.close()

import hashlib as _hashlib2

def _record_session(user_id: str, token: str, ip: str, ua: str):
    """Log a new session for the user after login."""
    try:
        _ensure_sessions_table()
        token_hash = _hashlib2.sha256(token.encode()).hexdigest()
        sid = str(uuid.uuid4())
        expires = datetime.utcnow() + timedelta(hours=int(os.environ.get("TOKEN_EXPIRY_HOURS","12")))
        conn = get_db(); c = conn.cursor()
        # Parse a friendly device string
        device = "Unknown"
        for kw in ["Mobile","Android","iPhone","iPad","Windows","Mac","Linux","Chrome","Firefox","Safari"]:
            if kw.lower() in ua.lower():
                device = kw; break
        c.execute("""INSERT INTO user_sessions (id,user_id,token_hash,device_info,ip_address,user_agent,expires_at)
                     VALUES (%s,%s,%s,%s,%s,%s,%s)""",
                  (sid, user_id, token_hash, device, ip, ua[:200], expires.isoformat()))
        conn.commit(); conn.close()
    except Exception as e:
        app.logger.debug(f"session record failed: {e}")

@app.route("/api/auth/sessions")
@login_required
def list_sessions():
    _ensure_sessions_table()
    conn = get_db(); c = conn.cursor()
    c.execute("""SELECT id, device_info, ip_address, created_at, last_seen, expires_at, is_active
                 FROM user_sessions WHERE user_id=%s ORDER BY last_seen DESC""", (g.user_id,))
    result = rows(c.fetchall()); conn.close()
    return jsonify(result)

@app.route("/api/auth/sessions/<sid>/revoke", methods=["POST"])
@login_required
def revoke_session(sid):
    _ensure_sessions_table()
    conn = get_db(); c = conn.cursor()
    c.execute("UPDATE user_sessions SET is_active=0 WHERE id=%s AND user_id=%s", (sid, g.user_id))
    conn.commit(); conn.close()
    return jsonify({"success": True})

@app.route("/api/auth/sessions/revoke-all", methods=["POST"])
@login_required
def revoke_all_sessions():
    _ensure_sessions_table()
    conn = get_db(); c = conn.cursor()
    c.execute("UPDATE user_sessions SET is_active=0 WHERE user_id=%s", (g.user_id,))
    conn.commit(); conn.close()
    return jsonify({"success": True})


# ── P1 #14 — JWT → httpOnly Cookie ───────────────────────────────────────────
_COOKIE_NAME   = "cms_token"
_COOKIE_SECURE = os.environ.get("COOKIE_SECURE", "0") == "1"  # set "1" in prod (HTTPS)

@app.route("/api/auth/login-cookie", methods=["POST"])
def login_cookie():
    """
    Same as /api/auth/login but sets token in httpOnly cookie instead of JSON body.
    Prevents XSS token theft. Frontend should call this endpoint and omit Authorization header.
    """
    _ip = request.headers.get("X-Forwarded-For", request.remote_addr or "?").split(",")[0].strip()
    if not rate_limit_login(_ip):
        return jsonify({"error": "Too many login attempts — try again in 1 minute"}), 429
    d = request.get_json(silent=True, force=True) or {}
    email = (d.get("email") or "").strip().lower()
    pw    = d.get("password") or ""
    if not email or not pw: return jsonify({"error": "Email and password required"}), 400
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM users WHERE email=%s AND is_active=1", (email,))
    user = row(c.fetchone())
    if not user or not verify_pw(pw, user["password"]):
        conn.close(); return jsonify({"error": "Invalid credentials"}), 401
    # Auto-upgrade legacy SHA-256 hash to bcrypt
    stored = user["password"]
    if not (stored.startswith("$2b$") or stored.startswith("$2a$")):
        try:
            c.execute("UPDATE users SET password=%s WHERE id=%s", (hash_pw(pw), user["id"]))
            conn.commit()
        except Exception: pass
    is_pa = bool(user.get("is_platform_admin", 0))
    token = make_token(user["id"], user["role"], user["name"], user.get("tenant_id"), is_pa)
    c.execute("UPDATE users SET last_login=NOW() WHERE id=%s", (user["id"],))
    conn.commit(); conn.close()
    # Record session
    ua = request.headers.get("User-Agent", "")
    _record_session(user["id"], token, _ip, ua)
    resp = jsonify({
        "user": {"id": user["id"], "name": user["name"], "email": user["email"], "role": user["role"]},
        "is_platform_admin": is_pa, "tenant_id": user.get("tenant_id"),
    })
    resp.set_cookie(
        _COOKIE_NAME, token,
        httponly=True, secure=_COOKIE_SECURE, samesite="Lax",
        max_age=int(os.environ.get("TOKEN_EXPIRY_HOURS","12")) * 3600,
        path="/"
    )
    return resp

@app.route("/api/auth/logout-cookie", methods=["POST"])
def logout_cookie():
    resp = jsonify({"success": True})
    resp.delete_cookie(_COOKIE_NAME, path="/")
    return resp


# ── P1 #6 — Bulk Operations (select + assign/delete) ─────────────────────────
@app.route("/api/tasks/bulk", methods=["POST"])
@login_required
def bulk_tasks():
    """Bulk assign or update status for multiple tasks."""
    d = request.get_json(silent=True) or {}
    ids    = d.get("ids", [])
    action = d.get("action", "")  # assign | status | delete
    if not ids or not action:
        return jsonify({"error": "ids and action required"}), 400
    conn = get_db(); c = conn.cursor()
    affected = 0
    if action == "assign":
        assignee = d.get("assigned_to")
        manager  = d.get("task_manager")
        leader   = d.get("task_leader")
        for tid in ids:
            fields = {}
            if assignee is not None: fields["assigned_to"]  = assignee or None
            if manager  is not None: fields["task_manager"] = manager  or None
            if leader   is not None: fields["task_leader"]  = leader   or None
            if fields:
                c.execute(f"UPDATE tasks SET {','.join(k+'=%s' for k in fields)} WHERE id=%s AND tenant_id=%s",
                          list(fields.values()) + [tid, g.tenant_id])
                affected += c.rowcount if hasattr(c, 'rowcount') else 1
    elif action == "status":
        new_status = d.get("status", "pending")
        for tid in ids:
            extra = ",completed_at=NOW()" if new_status == "completed" else ""
            c.execute(f"UPDATE tasks SET status=%s{extra} WHERE id=%s AND tenant_id=%s",
                      (new_status, tid, g.tenant_id))
            affected += 1
    elif action == "delete" and g.role == "superadmin":
        for tid in ids:
            c.execute("DELETE FROM tasks WHERE id=%s AND tenant_id=%s", (tid, g.tenant_id))
            affected += 1
    else:
        conn.close(); return jsonify({"error": f"Unknown action or insufficient role: {action}"}), 400
    conn.commit(); conn.close()
    return jsonify({"success": True, "affected": affected})


# ── P1 #2 — Dashboard analytics data endpoint ────────────────────────────────
@app.route("/api/analytics")
@login_required
def analytics():
    """Returns chart-ready analytics data for the dashboard."""
    conn = get_db(); c = conn.cursor()
    today = date.today()
    tid = g.tenant_id

    # Tasks by status (pie)
    c.execute("""SELECT status, COUNT(*) AS cnt FROM tasks WHERE tenant_id=%s
                 GROUP BY status""", (tid,))
    tasks_by_status = {r["status"] if isinstance(r,dict) else r[0]:
                       int(r["cnt"] if isinstance(r,dict) else r[1]) for r in c.fetchall()}

    # Tasks by priority
    c.execute("""SELECT priority, COUNT(*) AS cnt FROM tasks WHERE tenant_id=%s
                 AND status NOT IN ('completed','cancelled') GROUP BY priority""", (tid,))
    tasks_by_priority = {r["priority"] if isinstance(r,dict) else r[0]:
                         int(r["cnt"] if isinstance(r,dict) else r[1]) for r in c.fetchall()}

    # Tasks completed per month (last 6 months)
    months = []
    for i in range(5, -1, -1):
        m = (today.replace(day=1) - timedelta(days=i*30))
        months.append(m.strftime("%Y-%m"))
    monthly_completed = {}
    for m in months:
        y, mo = m.split("-")
        if USE_POSTGRES:
            c.execute("""SELECT COUNT(*) FROM tasks WHERE tenant_id=%s AND status='completed'
                         AND DATE_TRUNC('month', completed_at) = make_date(%s::int, %s::int, 1)""",
                      (tid, int(y), int(mo)))
        else:
            c.execute("""SELECT COUNT(*) FROM tasks WHERE tenant_id=%s AND status='completed'
                         AND strftime('%Y-%m', completed_at) = %s""", (tid, m))
        r = c.fetchone()
        monthly_completed[m] = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0

    # Alerts by severity
    c.execute("""SELECT severity, COUNT(*) AS cnt FROM alerts WHERE tenant_id=%s
                 AND status='active' GROUP BY severity""", (tid,))
    alerts_by_sev = {r["severity"] if isinstance(r,dict) else r[0]:
                     int(r["cnt"] if isinstance(r,dict) else r[1]) for r in c.fetchall()}

    # Company compliance health scores
    c.execute("SELECT id, name FROM companies WHERE tenant_id=%s AND status='active' LIMIT 10", (tid,))
    companies_list = rows(c.fetchall())
    conn.close()

    return jsonify({
        "tasks_by_status":   tasks_by_status,
        "tasks_by_priority": tasks_by_priority,
        "monthly_completed": {"labels": months, "data": [monthly_completed[m] for m in months]},
        "alerts_by_severity": alerts_by_sev,
        "companies": [{"id": co["id"], "name": co["name"]} for co in companies_list],
    })


# ── P1 #3 — Global Search ─────────────────────────────────────────────────────
@app.route("/api/search")
@login_required
def global_search():
    q = (request.args.get("q") or "").strip()
    if len(q) < 2: return jsonify({"results": []})
    like = f"%{q}%"
    tid  = g.tenant_id
    conn = get_db(); c = conn.cursor()
    results = []

    # Companies
    c.execute("SELECT id,'company' AS type,name AS title,cin AS subtitle FROM companies WHERE tenant_id=%s AND (name ILIKE %s OR cin ILIKE %s) LIMIT 5", (tid, like, like))
    results += rows(c.fetchall())

    # Directors
    c.execute("SELECT id,'director' AS type,name AS title,din AS subtitle FROM directors WHERE tenant_id=%s AND (name ILIKE %s OR din ILIKE %s) LIMIT 5", (tid, like, like))
    results += rows(c.fetchall())

    # Tasks
    c.execute("SELECT id,'task' AS type,title,status AS subtitle FROM tasks WHERE tenant_id=%s AND title ILIKE %s LIMIT 5", (tid, like))
    results += rows(c.fetchall())

    # Documents
    c.execute("SELECT id,'document' AS type,doc_name AS title,doc_type AS subtitle FROM documents WHERE tenant_id=%s AND doc_name ILIKE %s LIMIT 5", (tid, like))
    results += rows(c.fetchall())

    # Alerts
    c.execute("SELECT id,'alert' AS type,title,severity AS subtitle FROM alerts WHERE tenant_id=%s AND status='active' AND title ILIKE %s LIMIT 5", (tid, like))
    results += rows(c.fetchall())

    conn.close()
    return jsonify({"query": q, "results": results[:20]})


# ── P1 #12 — AGM Countdown & Annual Timeline ─────────────────────────────────
@app.route("/api/companies/<cid>/annual-timeline")
@login_required
def annual_timeline(cid):
    """
    Returns the full annual compliance calendar for a company.
    Covers: Board Meetings, AGM, MCA filings, GST (GSTR-1 & GSTR-3B monthly),
    MSME Form-1, TDS Returns (quarterly), Tax Audit, Income Tax Return,
    DIR-3 KYC, DPT-3.

    Indian Financial Year: April 1 → March 31.
    All dates computed relative to the CURRENT running FY so they always
    show the next upcoming deadline, never one that's already > 12 months past.
    """
    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s AND tenant_id=%s", (cid, g.tenant_id))
    co = row(c.fetchone())
    if not co: conn.close(); return jsonify({"error": "Company not found"}), 404
    conn.close()

    today = date.today()
    # Indian FY starts April 1. If today is Jan–Mar we're in the tail of last FY.
    fy_start_year  = today.year if today.month >= 4 else today.year - 1
    fy_end_year    = fy_start_year + 1            # e.g. 2026 → 2027
    fy_label       = f"{fy_start_year}-{str(fy_end_year)[-2:]}"
    cy             = today.year                   # calendar year shorthand

    events = []

    def _add(ev_name, ev_date, ev_type, ev_icon, category=""):
        """Append one event, rolling to next occurrence if already > 90 days past."""
        days = (ev_date - today).days
        events.append({
            "event":    ev_name,
            "date":     ev_date.isoformat(),
            "days_left": days,
            "type":     ev_type,
            "icon":     ev_icon,
            "category": category,
        })

    def _next_month_date(mo, dy, roll_year=None):
        """Return date(year, mo, dy) advancing year if that date is already > 60 days past."""
        y = roll_year or cy
        try:
            d_ = date(y, mo, dy)
            if (d_ - today).days < -60:
                d_ = date(y + 1, mo, dy)
            return d_
        except ValueError:
            # Handle month-end overflow (e.g. Feb 30 → Feb 28)
            import calendar
            last = calendar.monthrange(y, mo)[1]
            return date(y, mo, min(dy, last))

    # ─────────────────────────────────────────────────────────────────────────
    # 1. BOARD MEETINGS (quarterly — due within 45 days of quarter end)
    #    Q1: Apr–Jun  → by Aug 14  | Q2: Jul–Sep  → by Nov 14
    #    Q3: Oct–Dec  → by Feb 14  | Q4: Jan–Mar  → by May 14
    # ─────────────────────────────────────────────────────────────────────────
    board_qtrs = [
        (f"Q1 Board Meeting (Apr–Jun)", date(fy_start_year, 8, 14)),
        (f"Q2 Board Meeting (Jul–Sep)", date(fy_start_year, 11, 14)),
        (f"Q3 Board Meeting (Oct–Dec)", date(fy_end_year,   2, 14)),
        (f"Q4 Board Meeting (Jan–Mar)", date(fy_end_year,   5, 14)),
    ]
    for label, d_ in board_qtrs:
        _add(label, d_, "meeting", "🗓️", "MCA")

    # ─────────────────────────────────────────────────────────────────────────
    # 2. AGM — within 6 months of FY end (Sep 30)
    # ─────────────────────────────────────────────────────────────────────────
    _add("AGM Deadline", date(fy_end_year, 9, 30), "agm", "🏛️", "MCA")

    # ─────────────────────────────────────────────────────────────────────────
    # 3. MCA ANNUAL FILINGS
    # ─────────────────────────────────────────────────────────────────────────
    _add("AOC-4 Filing (Annual Accounts)",  date(fy_end_year, 10, 29), "filing", "📄", "MCA")
    _add("MGT-7 / MGT-7A (Annual Return)", date(fy_end_year, 11, 29), "filing", "📋", "MCA")
    _add("DPT-3 (Deposits Return)",         date(fy_end_year,  6, 30), "filing", "📋", "MCA")
    _add("DIR-3 KYC (Director KYC)",        date(fy_end_year,  9, 30), "kyc",    "🪪", "MCA")
    _add("BEN-2 (Beneficial Ownership)",    date(fy_end_year,  4, 30), "filing", "📋", "MCA")

    # ─────────────────────────────────────────────────────────────────────────
    # 4. GST — GSTR-3B Monthly (20th of next month for regular taxpayers)
    #    and GSTR-1 Monthly (11th of next month)
    #    Show next 6 months of each
    # ─────────────────────────────────────────────────────────────────────────
    gst_months = []
    for i in range(6):
        # compute the period month = today's month + i
        base = date(today.year, today.month, 1)
        # advance i months
        mo = today.month + i
        yr = today.year + (mo - 1) // 12
        mo = ((mo - 1) % 12) + 1
        period = date(yr, mo, 1)

        # GSTR-1 due: 11th of month following period
        gstr1_mo = mo + 1; gstr1_yr = yr
        if gstr1_mo > 12: gstr1_mo -= 12; gstr1_yr += 1
        gstr1_due = date(gstr1_yr, gstr1_mo, 11)

        # GSTR-3B due: 20th of month following period
        gstr3b_mo = mo + 1; gstr3b_yr = yr
        if gstr3b_mo > 12: gstr3b_mo -= 12; gstr3b_yr += 1
        gstr3b_due = date(gstr3b_yr, gstr3b_mo, 20)

        period_str = period.strftime("%b %Y")
        if (gstr1_due - today).days > -15:   # only show if not too far past
            _add(f"GSTR-1 — {period_str}", gstr1_due,  "gst", "📊", "GST")
        if (gstr3b_due - today).days > -15:
            _add(f"GSTR-3B — {period_str}", gstr3b_due, "gst", "📊", "GST")

    # GSTR-9 (Annual GST Return) — Dec 31
    _add("GSTR-9 (Annual GST Return)",    date(fy_end_year, 12, 31), "gst", "📊", "GST")
    _add("GSTR-9C (GST Reconciliation)",  date(fy_end_year, 12, 31), "gst", "📊", "GST")

    # ─────────────────────────────────────────────────────────────────────────
    # 5. TDS RETURNS (quarterly — Form 24Q / 26Q / 27Q / 27EQ)
    #    Q1 Apr–Jun → Jul 31  | Q2 Jul–Sep → Oct 31
    #    Q3 Oct–Dec → Jan 31  | Q4 Jan–Mar → May 31
    # ─────────────────────────────────────────────────────────────────────────
    tds_qtrs = [
        ("TDS Return Q1 (Apr–Jun) — 26Q/24Q", date(fy_start_year, 7,  31)),
        ("TDS Return Q2 (Jul–Sep) — 26Q/24Q", date(fy_start_year, 10, 31)),
        ("TDS Return Q3 (Oct–Dec) — 26Q/24Q", date(fy_end_year,   1,  31)),
        ("TDS Return Q4 (Jan–Mar) — 26Q/24Q", date(fy_end_year,   5,  31)),
    ]
    for label, d_ in tds_qtrs:
        _add(label, d_, "tds", "💰", "TDS")

    # TDS Certificate issuance deadlines (Form 16 / 16A)
    _add("Form 16 Issue (Salary TDS Cert)",   date(fy_end_year, 6, 15), "tds", "💰", "TDS")
    _add("Form 16A Issue (Non-Salary Cert) Q1", date(fy_start_year, 8, 15), "tds", "💰", "TDS")
    _add("Form 16A Issue (Non-Salary Cert) Q2", date(fy_start_year, 11, 15), "tds", "💰", "TDS")
    _add("Form 16A Issue (Non-Salary Cert) Q3", date(fy_end_year, 2, 15), "tds", "💰", "TDS")

    # ─────────────────────────────────────────────────────────────────────────
    # 6. INCOME TAX
    # ─────────────────────────────────────────────────────────────────────────
    _add("Tax Audit (Form 3CA/3CB/3CD)",  date(fy_end_year, 9,  30), "tax", "🏛️", "Income Tax")
    _add("Income Tax Return (Companies)", date(fy_end_year, 10, 31), "tax", "🏛️", "Income Tax")
    _add("ITR with Transfer Pricing",     date(fy_end_year, 11, 30), "tax", "🏛️", "Income Tax")
    # Advance Tax instalments
    _add("Advance Tax — 1st Instalment (15%)",  date(fy_start_year, 6, 15), "tax", "💸", "Income Tax")
    _add("Advance Tax — 2nd Instalment (45%)",  date(fy_start_year, 9, 15), "tax", "💸", "Income Tax")
    _add("Advance Tax — 3rd Instalment (75%)",  date(fy_start_year, 12, 15), "tax", "💸", "Income Tax")
    _add("Advance Tax — 4th Instalment (100%)", date(fy_end_year,   3, 15), "tax", "💸", "Income Tax")

    # ─────────────────────────────────────────────────────────────────────────
    # 7. MSME FILINGS
    #    MSME Form-1: Half-yearly — Apr 30 (Oct–Mar period) & Oct 31 (Apr–Sep period)
    #    MSME Udyam Registration renewal: annually
    # ─────────────────────────────────────────────────────────────────────────
    _add("MSME Form-1 (Apr–Sep period)",      date(fy_start_year, 10, 31), "msme", "🏭", "MSME")
    _add("MSME Form-1 (Oct–Mar period)",      date(fy_end_year,    4, 30), "msme", "🏭", "MSME")
    _add("MSME Samadhaan (Payment delayed?)", date(fy_start_year, 10, 15), "msme", "🏭", "MSME")

    # ─────────────────────────────────────────────────────────────────────────
    # 8. ESI / PF (monthly — 15th)
    # ─────────────────────────────────────────────────────────────────────────
    for i in range(3):   # next 3 months
        mo = today.month + i + 1
        yr = today.year + (mo - 1) // 12
        mo = ((mo - 1) % 12) + 1
        d_ = date(yr, mo, 15)
        period = d_.strftime("%b %Y")
        if (d_ - today).days >= -5:
            _add(f"PF/ESI Challan — {period}", d_, "payroll", "👷", "Payroll")

    # ─────────────────────────────────────────────────────────────────────────
    # Filter: remove events > 13 months out and deduplicate
    # Sort chronologically
    # ─────────────────────────────────────────────────────────────────────────
    seen = set()
    filtered = []
    for ev in events:
        key = (ev["event"], ev["date"])
        if key not in seen and ev["days_left"] >= -30 and ev["days_left"] <= 395:
            seen.add(key)
            filtered.append(ev)
    filtered.sort(key=lambda e: e["date"])

    return jsonify({
        "company_id":    cid,
        "company_name":  co["name"],
        "financial_year": fy_label,
        "events":        filtered,
        "categories":    ["MCA", "GST", "TDS", "Income Tax", "MSME", "Payroll"],
    })


# ── P1 #12b — Annual Timeline PDF download ───────────────────────────────────
@app.route("/api/companies/<cid>/annual-timeline-pdf")
@login_required
def annual_timeline_pdf(cid):
    """Board-ready Annual Compliance Timeline PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch, mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import io as _io

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s AND tenant_id=%s", (cid, g.tenant_id))
    co = row(c.fetchone())
    if not co: conn.close(); return jsonify({"error": "Company not found"}), 404
    conn.close()

    today = date.today()
    year  = today.year

    # ── Build events: call the JSON route logic directly ─────────────────────
    fy_start_year = today.year if today.month >= 4 else today.year - 1
    fy_end_year   = fy_start_year + 1
    fy_label      = f"{fy_start_year}-{str(fy_end_year)[-2:]}"

    def _ev_pdf(name, ev_date, ev_type, category=""):
        events.append({"event": name, "date": ev_date,
                       "days_left": (ev_date - today).days,
                       "type": ev_type, "category": category})

    events = []
    board = [("Q1 Board Meeting (Apr–Jun)", date(fy_start_year, 8,14)),
             ("Q2 Board Meeting (Jul–Sep)", date(fy_start_year,11,14)),
             ("Q3 Board Meeting (Oct–Dec)", date(fy_end_year,   2,14)),
             ("Q4 Board Meeting (Jan–Mar)", date(fy_end_year,   5,14))]
    for nm, dt in board:
        _ev_pdf(nm, dt, "meeting", "MCA")

    # MCA filings
    _ev_pdf("AGM Deadline",                  date(fy_end_year, 9,30),  "agm",     "MCA")
    _ev_pdf("AOC-4 (Annual Accounts)",       date(fy_end_year,10,29),  "filing",  "MCA")
    _ev_pdf("MGT-7 / MGT-7A (Annual Return)",date(fy_end_year,11,29),  "filing",  "MCA")
    _ev_pdf("DPT-3 (Deposits Return)",       date(fy_end_year, 6,30),  "filing",  "MCA")
    _ev_pdf("DIR-3 KYC",                     date(fy_end_year, 9,30),  "kyc",     "MCA")
    _ev_pdf("BEN-2 (Beneficial Ownership)",  date(fy_end_year, 4,30),  "filing",  "MCA")

    # GST — next 6 monthly GSTR-1 and GSTR-3B
    for i in range(6):
        mo = today.month + i;  yr = today.year + (mo-1)//12;  mo = ((mo-1)%12)+1
        g1mo = mo+1; g1yr = yr;  g3mo = mo+1; g3yr = yr
        if g1mo>12: g1mo-=12; g1yr+=1
        if g3mo>12: g3mo-=12; g3yr+=1
        ps = date(yr,mo,1).strftime("%b %Y")
        g1d = date(g1yr,g1mo,11); g3d = date(g3yr,g3mo,20)
        if (g1d-today).days > -15: _ev_pdf(f"GSTR-1 — {ps}", g1d,  "gst","GST")
        if (g3d-today).days > -15: _ev_pdf(f"GSTR-3B — {ps}",g3d,  "gst","GST")
    _ev_pdf("GSTR-9 (Annual GST Return)",    date(fy_end_year,12,31), "gst","GST")
    _ev_pdf("GSTR-9C (GST Reconciliation)",  date(fy_end_year,12,31), "gst","GST")

    # TDS Returns
    tds = [("TDS Return Q1 (Apr–Jun)",date(fy_start_year,7,31)),
           ("TDS Return Q2 (Jul–Sep)",date(fy_start_year,10,31)),
           ("TDS Return Q3 (Oct–Dec)",date(fy_end_year,  1,31)),
           ("TDS Return Q4 (Jan–Mar)",date(fy_end_year,  5,31))]
    for nm,dt in tds: _ev_pdf(nm, dt, "tds","TDS")
    _ev_pdf("Form 16 Issue Deadline",    date(fy_end_year,6,15),     "tds","TDS")

    # Income Tax
    _ev_pdf("Tax Audit (Form 3CA/3CB)",  date(fy_end_year,9,30),     "tax","Income Tax")
    _ev_pdf("Income Tax Return",         date(fy_end_year,10,31),    "tax","Income Tax")
    _ev_pdf("Advance Tax — 1st (15%)",   date(fy_start_year,6,15),   "tax","Income Tax")
    _ev_pdf("Advance Tax — 2nd (45%)",   date(fy_start_year,9,15),   "tax","Income Tax")
    _ev_pdf("Advance Tax — 3rd (75%)",   date(fy_start_year,12,15),  "tax","Income Tax")
    _ev_pdf("Advance Tax — 4th (100%)",  date(fy_end_year,3,15),     "tax","Income Tax")

    # MSME
    _ev_pdf("MSME Form-1 (Apr–Sep)",     date(fy_start_year,10,31),  "msme","MSME")
    _ev_pdf("MSME Form-1 (Oct–Mar)",     date(fy_end_year,4,30),     "msme","MSME")

    # Filter & sort
    seen = set()
    events_clean = []
    for ev in events:
        k = (ev["event"], ev["date"].isoformat() if hasattr(ev["date"],"isoformat") else str(ev["date"]))
        if k not in seen and ev["days_left"] >= -30 and ev["days_left"] <= 395:
            seen.add(k); events_clean.append(ev)
    events_clean.sort(key=lambda e: e["date"])
    events = events_clean

    # ── PDF colours ──────────────────────────────────────────────────────────
    NAVY  = colors.HexColor("#0f2d5c")
    BLUE  = colors.HexColor("#1a56db")
    GREY  = colors.HexColor("#64748b")
    WHITE = colors.white
    RED   = colors.HexColor("#dc2626")
    AMBER = colors.HexColor("#f59e0b")
    GREEN = colors.HexColor("#16a34a")
    type_colors = {
        "meeting": colors.HexColor("#3b82f6"),
        "agm":     colors.HexColor("#7c3aed"),
        "kyc":     colors.HexColor("#0891b2"),
        "filing":  colors.HexColor("#2563eb"),
        "tax":     colors.HexColor("#d97706"),
        "gst":     colors.HexColor("#059669"),
        "tds":     colors.HexColor("#dc2626"),
        "msme":    colors.HexColor("#7c3aed"),
        "payroll": colors.HexColor("#0891b2"),
    }
    cat_colors = {
        "MCA":        colors.HexColor("#1a56db"),
        "GST":        colors.HexColor("#059669"),
        "TDS":        colors.HexColor("#dc2626"),
        "Income Tax": colors.HexColor("#d97706"),
        "MSME":       colors.HexColor("#7c3aed"),
        "Payroll":    colors.HexColor("#0891b2"),
    }

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=0.6*inch, bottomMargin=0.6*inch,
                            leftMargin=0.85*inch, rightMargin=0.85*inch)

    def sty(name, sz=9, bold=False, col=None, align=TA_LEFT):
        return ParagraphStyle(name, fontName="Helvetica-Bold" if bold else "Helvetica",
                               fontSize=sz, textColor=col or colors.HexColor("#1e293b"),
                               alignment=align, leading=sz*1.5, spaceAfter=0)

    story = []
    # Header
    story.append(Paragraph(co["name"].upper(), sty("h1", 15, True, NAVY, TA_CENTER)))
    story.append(Paragraph("ANNUAL COMPLIANCE CALENDAR", sty("sub", 11, True, BLUE, TA_CENTER)))
    story.append(Paragraph(
        f"Financial Year: {fy_label}  |  Generated: {today.strftime('%d %B %Y')}  |  CIN: {co.get('cin','—')}",
        sty("dt", 8, False, GREY, TA_CENTER)))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLUE, spaceAfter=14))

    # Summary counts
    overdue  = sum(1 for e in events if e["days_left"] < 0)
    upcoming = sum(1 for e in events if 0 <= e["days_left"] <= 30)
    total    = len(events)
    summary_rows = [
        [Paragraph("Metric", sty("sh", 9, True, WHITE)),
         Paragraph("Count", sty("sh", 9, True, WHITE)),
         Paragraph("Status", sty("sh", 9, True, WHITE))],
        [Paragraph("Total Compliance Events", sty("td", 9)),
         Paragraph(str(total), sty("v", 10, True, BLUE)),
         Paragraph(f"FY {fy_label}", sty("d", 9, col=GREY))],
        [Paragraph("Overdue / Missed", sty("td", 9)),
         Paragraph(str(overdue), sty("v", 10, True, RED if overdue else GREEN)),
         Paragraph("Requires immediate attention" if overdue else "None overdue", sty("d", 9, col=RED if overdue else GREEN))],
        [Paragraph("Due in Next 30 Days", sty("td", 9)),
         Paragraph(str(upcoming), sty("v", 10, True, AMBER if upcoming else GREEN)),
         Paragraph("Action required soon" if upcoming else "None imminent", sty("d", 9, col=AMBER if upcoming else GREEN))],
    ]
    sum_ts = TableStyle([
        ("BACKGROUND",(0,0),(-1,0),BLUE), ("TEXTCOLOR",(0,0),(-1,0),WHITE),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#cbd5e1")),
        ("ROWPADDING",(0,0),(-1,-1),7), ("FONTSIZE",(0,0),(-1,-1),9),
        ("BACKGROUND",(0,2),(-1,2),colors.HexColor("#f8fafc")),
    ])
    story.append(Table(summary_rows, colWidths=[doc.width*0.45, doc.width*0.15, doc.width*0.4],
                       style=sum_ts))
    story.append(Spacer(1, 14))

    # Events table
    story.append(Paragraph("Compliance Events — Chronological", sty("s2", 10, True, NAVY)))
    story.append(Spacer(1, 6))

    # Group events by category for PDF sections
    from collections import OrderedDict as _OD
    cat_order = ["MCA", "GST", "TDS", "Income Tax", "MSME", "Payroll"]
    grouped = _OD()
    for cat in cat_order: grouped[cat] = []
    for ev in events:
        cat = ev.get("category","Other")
        if cat not in grouped: grouped[cat] = []
        grouped[cat].append(ev)

    for cat, cat_events in grouped.items():
        if not cat_events: continue
        cat_col = cat_colors.get(cat, BLUE)
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"{cat} Deadlines", sty(f"cat_{cat}", 10, True, cat_col)))
        story.append(Spacer(1, 4))

        ev_rows = [[Paragraph(h, sty(f"ch_{cat}_{i}", 8, True, WHITE))
                    for i, h in enumerate(["Event / Filing","Due Date","Days Left","Status"])]]
        for j, ev in enumerate(cat_events):
            dl = ev["days_left"]
            if dl < 0:
                status_txt = f"OVERDUE {-dl}d"; status_col = RED
            elif dl == 0:
                status_txt = "TODAY!";           status_col = RED
            elif dl <= 30:
                status_txt = f"Due in {dl}d";   status_col = AMBER
            elif dl <= 90:
                status_txt = f"{dl} days";       status_col = colors.HexColor("#2563eb")
            else:
                status_txt = f"{dl} days";       status_col = GREEN

            ev_date = ev["date"]
            if hasattr(ev_date, "strftime"):
                date_str = ev_date.strftime("%d %b %Y")
            else:
                try: date_str = date.fromisoformat(str(ev_date)[:10]).strftime("%d %b %Y")
                except: date_str = str(ev_date)[:10]

            row_bg = colors.HexColor("#fef2f2") if dl < 0 else (
                     colors.HexColor("#fffbeb") if dl <= 30 else
                     (colors.HexColor("#f8fafc") if j%2==0 else WHITE))
            ev_rows.append([
                Paragraph(ev["event"], sty(f"en_{cat}_{j}", 9, True)),
                Paragraph(date_str,    sty(f"ed_{cat}_{j}", 9)),
                Paragraph(status_txt,  sty(f"el_{cat}_{j}", 9, True, status_col)),
                Paragraph("⚠ Overdue" if dl < 0 else ("⏳ Soon" if dl<=30 else "◷ Scheduled"),
                          sty(f"es_{cat}_{j}", 8, col=status_col)),
            ])

        ev_ts = TableStyle([
            ("BACKGROUND",(0,0),(-1,0),cat_col), ("TEXTCOLOR",(0,0),(-1,0),WHITE),
            ("GRID",(0,0),(-1,-1),0.25,colors.HexColor("#e2e8f0")),
            ("ROWPADDING",(0,0),(-1,-1),6), ("FONTSIZE",(0,0),(-1,-1),8.5),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ])
        for i in range(1, len(ev_rows)):
            ev_ts.add("BACKGROUND",(0,i),(-1,i),
                colors.HexColor("#fef2f2") if cat_events[i-1]["days_left"] < 0 else
                (colors.HexColor("#fffbeb") if cat_events[i-1]["days_left"]<=30 else
                 (colors.HexColor("#f8fafc") if i%2==0 else WHITE)))
        story.append(Table(ev_rows,
                           colWidths=[doc.width*0.42, doc.width*0.18,
                                       doc.width*0.18, doc.width*0.22],
                           style=ev_ts))
    story.append(Spacer(1, 14))

    # Footer
    story.append(HRFlowable(width="100%", thickness=0.5, color=GREY))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "CONFIDENTIAL — Generated by Taxly-CMS | This calendar is for planning purposes only. "
        "Verify actual due dates with MCA/IT notifications.",
        sty("foot", 7, False, GREY, TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    fname = f"{co['name'].replace(' ','_')}_Annual_Timeline_{fy_label.replace('-','_')}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)


# ── P1 #9 — Company Compliance Report PDF ─────────────────────────────────────
@app.route("/api/companies/<cid>/compliance-report-pdf")
@login_required
def compliance_report_pdf(cid):
    """Board-ready compliance report PDF for a company."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    import io as _io

    conn = get_db(); c = conn.cursor()
    c.execute("SELECT * FROM companies WHERE id=%s AND tenant_id=%s", (cid, g.tenant_id))
    co = row(c.fetchone())
    if not co: conn.close(); return jsonify({"error": "Company not found"}), 404

    today = date.today()

    # Fetch key compliance data
    c.execute("""SELECT d.name, d.din, k.kyc_status, k.next_due_date
                 FROM directors d LEFT JOIN director_kyc k ON d.id=k.director_id
                 WHERE d.company_id=%s AND d.is_active=1""", (cid,))
    directors_data = rows(c.fetchall())

    c.execute("""SELECT COUNT(*) FROM tasks WHERE company_id=%s
                 AND status NOT IN ('completed','cancelled')""", (cid,))
    r = c.fetchone(); open_tasks_count = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0

    c.execute("""SELECT COUNT(*) FROM tasks WHERE company_id=%s
                 AND status NOT IN ('completed','cancelled') AND due_date < %s""",
              (cid, today.isoformat()))
    r = c.fetchone(); overdue_count = int(list(r.values())[0] if isinstance(r, dict) else r[0]) if r else 0

    c.execute("""SELECT a.title, a.severity, a.due_date FROM alerts a
                 WHERE a.company_id=%s AND a.status='active' ORDER BY a.due_date LIMIT 10""", (cid,))
    active_alerts_data = rows(c.fetchall())

    c.execute("""SELECT meeting_type, meeting_date, status FROM meetings
                 WHERE company_id=%s AND meeting_date >= %s ORDER BY meeting_date LIMIT 5""",
              (cid, (today - timedelta(days=90)).isoformat()))
    meetings_data = rows(c.fetchall())

    c.execute("""SELECT holder_name, valid_to, dsc_class FROM dsc_records
                 WHERE company_id=%s AND is_active=1""", (cid,))
    dscs_data = rows(c.fetchall())
    conn.close()

    buf = _io.BytesIO()
    NAVY = colors.HexColor("#0f2d5c"); BLUE = colors.HexColor("#1a56db")
    GREY = colors.HexColor("#64748b"); WHITE = colors.white
    RED  = colors.HexColor("#d93535"); GREEN= colors.HexColor("#16a34a")

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=0.6*inch, bottomMargin=0.6*inch,
                            leftMargin=0.85*inch, rightMargin=0.85*inch)

    def sty(n, sz=9, bold=False, col=None, align=TA_LEFT):
        return ParagraphStyle(n, fontName="Helvetica-Bold" if bold else "Helvetica",
                               fontSize=sz, textColor=col or colors.HexColor("#1e293b"),
                               alignment=align, leading=sz*1.4)

    story = []
    story.append(Paragraph(co["name"].upper(), sty("h1", 16, True, NAVY, TA_CENTER)))
    story.append(Paragraph("COMPLIANCE STATUS REPORT", sty("sub", 11, True, BLUE, TA_CENTER)))
    story.append(Paragraph(f"Generated: {today.strftime('%d %B %Y')} | CIN: {co.get('cin','—')}",
                            sty("dt", 8, False, GREY, TA_CENTER)))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLUE, spaceAfter=12))

    # Summary scorecard
    overdue_col = RED if overdue_count > 0 else GREEN
    kyc_ok = sum(1 for d_ in directors_data if d_.get("kyc_status") == "filed")
    kyc_col = GREEN if kyc_ok == len(directors_data) else RED
    alert_col = RED if active_alerts_data else GREEN

    summary_data = [
        [Paragraph("Metric", sty("sh", 9, True, WHITE)),
         Paragraph("Status", sty("sh", 9, True, WHITE)),
         Paragraph("Detail", sty("sh", 9, True, WHITE))],
        [Paragraph("Open Tasks", sty("td", 9)),
         Paragraph(str(open_tasks_count), sty("v", 10, True)),
         Paragraph(f"{overdue_count} overdue", sty("d", 9, col=overdue_col))],
        [Paragraph("Director KYC", sty("td", 9)),
         Paragraph(f"{kyc_ok}/{len(directors_data)} filed", sty("v", 10, True, kyc_col)),
         Paragraph("DIR-3 KYC Status", sty("d", 9))],
        [Paragraph("Active Alerts", sty("td", 9)),
         Paragraph(str(len(active_alerts_data)), sty("v", 10, True, RED if active_alerts_data else GREEN)),
         Paragraph("Compliance Alerts", sty("d", 9))],
        [Paragraph("DSC Records", sty("td", 9)),
         Paragraph(str(len(dscs_data)), sty("v", 10, True)),
         Paragraph(f"{sum(1 for d_ in dscs_data if _pd(d_.get('valid_to')) and (_pd(d_.get('valid_to')) - today).days < 0)} expired",
                   sty("d", 9, col=RED))],
    ]
    def make_ts(data):
        """Build a fresh TableStyle for a table with len(data) rows."""
        _ts = TableStyle([
            ("BACKGROUND",(0,0),(-1,0),BLUE),
            ("TEXTCOLOR",(0,0),(-1,0),WHITE),
            ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#cbd5e1")),
            ("ROWPADDING",(0,0),(-1,-1),6),
            ("FONTSIZE",(0,0),(-1,-1),9),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ])
        for i in range(1, len(data)):
            if i % 2 == 0:
                _ts.add("BACKGROUND",(0,i),(-1,i),colors.HexColor("#f8fafc"))
        return _ts

    t = Table(summary_data, colWidths=[doc.width*0.3, doc.width*0.2, doc.width*0.5])
    t.setStyle(make_ts(summary_data)); story.append(t); story.append(Spacer(1, 12))

    # Alerts section
    if active_alerts_data:
        story.append(Paragraph("Active Compliance Alerts", sty("s2", 10, True, NAVY)))
        story.append(Spacer(1, 4))
        alert_rows = [[Paragraph(h, sty("ah"+str(i), 8, True, WHITE)) for i,h in enumerate(["Alert","Severity","Due Date"])]]
        for a in active_alerts_data:
            sc = RED if a.get("severity") == "critical" else colors.HexColor("#d97706")
            alert_rows.append([Paragraph(a.get("title","")[:80], sty("at"+str(len(alert_rows)), 8)),
                               Paragraph((a.get("severity") or "").upper(), sty("as"+str(len(alert_rows)), 8, True, sc)),
                               Paragraph(str(a.get("due_date","—"))[:10], sty("ad"+str(len(alert_rows)), 8))])
        at = Table(alert_rows, colWidths=[doc.width*0.55, doc.width*0.2, doc.width*0.25])
        at.setStyle(make_ts(alert_rows)); story.append(at); story.append(Spacer(1, 10))

    # Meetings
    if meetings_data:
        story.append(Paragraph("Recent / Upcoming Meetings", sty("s3", 10, True, NAVY)))
        story.append(Spacer(1, 4))
        m_rows = [[Paragraph(h, sty("mh"+str(i), 8, True, WHITE)) for i,h in enumerate(["Type","Date","Status"])]]
        for m_ in meetings_data:
            m_rows.append([Paragraph(m_.get("meeting_type",""), sty("mt"+str(len(m_rows)), 8)),
                           Paragraph(str(m_.get("meeting_date",""))[:10], sty("md"+str(len(m_rows)), 8)),
                           Paragraph(m_.get("status",""), sty("ms"+str(len(m_rows)), 8))])
        mt = Table(m_rows, colWidths=[doc.width*0.4, doc.width*0.3, doc.width*0.3])
        mt.setStyle(make_ts(m_rows)); story.append(mt); story.append(Spacer(1, 10))

    story.append(HRFlowable(width="100%", thickness=0.5, color=GREY))
    story.append(Spacer(1, 4))
    story.append(Paragraph("CONFIDENTIAL — Generated by Taxly-CMS | Taxly India Private Limited",
                            sty("foot", 7, False, GREY, TA_CENTER)))

    doc.build(story)
    buf.seek(0)
    fname = f"{co['name'].replace(' ','_')}_Compliance_{today.strftime('%Y%m%d')}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=fname)

